# Imports padrão do Django
from django.http import JsonResponse, HttpResponse, HttpResponseRedirect
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.models import User
from django.contrib.auth import authenticate, login as login_django, logout, update_session_auth_hash
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.urls import reverse
from django.conf import settings
from django.utils import timezone, dateparse
from django.views.decorators.cache import cache_page
from django.views.decorators.csrf import csrf_exempt, csrf_protect
from django.views.decorators.http import require_POST
from django.utils.timezone import now
from django.core.cache import cache

# Imports do Django relacionados a modelos e banco de dados
from django.db.models import Q, Sum, Count, Avg, F

# Imports de terceiros
import os
import pandas as pd
import shopify
import unidecode
import requests
import json
import random
import string
import threading
import mailtrap as mt
import cloudinary
from cloudinary.uploader import upload
from cloudinary.utils import cloudinary_url
import warnings
import re
import logging
import xml.etree.ElementTree as ET
from docx import Document

# Imports de tempo e datas
from datetime import datetime, timedelta
from itertools import chain
from operator import itemgetter
import time
from django.utils.dateparse import parse_date

# Imports de módulos internos do projeto

from shopbotfy.models import ZZTextosMensagem, Agencia, PerfilUsuario, Notificacao, OuvidoriaMensagem
from shop_api.models import AAFormuUserer, FilaDeEspera,  Pendencia
from .forms import CadastroAgenciaForm
from .tasks import processar_usuario
from .helpers import handle_produz_request

# Configuração de logging
logger = logging.getLogger(__name__)
task_status = {}


############################
############################################################
                                                #########################   ADMINISTRAÇÃO   
############################################################

@csrf_protect
@login_required
def administrador(request):
    template, created = ZZTextosMensagem.objects.get_or_create(id=1)
    agencias = Agencia.objects.all()  # Buscar todas as agências cadastradas
    print(agencias)
    if request.method == "POST":
        if "update" in request.POST:
            # Atualizar os campos de links
            template.link_1 = request.POST.get("link_1", template.link_1)
            template.link_2 = request.POST.get("link_2", template.link_2)
            template.link_3 = request.POST.get("link_3", template.link_3)
            template.link_4 = request.POST.get("link_4", template.link_4)
            template.link_5 = request.POST.get("link_5", template.link_5)
            template.link_6 = request.POST.get("link_6", template.link_6)
            template.link_7 = request.POST.get("link_7", template.link_7)
            template.link_8 = request.POST.get("link_8", template.link_8)
            template.show_link_8 = (
                "show_link_8" in request.POST
            )  # Checar se o checkbox foi marcado

            # Atualizar os campos de rodapé
            template.copy = request.POST.get("copy", template.copy)
            template.telefone = request.POST.get("telefone", template.telefone)
            template.show_telefone = (
                "show_telefone" in request.POST
            )  # Checar se o checkbox foi marcado

            template.contato = request.POST.get("contato", template.contato)
            template.show_contato = (
                "show_contato" in request.POST
            )  # Checar se o checkbox foi marcado

            template.link_site = request.POST.get("link_site", template.link_site)
            template.show_link_site = (
                "show_link_site" in request.POST
            )  # Checar se o checkbox foi marcado

            # Atualizar o link de login enviado no e-mail
            template.link_login = request.POST.get("link_login", template.link_login)

            # Atualizar campos de redes sociais e controle de visibilidade
            template.link_face = request.POST.get("link_face", template.link_face)
            template.show_link_face = (
                "show_link_face" in request.POST
            )  # Checar se o checkbox foi marcado
            template.link_insta = request.POST.get("link_insta", template.link_insta)
            template.show_link_insta = (
                "show_link_insta" in request.POST
            )  # Checar se o checkbox foi marcado
            template.link_youtube = request.POST.get(
                "link_youtube", template.link_youtube
            )
            template.show_link_youtube = (
                "show_link_youtube" in request.POST
            )  # Checar se o checkbox foi marcado
            template.link_x = request.POST.get("link_x", template.link_x)
            template.show_link_x = (
                "show_link_x" in request.POST
            )  # Checar se o checkbox foi marcado
            # Verificar e atualizar imagem de fundo (se fornecida)
            if "bg_image" in request.FILES:
                template.bg_image = request.FILES["bg_image"]

            # Atualizar cabeçalho
            template.bg_width = request.POST.get("bg_width", template.bg_width)
            template.bg_height = request.POST.get("bg_height", template.bg_height)
            template.text_color = request.POST.get("text_color", template.text_color)
            template.header_color = request.POST.get(
                "header_color", template.header_color
            )
            template.bg_position = request.POST.get("bg_position", template.bg_position)
            template.text_position = request.POST.get(
                "text_position", template.text_position
            )
            template.text_content = request.POST.get(
                "text_content", template.text_content
            )
            template.font_size = request.POST.get("font_size", template.font_size)
            template.main_title_top = request.POST.get(
                "main_title_top", template.main_title_top
            )
            template.main_title_left = request.POST.get(
                "main_title_left", template.main_title_left
            )
            template.title_top = request.POST.get("title_top", template.title_top)
            template.title_left = request.POST.get("title_left", template.title_left)

            # Atualizar footer
            template.footer_color = request.POST.get(
                "footer_color", template.footer_color
            )

            # Atualizar as cores e textos do email
            template.background_color = request.POST.get(
                "background_color", template.background_color
            )
            template.text_header_color = request.POST.get(
                "text_header_color", template.text_header_color
            )
            template.text_header_content = request.POST.get(
                "text_header_content", template.text_header_content
            )
            template.text_saudacao_color = request.POST.get(
                "text_saudacao_color", template.text_saudacao_color
            )
            template.text_saudacao_content = request.POST.get(
                "text_saudacao_content", template.text_saudacao_content
            )

            template.titulo = request.POST.get("titulo", template.titulo)

            # Atualizar favicon (se fornecido)
            if "favicon" in request.FILES:
                template.favicon = request.FILES["favicon"]

            # Salvar as alterações
            template.save()
            messages.success(request, "Links atualizados com sucesso.")

            return redirect("administrador")

    #return render(request, "administrador.html", {"template": template, "agencias": agencias})
    return render(request, "administracao/painel_admin.html", {"template": template, "agencias": agencias})

@login_required
def cartoesadm(request):
    # Obtém a lista de agências para exibir no template
    agencias = Agencia.objects.all()

    # Obtém o valor da agência selecionada no filtro
    agencia_filtro = request.GET.get("agencia")

    # Filtra os cartões de acordo com a agência selecionada, se houver
    if agencia_filtro:
        cartoes = AAFormuUserer.objects.filter(agencia__agencia=agencia_filtro).order_by('-data_atualizacao')
    else:
        cartoes = AAFormuUserer.objects.order_by('-data_atualizacao')

    # Contagem dos cartões por status
    total_cartoes = cartoes.count()
    total_reservado = cartoes.filter(status="reservado").count()
    total_respondido = cartoes.filter(status="respondido").count()
    total_pendencia = cartoes.filter(status="pendencia").count()
    total_iniciado = cartoes.filter(status="iniciado").count()
    total_produzindo = cartoes.filter(status="produzindo").count()
    total_finalizado = cartoes.filter(status="finalizado").count()
    total_erro = cartoes.filter(status="erro").count()

    context = {
        "cartoes": cartoes,
        "agencias": agencias,
        "agencia_selecionada": agencia_filtro,
        "total_cartoes": total_cartoes,
        "total_reservado": total_reservado,
        "total_respondido": total_respondido,
        "total_iniciado": total_iniciado,
        "total_pendencia": total_pendencia,
        "total_produzindo": total_produzindo,
        "total_finalizado": total_finalizado,
        "total_erro": total_erro,
    }

    return render(request, "administracao/cartoesadm.html", context)

def search_adm(request):
    query = request.GET.get("q", "")  # Termo de pesquisa
    # Captura o valor da data de início
    data_inicio = request.GET.get("data_inicio")
    data_fim = request.GET.get("data_fim")  # Captura o valor da data de fim


    results = AAFormuUserer.objects.all()

            # Filtro adicional com o termo de pesquisa, se fornecido
    if query:
        results = results.filter(
            Q(nome__icontains=query)
            | Q(email__icontains=query)
            | Q(pais__icontains=query)
            | Q(nicho__icontains=query)
            | Q(tema_base__icontains=query)
            | Q(url_loja__icontains=query)
            | Q(telefone__icontains=query)
            | Q(email_suporte__icontains=query)
            | Q(empresa__icontains=query)
            | Q(username__icontains=query)
            | Q(status__icontains=query)
        )
        print(
            f"Resultados após filtro de pesquisa: {results.count()}")  # Depuração

        # Filtro adicional de data
        if data_inicio:
            # Converte para formato de data
            data_inicio = parse_date(data_inicio)
            results = results.filter(data_atualizacao__gte=data_inicio)

        if data_fim:
            # Converte para formato de data
            data_fim = parse_date(data_fim)
            results = results.filter(data_atualizacao__lte=data_fim)

            total_results = results.count()  # Captura o total de resultados encontrados

            return render(
                request,
                "search_results.html",
                {
                    "results": results,
                    "total_results": total_results,
                    "data_inicio": data_inicio,
                    "data_fim": data_fim,
                },
            )
       
def retornaagencias(request):
    agencias = Agencia.objects.all()  # Recupera todas as agências do banco de dados
    return render(request, "administrador.html", {"agencias": agencias})

@csrf_exempt
def criaracessosadmin(request):
    if request.method == "POST":
        # Obtém o nome da agência diretamente do POST
        nome_agencia = request.POST.get("agencia", "")

        if not nome_agencia:
            return JsonResponse(
                {"status": "error", "message": "Nome da agência não fornecido"},
                status=400,
            )

        # Busca a agência pelo nome
        try:
            agencia = Agencia.objects.get(agencia=nome_agencia)
        except Agencia.DoesNotExist:
            return JsonResponse(
                {"status": "error", "message": "Agência não encontrada"}, status=404
            )

        # Cria 10 acessos
        acessos = []
        for _ in range(5):
            username = "".join(
                random.choices(string.ascii_letters + string.digits, k=8)
            )

            AAFormuUserer.objects.create(
                email=agencia.email,  # Presumindo que a classe Agencia tem um campo de email
                username=username,
                senha="Desconhecido",
                agencia=agencia,
                nome="Nome Padrão",
                pais="Brasil",
                cor="Desconhecido",
                nicho="Genérica",
                tema_base=agencia.tema_base,
                link_store="https://example.com/default-link",
                url_loja="example.com/",
                token_senha="Desconhecido",
                chave_de_api="Desconhecido",
                chave_secreta="Desconhecido",
                telefone="Desconhecido",
                email_suporte="suporte@example.com",
                empresa="Desconhecido",
                business_hours="Desconhecido",
                produtos=False,
                tema=False,
                paginas=False,
                banners=False,
                politicas=False,
                colecoes=False,
                estilizacao=False,
                lojaproduzida=False,
                status="reservado",
            )

            acessos.append({"username": username, "senha": "Desconhecido"})

        return redirect(administrador)

    return JsonResponse(
        {"status": "error", "message": "Método não permitido"}, status=405
    )

def cadastrar_agencia(request):
    if request.method == "POST":
        # Criar a instância do formulário com os dados enviados
        form = CadastroAgenciaForm(request.POST)
        
        if form.is_valid():  # Verifica se o formulário (incluindo o captcha) é válido
            # Dados do formulário
            agencia_nome = form.cleaned_data['agencia']
            email = form.cleaned_data['email']
            full_name = form.cleaned_data['full_name']
            username = form.cleaned_data['username']
            password = form.cleaned_data['password']
            telefone = form.cleaned_data.get('telefone', "")

            # Determinar o próximo template disponível
            total_templates = 9999  
            agencias_existentes = Agencia.objects.count()
            numero_template = agencias_existentes % total_templates + 1  # O próximo template será um número entre 1 e 9999
            template_usado = f"agencia{numero_template:04}"  # Gera 'agencia0001', 'agencia0002', ...

            # Criar a agência com o template associado
            agencia = Agencia.objects.create(
                agencia=agencia_nome,
                full_name=full_name,
                email=email,
                telefone=telefone,
                username=username,
                senha=password,
                template_usado=template_usado,  # Associa a agência ao template
                plano="ativo"
            )

            # Criar o usuário e salvar com senha criptografada
            user = User(username=username, email=email)
            user.set_password(password)
            user.save()

            PerfilUsuario.objects.create(user=user, agencia=agencia, full_name=full_name)

            # Criar 5 acessos adicionais
            for _ in range(2):
                username_acesso = "".join(
                    random.choices(string.ascii_letters + string.digits, k=8)
                )

                AAFormuUserer.objects.create(
                    email="cliente@example.com",  
                    username=username_acesso,
                    senha="Desconhecido",
                    agencia=agencia,
                    nome="Nome Padrão",
                    pais="Brasil",
                    cor="Desconhecido",
                    nicho="Genérica",
                    tema_base=agencia.tema_base,
                    link_store="https://example.com/default-link",
                    url_loja="example.com/",
                    token_senha="Desconhecido",
                    chave_de_api="Desconhecido",
                    chave_secreta="Desconhecido",
                    telefone="Desconhecido",
                    email_suporte="suporte@example.com",
                    empresa="Desconhecido",
                    business_hours="Desconhecido",
                    produtos=False,
                    tema=False,
                    paginas=False,
                    banners=False,
                    politicas=False,
                    colecoes=False,
                    estilizacao=False,
                    lojaproduzida=False,
                    status="reservado",
                )

            # Enviar e-mail de boas-vindas
            def enviar_email():
                email_subject = "Bem-vindo à ShopBotfy!"
                email_body = render_to_string(
                    "email_cadastro.html", {"agencia": agencia_nome}
                )
                mail = mt.Mail(
                    sender=mt.Address(
                        email="mailtrap@shopbotfy.com", name="ShopBotfy"
                    ),
                    to=[mt.Address(email=email)],
                    subject=email_subject,
                    html=email_body,
                )
                client = mt.MailtrapClient(token="141ed2a0f0b501224f4050c1878381f4")
                client.send(mail)

            email_thread = threading.Thread(target=enviar_email)
            email_thread.start()

            return redirect("home")
    else:
        form = CadastroAgenciaForm()

    return render(request, "cadastrar.html", {'form': form})

@csrf_exempt
def formulario(request): # RETORNA O LINK SHOPIFY DA AGENCIA
    if request.method == "GET":
        return render(request, "formulario.html")
    elif request.method == "POST":
        agencia_template = request.POST.get("agencia")  # Captura o valor do campo 'template_usado'
        if agencia_template:
            try:
                # Buscar a agência usando o campo 'template_usado'
                agencia = Agencia.objects.get(template_usado=agencia_template)
                link_impact = agencia.link_impact

                print(f"Agência: {agencia.agencia}")
                print(f"Link Impact: {link_impact}")

                return JsonResponse({"link_impact": link_impact})
            except Agencia.DoesNotExist:
                print("Agência não encontrada.")
                return JsonResponse({"error": "Agência não encontrada."}, status=404)
        else:
            print("Nenhuma agência informada.")
            return JsonResponse({"error": "Nenhuma agência informada."}, status=400)
    else:
        return JsonResponse({"error": "Método não permitido."}, status=405)

@csrf_exempt
def formulario_gringo(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            print("Dados recebidos:", data)

            # Captura os dados do formulário
            link_store = data.get("link_store")
            pais = data.get("pais")
            nome = data.get("nome")
            telefone = data.get("telefone")
            email_suporte = data.get("email_suporte")
            empresa = data.get("empresa")
            email = data.get("email_shopify")  # Corrigido aqui
            senha = data.get("senha")
            business_hours = data.get("business_hours")
            nicho = data.get("nicho")
            cor = data.get("cor")
            agencia_template = data.get("agencia", "").strip()

            print(f"Template da agência recebida: '{agencia_template}'")

            # Verifica se a agência existe
            try:
                agencia = Agencia.objects.get(template_usado=agencia_template)
                print(f"Agência encontrada: '{agencia.template_usado}'")
            except Agencia.DoesNotExist:
                print(f"Agência não encontrada: '{agencia_template}'")
                return JsonResponse(
                    {"status": "error", "message": "Agência não encontrada."},
                    status=400,
                )

            # Busca um formulário disponível para a agência
            formulario_user = AAFormuUserer.objects.filter(
                agencia=agencia, status="reservado"
            ).first()

            
            if formulario_user:
                # Extrai o identificador da loja do link fornecido
                match = re.match(
                    r"https://admin\.shopify\.com/store/([a-zA-Z0-9-]+)", link_store
                )
                if match:
                    loja_id = match.group(1)
                    url_loja = f"{loja_id}.myshopify.com"
                    cleaned_link_store = f"https://admin.shopify.com/store/{loja_id}"

                    # Atualiza o formulário com os dados recebidos
                    formulario_user.link_store = cleaned_link_store
                    formulario_user.url_loja = url_loja
                    formulario_user.pais = pais
                    formulario_user.nome = nome
                    formulario_user.telefone = telefone
                    formulario_user.cor = cor
                    formulario_user.nicho = nicho
                    formulario_user.email_suporte = email_suporte
                    formulario_user.empresa = empresa
                    formulario_user.email = email
                    formulario_user.senha = senha
                    formulario_user.business_hours = business_hours
                    formulario_user.status = "respondido"
                    formulario_user.save()

                    print(f"Formulário atualizado com sucesso: {formulario_user}")
                    return JsonResponse({"status": "success"}, status=200)
                else:
                    print("Link da loja inválido.")
                    return JsonResponse(
                        {"status": "error", "message": "Link da loja inválido."},
                        status=400,
                    )
            else:
                # Adiciona os dados à fila de espera
                FilaDeEspera.objects.create(agencia=agencia, dados_formulario=data)
                print("Formulário adicionado à fila de espera.")
                return JsonResponse(
                    {
                        "status": "success",
                        "message": "Sem cartões disponíveis. Formulário adicionado à fila.",
                    },
                    status=200,
                )

        except json.JSONDecodeError as e:
            print("Erro ao decodificar JSON:", e)
            return JsonResponse(
                {"status": "error", "message": "Erro ao processar os dados."},
                status=400,
            )
        except Exception as e:
            print("Erro inesperado:", e)
            return JsonResponse(
                {"status": "error", "message": "Erro inesperado."},
                status=500,
            )

    else:
        print("Método não permitido")
        return JsonResponse(
            {"status": "error", "message": "Método não permitido."}, status=405
        )

def agencia_dinamica(request, template_name):
    agencia = get_object_or_404(Agencia, template_usado=template_name)
    context = {
        "agencia": agencia,
    }
    return render(request, f"agencias/{template_name}.html", context)

@login_required
def painel_agencia(request):
    logger.debug("Início da view painel_agencia")
    print("Início da view painel_agencia")
    agencia_id = request.session.get("agencia_id")
    print(f"agencia_id: {agencia_id}")
    logger.debug(f"agencia_id: {agencia_id}")
    avatars = range(1, 31)  # Gera uma lista de números de 1 a 30

    if agencia_id:
        try:
            agencia = get_object_or_404(Agencia, id=agencia_id)
        except Exception as e:
            logger.error(f"Erro ao obter agência: {e}")
            return redirect("home")
        
        # Buscar notificações não lidas
        notificacoes = Notificacao.objects.filter(agencia=agencia, lida=False)

        # Verificar se já existe uma notificação para a fila de espera (sem necessidade de consultar a FilaDeEspera diretamente)
        if not notificacoes.filter(mensagem__contains="fila de espera").exists():
            # Contar os registros não processados na fila de espera
            fila_espera_count = FilaDeEspera.objects.filter(agencia=agencia, processado=False).count()
            
            if fila_espera_count > 0:
                # Verificar se já existe uma notificação de fila de espera
                notificacao_fila = Notificacao.objects.filter(agencia=agencia, mensagem__contains="fila de espera").first()
                
                if not notificacao_fila:
                    # Se a notificação ainda não foi criada, cria uma nova
                    Notificacao.objects.create(
                        agencia=agencia,
                        mensagem=f"Você tem {fila_espera_count} registros na fila de espera."
                    )
                
            # Atualizar a lista de notificações (caso alguma nova tenha sido criada)
            notificacoes = Notificacao.objects.filter(agencia=agencia, lida=False)


    if agencia_id:
        try:
            agencia = get_object_or_404(Agencia, id=agencia_id)
            logger.debug(f"Agencia: {agencia.agencia}")
        except Exception as e:
            logger.error(f"Erro ao obter agência: {e}")
            return redirect("home")

        if request.method == "POST":
            logger.debug("POST recebido")
            try:
                # Atualiza imagens
                agencia_bg_image_file = request.FILES.get("agencia_bg_image")
                favicon_agencia_file = request.FILES.get("favicon_agencia")
                generico_agencia_file = request.FILES.get("generico_agencia")
                eletronicos_agencia_file = request.FILES.get("eletronicos_agencia")
                casa_agencia_file = request.FILES.get("casa_agencia")
                moda_agencia_file = request.FILES.get("moda_agencia")
                kids_agencia_file = request.FILES.get("kids_agencia")
                pets_agencia_file = request.FILES.get("pets_agencia")
                fitness_agencia_file = request.FILES.get("fitness_agencia")
                saude_agencia_file = request.FILES.get("saude_agencia")
                masculino_agencia_file = request.FILES.get("masculino_agencia")
                feminino_agencia_file = request.FILES.get("feminino_agencia")
                oculos_agencia_file = request.FILES.get("oculos_agencia")
                relogio_agencia_file = request.FILES.get("relogio_agencia")
                joias_agencia_file = request.FILES.get("joias_agencia")

                
                if agencia_bg_image_file:
                    upload_result = cloudinary.uploader.upload(agencia_bg_image_file)
                    agencia.agencia_bg_image = upload_result["secure_url"]
                if favicon_agencia_file:
                    upload_result = cloudinary.uploader.upload(favicon_agencia_file)
                    agencia.favicon_agencia = upload_result["secure_url"]
                if generico_agencia_file:
                    upload_result = cloudinary.uploader.upload(generico_agencia_file)
                    agencia.generico_agencia = upload_result["secure_url"]
                if eletronicos_agencia_file:
                    upload_result = cloudinary.uploader.upload(eletronicos_agencia_file)
                    agencia.eletronicos_agencia = upload_result["secure_url"]
                if casa_agencia_file:
                    upload_result = cloudinary.uploader.upload(casa_agencia_file)
                    agencia.casa_agencia = upload_result["secure_url"]
                if moda_agencia_file:
                    upload_result = cloudinary.uploader.upload(moda_agencia_file)
                    agencia.moda_agencia = upload_result["secure_url"]
                if kids_agencia_file:
                    upload_result = cloudinary.uploader.upload(kids_agencia_file)
                    agencia.kids_agencia = upload_result["secure_url"]
                if pets_agencia_file:
                    upload_result = cloudinary.uploader.upload(pets_agencia_file)
                    agencia.pets_agencia = upload_result["secure_url"]
                if fitness_agencia_file:
                    upload_result = cloudinary.uploader.upload(fitness_agencia_file)
                    agencia.fitness_agencia = upload_result["secure_url"]
                if saude_agencia_file:
                    upload_result = cloudinary.uploader.upload(saude_agencia_file)
                    agencia.saude_agencia = upload_result["secure_url"]
                if masculino_agencia_file:
                    upload_result = cloudinary.uploader.upload(masculino_agencia_file)
                    agencia.masculino_agencia = upload_result["secure_url"]
                if feminino_agencia_file:
                    upload_result = cloudinary.uploader.upload(feminino_agencia_file)
                    agencia.feminino_agencia = upload_result["secure_url"]
                if oculos_agencia_file:
                    upload_result = cloudinary.uploader.upload(oculos_agencia_file)
                    agencia.oculos_agencia = upload_result["secure_url"]
                if relogio_agencia_file:
                    upload_result = cloudinary.uploader.upload(relogio_agencia_file)
                    agencia.relogio_agencia = upload_result["secure_url"]
                if joias_agencia_file:
                    upload_result = cloudinary.uploader.upload(joias_agencia_file)
                    agencia.joias_agencia = upload_result["secure_url"]

                

                # dimensões do banner logomarca
                agencia.agencia_bg_width = request.POST.get("agencia_bg_width")
                agencia.agencia_bg_height = request.POST.get("agencia_bg_height")
                
                # Atualiza os liks da agência
                agencia.titulo_agencia = request.POST.get("titulo_agencia")
                agencia.texto_cabecalho_agencia = request.POST.get("texto_cabecalho_agencia")
                agencia.cabecalho_font_size = request.POST.get("cabecalho_font_size")
                agencia.cabecalho_text_color = request.POST.get("cabecalho_text_color")
                agencia.agencia_header_color = request.POST.get("agencia_header_color")
                agencia.agencia_footer_color = request.POST.get("agencia_footer_color")
                agencia.agencia_footer_text_color = request.POST.get("agencia_footer_text_color")
                agencia.agencia_background_color = request.POST.get("agencia_background_color")
                agencia.agencia_explica_color = request.POST.get("agencia_explica_color")
                agencia.explicacao_text_color = request.POST.get("explicacao_text_color")
                agencia.apresentacao_font_size = request.POST.get("apresentacao_font_size")

                # copy da agência
                agencia.copy_agencia = request.POST.get("copy_agencia")

                # Atualiza os liks da agência
                agencia.link_impact = request.POST.get("link_impact")
                agencia.link_site_agencia = request.POST.get("link_site_agencia")
                agencia.telefone_agencia = request.POST.get("telefone_agencia")
                agencia.contato_agencia = request.POST.get("contato_agencia")
                agencia.link_face_agencia = request.POST.get("link_face_agencia")
                agencia.link_insta_agencia = request.POST.get("link_insta_agencia")
                agencia.link_youtube_agencia = request.POST.get("link_youtube_agencia")

                # Atualiza os liks dos nichos
                agencia.link_generico_agencia = request.POST.get("link_generico_agencia")
                agencia.link_eletronicos_agencia = request.POST.get("link_eletronicos_agencia")
                agencia.link_casa_agencia = request.POST.get("link_casa_agencia")
                agencia.link_moda_agencia = request.POST.get("link_moda_agencia")
                agencia.link_kids_agencia = request.POST.get("link_kids_agencia")
                agencia.link_pets_agencia = request.POST.get("link_pets_agencia")
                agencia.link_fitness_agencia = request.POST.get("link_fitness_agencia")
                agencia.link_saude_agencia = request.POST.get("link_saude_agencia")
                agencia.link_masculino_agencia = request.POST.get("link_masculino_agencia")
                agencia.link_feminino_agencia = request.POST.get("link_feminino_agencia")
                agencia.link_oculos_agencia = request.POST.get("link_oculos_agencia")
                agencia.link_relogio_agencia = request.POST.get("link_relogio_agencia")
                agencia.link_joias_agencia = request.POST.get("link_joias_agencia")


                # checkbox cabeçalo e rodapé
                agencia.show_texto_cabecalho_agencia = ("show_texto_cabecalho_agencia" in request.POST)
                agencia.show_link_site_agencia = ("show_link_site_agencia" in request.POST)  
                agencia.show_telefone_agencia = ("show_telefone_agencia" in request.POST)  
                agencia.show_contato_agencia = ("show_contato_agencia" in request.POST) 
                agencia.show_link_face_agencia = ("show_link_face_agencia" in request.POST) 
                agencia.show_link_insta_agencia = ("show_link_insta_agencia" in request.POST) 
                agencia.show_link_youtube_agencia = ("show_link_youtube_agencia" in request.POST)

                # checkbox nichos
                agencia.show_nicho_generico_agencia = "show_nicho_generico_agencia" in request.POST
                agencia.show_nicho_eletronicos_agencia = "show_nicho_eletronicos_agencia" in request.POST
                agencia.show_nicho_casa_agencia = "show_nicho_casa_agencia" in request.POST
                agencia.show_nicho_moda_agencia = "show_nicho_moda_agencia" in request.POST
                agencia.show_nicho_kids_agencia = "show_nicho_kids_agencia" in request.POST
                agencia.show_nicho_pets_agencia = "show_nicho_pets_agencia" in request.POST
                agencia.show_nicho_fitness_agencia = "show_nicho_fitness_agencia" in request.POST
                agencia.show_nicho_saude_agencia = "show_nicho_saude_agencia" in request.POST
                agencia.show_nicho_masculino_agencia = "show_nicho_masculino_agencia" in request.POST
                agencia.show_nicho_feminino_agencia = "show_nicho_feminino_agencia" in request.POST
                agencia.show_nicho_oculos_agencia = "show_nicho_oculos_agencia" in request.POST
                agencia.show_nicho_relogio_agencia = "show_nicho_relogio_agencia" in request.POST
                agencia.show_nicho_joias_agencia = "show_nicho_joias_agencia" in request.POST

                # checkbox paises
                agencia.show_pais_argentina_agencia = "show_pais_argentina_agencia" in request.POST
                agencia.show_pais_brasil_agencia = "show_pais_brasil_agencia" in request.POST
                agencia.show_pais_bolivia_agencia = "show_pais_bolivia_agencia" in request.POST
                agencia.show_pais_chile_agencia = "show_pais_chile_agencia" in request.POST
                agencia.show_pais_colombia_agencia = "show_pais_colombia_agencia" in request.POST
                agencia.show_pais_paraguai_agencia = "show_pais_paraguai_agencia" in request.POST
                agencia.show_pais_peru_agencia = "show_pais_peru_agencia" in request.POST
                agencia.show_pais_uruguai_agencia = "show_pais_uruguai_agencia" in request.POST
                agencia.show_pais_venezuela_agencia = "show_pais_venezuela_agencia" in request.POST
                agencia.show_pais_equador_agencia = "show_pais_equador_agencia" in request.POST
                agencia.show_pais_espanha_agencia = "show_pais_espanha_agencia" in request.POST
                agencia.show_pais_eua_agencia = "show_pais_eua_agencia" in request.POST
                agencia.show_pais_portugal_agencia = "show_pais_portugal_agencia" in request.POST



                agencia.save()
                print("ados atualizados")
                return redirect("painel_agencia")
            except Exception as e:
                logger.error(f"Erro ao salvar dados: {e}")
                return render(
                    request,
                    "painel_agencia.html",
                    {
                        "agencia": agencia,
                        "escassez_mensagem": "Erro ao salvar os dados. Tente novamente.",
                    },
                )
       
 
        agencia_nome = agencia.agencia.strip().lower().replace(" ", "-")  # Exemplo de transformação para URL
        template_usado = agencia.template_usado  # Nome do template usado
        template_url = f"/agencia/{template_usado}/"  # Ajuste para usar o nome do template

        print(f"Template URL: {template_url}")

        return render(
            request,
            "painel_agencia.html",
            {
                "agencia": agencia,
                "template_url": template_url,
                "avatars": range(1, 31),
                'notificacoes': notificacoes,
            }
        )
    return redirect("home")

############################
############################################################
                                                #########################   COMUNICAÇÃO  
############################################################
def notificacoes(request):
    agencias = Agencia.objects.all()

    if request.method == "POST":
        agencia_ids = request.POST.get("agencias", "").split(",")  # IDs das agências selecionadas
        mensagem = request.POST.get("mensagem")  # Mensagem a ser enviada

        if agencia_ids and mensagem:
            for agencia_id in agencia_ids:
                if agencia_id:  # Evitar IDs vazios
                    agencia = Agencia.objects.get(id=agencia_id)
                    Notificacao.objects.create(agencia=agencia, mensagem=mensagem, data_criacao=now())

        return redirect("minhas_notificacoes")  # Redireciona para a mesma página após envio

    return render(request, "administracao/notificacoes.html", {"agencias": agencias})

@csrf_exempt
def marcar_como_lida(request, notification_id):
    if request.method == 'POST':
        try:
            notificacao = Notificacao.objects.get(id=notification_id)
            notificacao.lida = True
            notificacao.save()
            logger.debug(f"Notificação {notification_id} marcada como lida.")  # Adicionando log para confirmar a alteração

            return JsonResponse({'status': 'success'})
        except Notificacao.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Notificação não encontrada'}, status=404)
    return JsonResponse({'status': 'error', 'message': 'Método não permitido'}, status=405)

def email_selecao(request):
    agencias = Agencia.objects.all()  # Obtém todas as agências
    if request.method == 'POST':
        selected_agencias = request.POST.getlist('agencias')  # Recebe as agências selecionadas
        template = request.POST.get('template')  # Obtém o template selecionado

        for agencia_id in selected_agencias:
            agencia = Agencia.objects.get(id=agencia_id)
            emailnotificacao(agencia, template)  # Envia o e-mail para cada agência

        return redirect('email_selecao')  # Redireciona após o envio

    return render(request, 'administracao/email_selecao.html', {'agencias': agencias})

def emailnotificacao(agencia, template):
    email_subject = "Bem-vindo à Plataforma"
    
    # Renderiza o corpo do e-mail com o template
    email_body = render_to_string(
        "administracao/emailnotificacao.html",
        {
            "agencia": agencia.agencia,
            "template": template,
        },
    )

    recipient_email = agencia.email

    mail = mt.Mail(
        sender=mt.Address(email="mailtrap@shopbotfy.com", name="ShopBotfy"),
        to=[mt.Address(email=recipient_email)],
        subject=email_subject,
        html=email_body,
    )

    client = mt.MailtrapClient(token="141ed2a0f0b501224f4050c1878381f4")

    try:
        # Tentativa de envio de e-mail
        response = client.send(mail)
        print(f"Email sent successfully: {response.status_code}")
    except Exception as e:
        print(f"Error sending email: {e}")

def enviar_email(agencia, username, senha, template):
    email_subject = "Bem-vindo à Plataforma"

    # Renderiza o corpo do e-mail com o template
    email_body = render_to_string(
        "email_agencia.html",
        {
            "agencia": agencia,
            "username": username,
            "senha": senha,
            "template": template,
        },
    )

    recipient_email = agencia.email

    mail = mt.Mail(
        sender=mt.Address(email="mailtrap@shopbotfy.com", name="ShopBotfy"),
        to=[mt.Address(email=recipient_email)],
        subject=email_subject,
        html=email_body,
    )

    client = mt.MailtrapClient(token="141ed2a0f0b501224f4050c1878381f4")

    try:
        # Tentativa de envio de e-mail
        response = client.send(mail)
        print(
            f"Email sent successfully: {response.status_code}"
        )  # Verifique o código de status do envio
    except Exception as e:
        print(f"Error sending email: {e}")

def mail(request):
    return render(request, "mail.html")

def email(request):
    return render(request, "email.html")

def email_cadastro(request):
    return render(request, "email_cadastro.html")

def email_agencia(request):
    return render(request, "email_agencia.html")


def emailgl(request):
    return render(request, "emailgl.html")

@csrf_exempt
def pedidogl(request):
    if request.method == "POST":
        try:
            # Extrair dados do corpo da solicitação POST (JSON)
            payload = json.loads(request.body)

            # Extrair nome e email do payload
            customer = payload.get("Customer", {})
            nome = customer.get("full_name", "")
            email = customer.get("email", "")

            if nome and email:
                # Enviar resposta ao webhook
                response_data = {
                    "status": "success",
                    "message": "E-mail será enviado!",
                }
                response = JsonResponse(response_data, status=200)

                # Enviar o e-mail após enviar a resposta ao webhook
                def enviar_email():
                    email_subject = "Pedido Aprovado"
                    email_body = render_to_string(
                        "emailgl.html", {"nome": nome, "email": email}
                    )

                    mail = mt.Mail(
                        sender=mt.Address(
                            email="mailtrap@shopbotfy.com", name="ShopBotfy"
                        ),
                        to=[mt.Address(email=email)],
                        subject=email_subject,
                        html=email_body,
                    )

                    client = mt.MailtrapClient(token="141ed2a0f0b501224f4050c1878381f4")
                    client.send(mail)

                # Usar threading para enviar o e-mail após responder ao webhook
                email_thread = threading.Thread(target=enviar_email)
                email_thread.start()

                return response
            else:
                return JsonResponse(
                    {"status": "error", "message": "Dados incompletos"}, status=400
                )

        except json.JSONDecodeError:
            return JsonResponse(
                {"status": "error", "message": "Formato de JSON inválido"}, status=400
            )

    # Se a solicitação não for POST, retornar erro
    return JsonResponse(
        {"status": "error", "message": "Método não permitido"}, status=405
    )


############################
############################################################
                                                #########################   HOME E PAGINAS DO SITE  
############################################################
def login(request):
    if request.method == "GET":
        return render(request, "login.html")

    # Capturar dados do formulário
    username = request.POST.get("username")
    senha = request.POST.get("senha")
    user = authenticate(username=username, password=senha)

    mensagens = []  # Lista para mensagens a serem exibidas no popup
    redirect_url = None  # URL para redirecionamento após o popup

    # Verificar se os dados de login estão corretos
    if user:
        login_django(request, user)

        # Verificar se é superuser
        if user.is_superuser:
            return redirect("administrador")

        # Obter a agência associada ao username
        agencia = Agencia.objects.filter(username=username).first()
        if agencia:
            # Verificação do plano 'ativo' (essa lógica é executada primeiro)
            if agencia.plano == "ativo":
                # Redireciona imediatamente para 'painel_agencia'

                request.session["agencia_id"] = (
                    agencia.id
                )  # Salva a agencia_id na sessão
                request.session["agencia_nome"] = agencia.agencia
                request.session["agencia_plano"] = agencia.plano
                print(
                    f"agencia_id: {request.session['agencia_id']}"
                )  # Confirmação no log
                return redirect("painel_agencia")

            try:
                # Lógica para planos 'teste', 'limite', 'atraso', etc.
                if agencia.plano == "teste":
                    mensagens.append(
                        "Você está usando um período de avaliação gratuita."
                    )
                    redirect_url = "painel_agencia"

                elif agencia.plano == "limite":
                    mensagens.append(
                        "Seu período de teste expirou. Adquira um plano para continuar."
                    )
                    redirect_url = "painel_agencia"

                elif agencia.plano.startswith("atraso"):
                    dias_atraso = int(agencia.plano.replace("atraso", ""))
                    mensagens.append(
                        f"Seu plano está atrasado em {dias_atraso} dias. Regularize seu pagamento."
                    )
                    redirect_url = "painel_agencia"

                elif agencia.plano == "vencido":
                    mensagens.append(
                        "Seu plano está vencido. Por favor, renove para continuar."
                    )
                    redirect_url = "home"

                else:
                    mensagens.append("Plano inválido. Entre em contato com o suporte.")
                    redirect_url = "home"

                # Salvar informações da agência
                request.session["agencia_id"] = agencia.id
                request.session["agencia_nome"] = agencia.agencia
                request.session["agencia_plano"] = agencia.plano
                print(
                    f"agencia_id: {request.session['agencia_id']}"
                )  # Confirmação no log

            except Exception as e:
                print(f"Erro ao processar plano da agência: {e}")
                mensagens.append("Ocorreu um erro. Tente novamente mais tarde.")

        else:
            mensagens.append("Agência não encontrada.")
            redirect_url = "home"

    else:
        # Dados de login inválidos
        mensagens.append("Credenciais inválidas. Tente novamente.")
        redirect_url = "login"  # Não precisa de redirecionamento específico

    return render(
        request,
        "login.html",
        {
            "mensagens": mensagens,
            "redirect_url": redirect_url,
        },
    )

def user_logout(request):
    logout(request)  # Faz o logout do usuário
    return redirect("home")  # Redireciona para a página inicial

def home(request):
    return render(request, "home.html")

def teste(request):
    return render(request, "teste.html")

def politica_privacidade(request): # POLITICA DA EXTENSÃO
    return render(request, "politica_privacidade.html")

def politica_site(request):
    return render(request, "politica.html")

def termos_de_uso(request):
    return render(request, "termos.html")

def ouvidoria(request):
    return render(request, "ouvidoria.html")

def sobrenos(request):
    return render(request, "sobrenos.html")

def historia(request):
    return render(request, "historia.html")

def enviar_ouvidoria(request):
    if request.method == "POST":
        nome = request.POST.get('nome')
        email = request.POST.get('email')
        telefone = request.POST.get('telefone')
        mensagem = request.POST.get('mensagem')

        # Cria uma nova mensagem de ouvidoria e salva no banco de dados
        nova_mensagem = OuvidoriaMensagem(
            nome=nome,
            email=email,
            telefone=telefone,
            mensagem=mensagem
        )
        nova_mensagem.save()

        # Redireciona para uma página de agradecimento
        return redirect('agradecimento')
    else:
        return HttpResponse("Método não permitido.", status=405)

def agradecimento(request):
    return render(request, "agradecimento.html")


############################
############################################################
                                                #########################   WEBHOOKS  
############################################################

@csrf_exempt
def acesso1(request):
    if request.method == "POST":
        try:
            payload = json.loads(request.body)
            customer = payload.get("Customer", {})
            nome = customer.get("full_name", "")
            email = customer.get("email", "")

            if nome and email:
                try:
                    agencia = Agencia.objects.get(email=email)
                except Agencia.DoesNotExist:
                    return JsonResponse(
                        {"status": "error", "message": "Agência não encontrada"},
                        status=404,
                    )

                tema_base = agencia.tema_base  # Recupera o tema base da agência

                request.session["agencia_nome"] = agencia.agencia

                acessos = []
                for _ in range(1):
                    username = "".join(
                        random.choices(string.ascii_letters + string.digits, k=8)
                    )

                    formu_user = AAFormuUserer.objects.create(
                        email="cliente@example.com",
                        username=username,
                        senha="Desconhecido",
                        agencia=agencia,  # Usar a instância de agencia
                        nome="Nome Padrão",
                        pais="Brasil",
                        cor="Desconhecido",
                        nicho="Genérica",
                        tema_base=tema_base,  # Usar o tema base da agência
                        link_store="https://example.com/default-link",
                        url_loja="example.com/",
                        token_senha="Desconhecido",
                        chave_de_api="Desconhecido",
                        chave_secreta="Desconhecido",
                        telefone="Desconhecido",
                        email_suporte="suporte@example.com",
                        empresa="Desconhecido",
                        business_hours="Desconhecido",
                        produtos=False,
                        tema=False,
                        paginas=False,
                        banners=False,
                        politicas=False,
                        colecoes=False,
                        estilizacao=False,
                        lojaproduzida=False,
                        status="reservado",
                    )
                    acessos.append({"username": username, "senha": "Desconhecido"})

                response_data = {
                    "status": "success",
                    "message": "1 formulário criado com sucesso!",
                    "acessos": acessos,
                }
                response = JsonResponse(response_data, status=200)
                response["Content-Type"] = "application/json"

                

                return response
            else:
                return JsonResponse(
                    {"status": "error", "message": "Dados incompletos"}, status=400
                )
        except json.JSONDecodeError:
            return JsonResponse(
                {"status": "error", "message": "Formato de JSON inválido"}, status=400
            )

    return JsonResponse(
        {"status": "error", "message": "Método não permitido"}, status=405
    )


@csrf_exempt
def acesso5(request):
    if request.method == "POST":
        try:
            payload = json.loads(request.body)
            customer = payload.get("Customer", {})
            nome = customer.get("full_name", "")
            email = customer.get("email", "")

            if nome and email:
                try:
                    agencia = Agencia.objects.get(email=email)
                except Agencia.DoesNotExist:
                    return JsonResponse(
                        {"status": "error", "message": "Agência não encontrada"},
                        status=404,
                    )

                tema_base = agencia.tema_base  # Recupera o tema base da agência

                request.session["agencia_nome"] = agencia.agencia

                acessos = []
                for _ in range(5):
                    username = "".join(
                        random.choices(string.ascii_letters + string.digits, k=8)
                    )

                    formu_user = AAFormuUserer.objects.create(
                        email="cliente@example.com",
                        username=username,
                        senha="Desconhecido",
                        agencia=agencia,  # Usar a instância de agencia
                        nome="Nome Padrão",
                        pais="Brasil",
                        cor="Desconhecido",
                        nicho="Genérica",
                        tema_base=tema_base,  # Usar o tema base da agência
                        link_store="https://example.com/default-link",
                        url_loja="example.com/",
                        token_senha="Desconhecido",
                        chave_de_api="Desconhecido",
                        chave_secreta="Desconhecido",
                        telefone="Desconhecido",
                        email_suporte="suporte@example.com",
                        empresa="Desconhecido",
                        business_hours="Desconhecido",
                        produtos=False,
                        tema=False,
                        paginas=False,
                        banners=False,
                        politicas=False,
                        colecoes=False,
                        estilizacao=False,
                        lojaproduzida=False,
                        status="reservado",
                    )
                    acessos.append({"username": username, "senha": "Desconhecido"})

                response_data = {
                    "status": "success",
                    "message": "5 formulários criados com sucesso!",
                    "acessos": acessos,
                }
                response = JsonResponse(response_data, status=200)
                response["Content-Type"] = "application/json"


                return response
            else:
                return JsonResponse(
                    {"status": "error", "message": "Dados incompletos"}, status=400
                )
        except json.JSONDecodeError:
            return JsonResponse(
                {"status": "error", "message": "Formato de JSON inválido"}, status=400
            )

    return JsonResponse(
        {"status": "error", "message": "Método não permitido"}, status=405
    )


@csrf_exempt
def acesso10(request):
    if request.method == "POST":
        try:
            payload = json.loads(request.body)
            customer = payload.get("Customer", {})
            nome = customer.get("full_name", "")
            email = customer.get("email", "")

            if nome and email:
                try:
                    agencia = Agencia.objects.get(email=email)
                except Agencia.DoesNotExist:
                    return JsonResponse(
                        {"status": "error", "message": "Agência não encontrada"},
                        status=404,
                    )

                tema_base = agencia.tema_base  # Recupera o tema base da agência

                request.session["agencia_nome"] = agencia.agencia

                acessos = []
                for _ in range(10):
                    username = "".join(
                        random.choices(string.ascii_letters + string.digits, k=8)
                    )

                    formu_user = AAFormuUserer.objects.create(
                        email="cliente@example.com",
                        username=username,
                        senha="Desconhecido",
                        agencia=agencia,  # Usar a instância de agencia
                        nome="Nome Padrão",
                        pais="Brasil",
                        cor="Desconhecido",
                        nicho="Genérica",
                        tema_base=tema_base,  # Usar o tema base da agência
                        link_store="https://example.com/default-link",
                        url_loja="example.com/",
                        token_senha="Desconhecido",
                        chave_de_api="Desconhecido",
                        chave_secreta="Desconhecido",
                        telefone="Desconhecido",
                        email_suporte="suporte@example.com",
                        empresa="Desconhecido",
                        business_hours="Desconhecido",
                        produtos=False,
                        tema=False,
                        paginas=False,
                        banners=False,
                        politicas=False,
                        colecoes=False,
                        estilizacao=False,
                        lojaproduzida=False,
                        status="reservado",
                    )
                    acessos.append({"username": username, "senha": "Desconhecido"})

                response_data = {
                    "status": "success",
                    "message": "10 formulários criados com sucesso!",
                    "acessos": acessos,
                }
                response = JsonResponse(response_data, status=200)
                response["Content-Type"] = "application/json"

              

                return response
            else:
                return JsonResponse(
                    {"status": "error", "message": "Dados incompletos"}, status=400
                )
        except json.JSONDecodeError:
            return JsonResponse(
                {"status": "error", "message": "Formato de JSON inválido"}, status=400
            )

    return JsonResponse(
        {"status": "error", "message": "Método não permitido"}, status=405
    )


@csrf_exempt
def acesso30(request):
    if request.method == "POST":
        try:
            payload = json.loads(request.body)
            customer = payload.get("Customer", {})
            nome = customer.get("full_name", "")
            email = customer.get("email", "")

            if nome and email:
                try:
                    agencia = Agencia.objects.get(email=email)
                except Agencia.DoesNotExist:
                    return JsonResponse(
                        {"status": "error", "message": "Agência não encontrada"},
                        status=404,
                    )

                tema_base = agencia.tema_base  # Recupera o tema base da agência

                request.session["agencia_nome"] = agencia.agencia

                acessos = []
                for _ in range(30):
                    username = "".join(
                        random.choices(string.ascii_letters + string.digits, k=8)
                    )

                    formu_user = AAFormuUserer.objects.create(
                        email="cliente@example.com",
                        username=username,
                        senha="Desconhecido",
                        agencia=agencia,  # Usar a instância de agencia
                        nome="Nome Padrão",
                        pais="Brasil",
                        cor="Desconhecido",
                        nicho="Genérica",
                        tema_base=tema_base,  # Usar o tema base da agência
                        link_store="https://example.com/default-link",
                        url_loja="example.com/",
                        token_senha="Desconhecido",
                        chave_de_api="Desconhecido",
                        chave_secreta="Desconhecido",
                        telefone="Desconhecido",
                        email_suporte="suporte@example.com",
                        empresa="Desconhecido",
                        business_hours="Desconhecido",
                        produtos=False,
                        tema=False,
                        paginas=False,
                        banners=False,
                        politicas=False,
                        colecoes=False,
                        estilizacao=False,
                        lojaproduzida=False,
                        status="reservado",
                    )
                    acessos.append({"username": username, "senha": "Desconhecido"})

                response_data = {
                    "status": "success",
                    "message": "30 formulários criados com sucesso!",
                    "acessos": acessos,
                }
                response = JsonResponse(response_data, status=200)
                response["Content-Type"] = "application/json"

                
                return response
            else:
                return JsonResponse(
                    {"status": "error", "message": "Dados incompletos"}, status=400
                )
        except json.JSONDecodeError:
            return JsonResponse(
                {"status": "error", "message": "Formato de JSON inválido"}, status=400
            )

    return JsonResponse(
        {"status": "error", "message": "Método não permitido"}, status=405
    )


@csrf_exempt
def acesso50(request):
    if request.method == "POST":
        try:
            payload = json.loads(request.body)
            customer = payload.get("Customer", {})
            nome = customer.get("full_name", "")
            email = customer.get("email", "")

            if nome and email:
                try:
                    agencia = Agencia.objects.get(email=email)
                except Agencia.DoesNotExist:
                    return JsonResponse(
                        {"status": "error", "message": "Agência não encontrada"},
                        status=404,
                    )

                tema_base = agencia.tema_base  # Recupera o tema base da agência

                request.session["agencia_nome"] = agencia.agencia

                acessos = []
                for _ in range(50):
                    username = "".join(
                        random.choices(string.ascii_letters + string.digits, k=8)
                    )

                    formu_user = AAFormuUserer.objects.create(
                        email="cliente@example.com",
                        username=username,
                        senha="Desconhecido",
                        agencia=agencia,  # Usar a instância de agencia
                        nome="Nome Padrão",
                        pais="Brasil",
                        cor="Desconhecido",
                        nicho="Genérica",
                        tema_base=tema_base,  # Usar o tema base da agência
                        link_store="https://example.com/default-link",
                        url_loja="example.com/",
                        token_senha="Desconhecido",
                        chave_de_api="Desconhecido",
                        chave_secreta="Desconhecido",
                        telefone="Desconhecido",
                        email_suporte="suporte@example.com",
                        empresa="Desconhecido",
                        business_hours="Desconhecido",
                        produtos=False,
                        tema=False,
                        paginas=False,
                        banners=False,
                        politicas=False,
                        colecoes=False,
                        estilizacao=False,
                        lojaproduzida=False,
                        status="reservado",
                    )
                    acessos.append({"username": username, "senha": "Desconhecido"})

                response_data = {
                    "status": "success",
                    "message": "50 formulários criados com sucesso!",
                    "acessos": acessos,
                }
                response = JsonResponse(response_data, status=200)
                response["Content-Type"] = "application/json"

                

                return response
            else:
                return JsonResponse(
                    {"status": "error", "message": "Dados incompletos"}, status=400
                )
        except json.JSONDecodeError:
            return JsonResponse(
                {"status": "error", "message": "Formato de JSON inválido"}, status=400
            )

    return JsonResponse(
        {"status": "error", "message": "Método não permitido"}, status=405
    )


@csrf_exempt
def acesso100(request):
    if request.method == "POST":
        try:
            payload = json.loads(request.body)
            customer = payload.get("Customer", {})
            nome = customer.get("full_name", "")
            email = customer.get("email", "")

            if nome and email:
                try:
                    agencia = Agencia.objects.get(email=email)
                except Agencia.DoesNotExist:
                    return JsonResponse(
                        {"status": "error", "message": "Agência não encontrada"},
                        status=404,
                    )

                tema_base = agencia.tema_base  # Recupera o tema base da agência

                request.session["agencia_nome"] = agencia.agencia

                acessos = []
                for _ in range(100):
                    username = "".join(
                        random.choices(string.ascii_letters + string.digits, k=8)
                    )

                    formu_user = AAFormuUserer.objects.create(
                        email="cliente@example.com",
                        username=username,
                        senha="Desconhecido",
                        agencia=agencia,  # Usar a instância de agencia
                        nome=nome,
                        pais="Brasil",
                        cor="Desconhecido",
                        nicho="Genérica",
                        tema_base=tema_base,  # Usar o tema base da agência
                        link_store="https://example.com/default-link",
                        url_loja="example.com/",
                        token_senha="Desconhecido",
                        chave_de_api="Desconhecido",
                        chave_secreta="Desconhecido",
                        telefone="Desconhecido",
                        email_suporte="suporte@example.com",
                        empresa="Desconhecido",
                        business_hours="Desconhecido",
                        produtos=False,
                        tema=False,
                        paginas=False,
                        banners=False,
                        politicas=False,
                        colecoes=False,
                        estilizacao=False,
                        lojaproduzida=False,
                        status="reservado",
                    )
                    acessos.append({"username": username, "senha": "Desconhecido"})

                response_data = {
                    "status": "success",
                    "message": "100 formulários criados com sucesso!",
                    "acessos": acessos,
                }
                response = JsonResponse(response_data, status=200)
                response["Content-Type"] = "application/json"


                return response
            else:
                return JsonResponse(
                    {"status": "error", "message": "Dados incompletos"}, status=400
                )
        except json.JSONDecodeError:
            return JsonResponse(
                {"status": "error", "message": "Formato de JSON inválido"}, status=400
            )

    return JsonResponse(
        {"status": "error", "message": "Método não permitido"}, status=405
    )

@csrf_exempt
def webhook_renovacao(request):
    if request.method == "POST":
        try:
            payload = json.loads(request.body)
            customer = payload.get("Customer", {})
            email = customer.get("email", "")

            if not email:
                return JsonResponse({"error": "Email não fornecido"}, status=400)

            agencia = Agencia.objects.filter(email=email).first()
            if agencia:
                agencia.plano = "ativo"
                agencia.data_renovacao = timezone.now()  # Atualiza a data de renovação
                agencia.save()
                return JsonResponse(
                    {"status": "Plano atualizado com sucesso"}, status=200
                )
            else:
                return JsonResponse({"error": "Agência não encontrada"}, status=404)
        except json.JSONDecodeError:
            return JsonResponse({"error": "Erro ao decodificar o JSON"}, status=400)
        except Exception as e:
            return JsonResponse({"error": str(e)}, status=500)
    return JsonResponse({"error": "Método não permitido"}, status=405)

############################
############################################################
                                                #########################   FUNÇÕES AGÊNCIAS 
############################################################
def agenciaestyle(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'agencias/agenciaestyle.html', {'agencia': agencia})

def perfil_view(request):
    # Caminho da pasta de avatares
    avatar_dir = os.path.join(settings.STATIC_ROOT, "home/img/avatar/")
    avatar_urls = []

    # Lista os arquivos no diretório
    if os.path.exists(avatar_dir):
        for filename in os.listdir(avatar_dir):
            if filename.endswith((".png", ".jpg", ".jpeg", ".svg")):
                avatar_urls.append(f"/static/home/img/avatar/{filename}")

    # Verifica se foi enviado algum avatar no formulário
    if request.method == "POST":
        selected_avatar = request.POST.get("selected_avatar")
        uploaded_avatar = request.FILES.get("avatar")

        if selected_avatar:
            # Salvar o avatar selecionado no perfil do usuário
            request.user.profile.avatar = selected_avatar
            request.user.profile.save()
        elif uploaded_avatar:
            # Salvar o avatar enviado pelo usuário
            request.user.profile.avatar.save(uploaded_avatar.name, uploaded_avatar)

    return render(request, "perfil.html", {"avatars": avatar_urls})

def atualizar_perfil(request): 
    agencia_id = request.session.get("agencia_id")

    if agencia_id:
        try:
            agencia = get_object_or_404(Agencia, id=agencia_id)
            user = User.objects.get(username=agencia.username)
        except Exception as e:
            logger.error(f"Erro ao obter agência ou usuário: {e}")
            return redirect("home")

        if request.method == "POST":
            try:
                senha_atual = request.POST.get("senha_atual")
                nova_senha = request.POST.get("nova_senha")
                novo_login = request.POST.get("novo_login")
                novo_email = request.POST.get("novo_email")
                novo_telefone = request.POST.get("novo_telefone")
                novo_nome_agencia = request.POST.get("novo_nome_agencia")
                avatar_file = request.FILES.get("avatar")

                # Verificar se a senha atual está correta
                if senha_atual and not user.check_password(senha_atual):
                    messages.error(request, "Senha atual incorreta.")
                    return redirect("painel_agencia")

                # Atualizar login se fornecido
                if novo_login:
                    user.username = novo_login
                    agencia.username = novo_login

                # Atualizar senha se fornecido
                if nova_senha:
                    user.set_password(nova_senha)
                    agencia.senha = nova_senha

                # Atualizar email se fornecido
                if novo_email:
                    user.email = novo_email
                    agencia.email = novo_email

                # Atualizar telefone se fornecido
                if novo_telefone:
                    agencia.telefone = novo_telefone

                # Atualizar nome da agência se fornecido
                if novo_nome_agencia:
                    agencia.agencia = novo_nome_agencia
                    request.session['agencia_nome'] = novo_nome_agencia  # Atualiza o nome da agência na sessão

                # Atualizar avatar se fornecido
                if avatar_file:
                    upload_result = cloudinary.uploader.upload(avatar_file)
                    agencia.avatar = upload_result["secure_url"]

                # Salvar alterações
                user.save()
                agencia.save()

                # Se a senha foi alterada, redirecionar para a página de login
                if nova_senha:
                    logout(request)  # Invalidar a sessão atual
                    #messages.success(request, "Senha alterada com sucesso.")
                    return redirect("login")  # Redirecionar para a página de login

                messages.success(request, "Perfil atualizado com sucesso!")
                return redirect("painel_agencia")

            except Exception as e:
                logger.error(f"Erro ao salvar dados: {e}")
                messages.error(request, "Erro ao atualizar o perfil.")
                return redirect("painel_agencia")

    return render(request, "painel_agencia.html", {"agencia": agencia})

def search_view(request):
    query = request.GET.get("q", "")  # Termo de pesquisa
    agencia_nome = request.session.get("agencia_nome")  # Obtém a agência da sessão
    # Captura o valor da data de início
    data_inicio = request.GET.get("data_inicio")
    data_fim = request.GET.get("data_fim")  # Captura o valor da data de fim

    print(f"agencia_nome: {agencia_nome}")  # Depuração

    if agencia_nome:
        # Verifica se o nome da agência existe na tabela Agencia
        agencia = Agencia.objects.filter(agencia=agencia_nome).first()

        print(f"agencia encontrada: {agencia}")  # Depuração

        if agencia:
            # Filtrando pela agência no modelo AAFormuUserer
            results = AAFormuUserer.objects.filter(agencia=agencia)

            print(
                f"Resultados após filtro de agência: {results.count()}")  # Depuração

            # Filtro adicional com o termo de pesquisa, se fornecido
            if query:
                results = results.filter(
                    Q(nome__icontains=query)
                    | Q(email__icontains=query)
                    | Q(pais__icontains=query)
                    | Q(nicho__icontains=query)
                    | Q(tema_base__icontains=query)
                    | Q(url_loja__icontains=query)
                    | Q(telefone__icontains=query)
                    | Q(email_suporte__icontains=query)
                    | Q(empresa__icontains=query)
                    | Q(status__icontains=query)
                )
                print(
                    f"Resultados após filtro de pesquisa: {results.count()}")  # Depuração

            # Filtro adicional de data
            if data_inicio:
                # Converte para formato de data
                data_inicio = parse_date(data_inicio)
                results = results.filter(data_atualizacao__gte=data_inicio)

            if data_fim:
                # Converte para formato de data
                data_fim = parse_date(data_fim)
                results = results.filter(data_atualizacao__lte=data_fim)

            total_results = results.count()  # Captura o total de resultados encontrados

            return render(
                request,
                "search_results.html",
                {
                    "results": results,
                    "total_results": total_results,
                    "data_inicio": data_inicio,
                    "data_fim": data_fim,
                },
            )
        else:
            print("Agência não encontrada com o nome na sessão.")  # Depuração
            return redirect("login")  # Ou outra URL de erro
    else:
        print("Agência não definida na sessão.")  # Depuração
        return redirect("login")  # Ajuste conforme a URL apropriada


def pagefinal(request):
    return render(request, "pagefinal.html")

@login_required
def list_cards(request):
    query = request.GET.get("q")
    formularios_list = AAFormuUserer.objects.all().order_by("-data_criacao")
    if query:
        formularios_list = formularios_list.filter(
            Q(nome__icontains=query)
            | Q(email__icontains=query)
            | Q(username__icontains=query)
        )
    # Mostra 10 formulários por página
    paginator = Paginator(formularios_list, 10)
    page = request.GET.get("page")
    try:
        formularios = paginator.page(page)
    except PageNotAnInteger:
        formularios = paginator.page(1)
    except EmptyPage:
        formularios = paginator.page(paginator.num_pages)

    return render(
        request, "totalcards.html", {"formularios": formularios, "query": query}
    )

@csrf_exempt
@login_required
def style_forms(request):
    if request.method == "GET":
        return render(request, "style_forms.html")
    if request.method == "POST":
        data = json.loads(request.body.decode("utf-8"))
        agencia_nome = data.get("agencia")
        try:
            agencia_user = Agencia.objects.get(agencia=agencia_nome)
            print(f"[INFO] Agência encontrada: {agencia_user}")
        except Agencia.DoesNotExist:
            print("[ERROR] Agência não encontrada.")
            return JsonResponse({"error": "Agência não encontrada!"}, status=404)

        # Atualiza as cores associadas à agência
        agencia_user.agencia_header_color = data.get("agencia_header_color")
        agencia_user.agencia_background_color = data.get("agencia_background_color")
        agencia_user.agencia_explica_color = data.get("agencia_explica_color")
        agencia_user.agencia_footer_color = data.get("agencia_footer_color")
        agencia_user.save()

        return JsonResponse({"success": True})

    return redirect("painelagencia")

@csrf_exempt
@login_required
def temaagencia(request):
    if request.method == "GET":
        return render(request, "temaagencia.html")

    if request.method == "POST":
        try:
            # Decodifica a string de bytes para uma string utf-8 e carrega como
            # JSON
            data = json.loads(request.body.decode("utf-8"))
            tema_base = data.get("tema_base")
            agencia_nome = data.get("agencia")

            print("Dados recebidos:", data)

            try:
                agencia_user = Agencia.objects.get(agencia=agencia_nome)
                print(f"[INFO] Agência encontrada: {agencia_user}")
            except Agencia.DoesNotExist:
                print("[ERROR] Agência não encontrada.")
                return JsonResponse({"error": "Agência não encontrada!"}, status=404)

            # Atualiza o tema_base da agência
            agencia_user.tema_base = tema_base
            agencia_user.save()

            print(f"[INFO] Tema atualizado com sucesso: {agencia_user}")

            return JsonResponse({"status": "success"}, status=200)
        except json.JSONDecodeError:
            print("Erro ao decodificar JSON")
            return JsonResponse({"error": "Formato de JSON inválido!"}, status=400)

    return JsonResponse(
        {"status": "error", "message": "Método não permitido!"}, status=405
    )

@csrf_exempt
def formularioerro(request):
    if request.method == "POST":
        username = request.POST.get("username")

        data = json.loads(request.body)
        print("Dados recebidos:", data)
        url_loja = data.get("url_loja")
        telefone = data.get("telefone")
        email_suporte = data.get("email_suporte")
        empresa = data.get("empresa")

        token_senha = data.get("token_senha")
        chave_de_api = data.get("chave_de_api")
        chave_secreta = data.get("chave_secreta")
        business_hours = data.get("business_hours")
        nicho = data.get("nicho")
        cor = data.get("cor")
        username = data.get("username")
        try:
            formulario_user = AAFormuUserer.objects.get(username=username)
            print(f"[INFO] Formulário do usuário encontrado: {formulario_user}")
        except AAFormuUserer.DoesNotExist:
            print("[ERROR] Formulário do usuário não encontrado.")
            return JsonResponse({"error": "Usuário não encontrado!"}, status=404)
        # Se o formulário foi encontrado, salva os dados no banco de dados

        formulario_user.url_loja = url_loja
        formulario_user.telefone = telefone
        formulario_user.email_suporte = email_suporte
        formulario_user.empresa = empresa
        formulario_user.token_senha = token_senha
        formulario_user.chave_de_api = chave_de_api
        formulario_user.chave_secreta = chave_secreta
        formulario_user.business_hours = business_hours
        formulario_user.nicho = nicho
        formulario_user.cor = cor
        formulario_user.status = "produzindo"
        formulario_user.save()
        print(f"Formulário atualizado com sucesso: {formulario_user}")
        return JsonResponse(
            {"status": "success", "message": "Formulário atualizado com sucesso!"}
        )

    else:
        print("Método não permitido")
        return JsonResponse(
            {"status": "error", "message": "Método não permitido."}, status=405
        )

@login_required
def acesapp(request, section_id):
    formulario_user = get_object_or_404(AAFormuUserer, section_id=section_id)
    return render(request, "criarapp.html", {"formulario_user": formulario_user})

def remarcar(request):
    if request.method == "POST":
        username = request.POST.get("username")

        try:
            formulario_user = AAFormuUserer.objects.get(username=username)
            print(f"[INFO] Formulário do usuário encontrado: {formulario_user}")
        except AAFormuUserer.DoesNotExist:
            print("[ERROR] Formulário do usuário não encontrado.")
            return JsonResponse(
                {"status": "error", "message": "Usuário não encontrado!"}, status=404
            )

        # Atualiza os campos booleanos com base nos checkboxes
        formulario_user.produtos = "produtos" in request.POST
        formulario_user.tema = "tema" in request.POST
        formulario_user.banners = "banners" in request.POST
        formulario_user.colecoes = "colecoes" in request.POST
        formulario_user.paginas = "paginas" in request.POST
        formulario_user.politicas = "politicas" in request.POST
        formulario_user.prod_col = "prod_col" in request.POST
        formulario_user.estilizacao = "estilizacao" in request.POST
        formulario_user.status = "produzindo"
        formulario_user.save()  # Salva as alterações no modelo
        print(f"Status atualizado com sucesso: {formulario_user}")

        # Retorna uma resposta de sucesso para o frontend
        return JsonResponse(
            {"status": "success", "message": "Status atualizado com sucesso!"}
        )

    # Se qualquer outro método for utilizado, retorna um erro 405
    return JsonResponse({"error": "Método não permitido!"}, status=405)

def acesapperro(request, section_id):
    formulario_user = get_object_or_404(AAFormuUserer, section_id=section_id)
    return render(request, "criarapperro.html", {"formulario_user": formulario_user})

@csrf_exempt
def edtema(request):
    if request.method == "POST":
        # Coleta os dados do POST
        # Carregar dados JSON do corpo da requisição
        data = json.loads(request.body)
        tema_base = data.get("tema_base")
        username = data.get("username")
        # Exibe prints para depuração
        print(f"[DEBUG] Tema recebido: {tema_base}")

        # Busca o formulário do usuário com base no username
        try:
            formulario_user = AAFormuUserer.objects.get(username=username)
            print(f"[INFO] Formulário do usuário encontrado: {formulario_user}")
        except AAFormuUserer.DoesNotExist:
            print("[ERROR] Formulário do usuário não encontrado.")
            return JsonResponse({"error": "Usuário não encontrado!"}, status=404)

        # Se o formulário foi encontrado, salva os dados no banco de dados
        formulario_user.tema_base = tema_base
        formulario_user.save()

        print("[INFO] Tema Atualizado com sucesso no formulário do usuário.")

        # Retorna uma resposta de sucesso em JSON
        return JsonResponse({"success": "Tema Atualizado!"}, status=200)

    # Se qualquer outro método for utilizado, retorna um erro 405
    return JsonResponse({"error": "Método não permitido!"}, status=405)

@csrf_exempt
def criarapp(request):
    # Exibe o formulário na requisição GET
    if request.method == "GET":
        formulario_user = None
        return render(request, "criarapp.html", {"formulario_user": formulario_user})

    # Processa os dados na requisição POST
    if request.method == "POST":
        # Coleta os dados do POST
        # Carregar dados JSON do corpo da requisição
        data = json.loads(request.body)
        token_senha = data.get("PRIVATE_APP_PASSWORD")
        chave_de_api = data.get("API_KEY")
        chave_secreta = data.get("API_SECRET")
        username = data.get("username")
        url_loja = data.get("url_loja")

        # Busca o formulário do usuário com base no username
        try:
            formulario_user = AAFormuUserer.objects.get(username=username)
            print(f"[INFO] Formulário do usuário encontrado: {formulario_user}")
            # Atualiza os dados no formulário
            formulario_user.token_senha = token_senha
            formulario_user.chave_de_api = chave_de_api
            formulario_user.chave_secreta = chave_secreta
            formulario_user.save()
        except AAFormuUserer.DoesNotExist:
            print("[ERROR] Formulário do usuário não encontrado.")
            return JsonResponse({"error": "Usuário não encontrado!"}, status=404)

        # Define variáveis para a API com dados salvos no formulário
        SHOP_URL = url_loja
        PRIVATE_APP_PASSWORD = token_senha
        API_KEY = chave_de_api
        API_VERSION = "2024-10"
        print(
            f"SHOP_URL: {SHOP_URL}, API_KEY: {API_KEY}, PRIVATE_APP_PASSWORD: {PRIVATE_APP_PASSWORD}")

        # Conecta com a API Shopify para verificar o plano
        endpoint = f"https://{SHOP_URL}/admin/api/{API_VERSION}/shop.json"
        headers = {"Content-Type": "application/json"}
        auth = (API_KEY, PRIVATE_APP_PASSWORD)
        try:
            response = requests.get(endpoint, headers=headers, auth=auth)
            if response.status_code == 200:
                shop_info = response.json()
                plan_name = shop_info.get("shop", {}).get("plan_name")
                print(plan_name )
                
                # Salva o status do usuário com base no plano
                if plan_name in ["affiliate", "basic", "Shopify", "Advanced", "unlimited"]:
                    formulario_user.status = "produzindo"
                    
                    print("Plano ativo.")
                    success_message = {"success": "Plano Ativo"}
                else:
                    # Para qualquer plano que não seja ativo
                    formulario_user.status = "pendencia"
                    success_message = {"error": "Plano Não encontrado"}

                    # Aqui adicionamos a lógica para copiar os dados para
                    # Pendencia
                    copiar_para_pendencia(formulario_user)

                formulario_user.save()  # Salva o status final do usuário
                return JsonResponse(
                    success_message,
                    status=(
                        200
                        if plan_name in ["affiliate", "basic", "Shopify", "Advanced"]
                        else 400
                    ),
                )

            else:
                print(
                    f"Erro na resposta da API Shopify: {response.status_code}")
                return JsonResponse({"error": "Erro ao conectar com a Shopify"},status=response.status_code,)

        except requests.exceptions.RequestException as e:
            print(f"Erro ao conectar à API da Shopify: {e}")
            return JsonResponse({"error": "Erro na conexão com a Shopify"}, status=500)

    # Se qualquer outro método for utilizado, retorna um erro 405
    return JsonResponse({"error": "Método não permitido!"}, status=405)

def copiar_para_pendencia(formulario_user):
    # Cria uma cópia dos dados do formulário atual para o modelo Pendencia
    pendencia = Pendencia(
        agencia=formulario_user.agencia,
        username=formulario_user.username,
        # Adicione outros campos conforme necessário
    )
    pendencia.save()
    print(f"[INFO] Dados copiados para Pendencia: {pendencia}")

@login_required
def CartoesView(request):
    perfil_usuario = get_object_or_404(PerfilUsuario, user=request.user)
    agencia = perfil_usuario.agencia

    # Filtra os cartões de acordo com a agência do usuário logado e ordena pela última atualização
    cartoes = AAFormuUserer.objects.filter(agencia=agencia).order_by('-data_atualizacao')

    # Contagem dos cartões por status
    total_cartoes = cartoes.count()
    total_reservado = cartoes.filter(status="reservado").count()
    total_respondido = cartoes.filter(status="respondido").count()
    total_pendencia = cartoes.filter(status="pendencia").count()
    total_iniciado = cartoes.filter(status="iniciado").count()
    total_produzindo = cartoes.filter(status="produzindo").count()
    total_finalizado = cartoes.filter(status="finalizado").count()
    total_erro = cartoes.filter(status="erro").count()

    # Verificação de escassez de cartões reservados
    escassez_mensagem = None
    if total_reservado <= 9:
        escassez_mensagem = f"Restam apenas {total_reservado} cartões reservados!"
    elif total_reservado <= 5:
        escassez_mensagem = f"Sua agência possui apenas {total_reservado} cartões disponíveis!"
    elif total_reservado == 0:
        escassez_mensagem = "Seus cartões acabaram! Adquira mais para continuar produzindo!"

    context = {
        "cartoes": cartoes,
        "total_cartoes": total_cartoes,
        "total_reservado": total_reservado,
        "total_respondido": total_respondido,
        "total_iniciado": total_iniciado,
        "total_pendencia": total_pendencia,
        "total_produzindo": total_produzindo,
        "total_finalizado": total_finalizado,
        "total_erro": total_erro,
        "escassez_mensagem": escassez_mensagem,
    }

    return render(request, "cartoes.html", context)


@login_required
def atualizar_cartoes(request):
    perfil_usuario = get_object_or_404(PerfilUsuario, user=request.user)
    agencia = perfil_usuario.agencia

    # Filtra os cartões de acordo com a agência do usuário logado
    cartoes = AAFormuUserer.objects.filter(agencia=agencia)

    # Contagem dos cartões por status
    total_cartoes = cartoes.count()
    total_reservado = cartoes.filter(status="reservado").count()
    total_respondido = cartoes.filter(status="respondido").count()
    total_iniciado = cartoes.filter(status="iniciado").count()
    total_pendencia = cartoes.filter(status="pendencia").count()
    total_produzindo = cartoes.filter(status="produzindo").count()
    total_finalizado = cartoes.filter(status="finalizado").count()
    total_erro = cartoes.filter(status="erro").count()

    # Preparando a resposta para o AJAX
    cartoes_data = []
    for cartao in cartoes:
        cartoes_data.append(
            {
                "username": cartao.username,
                "status": cartao.status,
                "nome": cartao.nome,
                "empresa": cartao.empresa,
            }
        )

    return JsonResponse(
        {
            "cartoes": cartoes_data,
            "total_reservado": total_reservado,
            "total_respondido": total_respondido,
            "total_iniciado": total_iniciado,
            "total_pendencia": total_pendencia,
            "total_produzindo": total_produzindo,
            "total_finalizado": total_finalizado,
            "total_erro": total_erro,
        }
    )

@login_required
def DetalhesCartaoView(request, section_id):
    formulario_user = get_object_or_404(AAFormuUserer, section_id=section_id)
    return render(request, "detalhes_cartao.html", {"formulario_user": formulario_user})

@login_required
def detalhes(request, section_id):
    formulario_user = get_object_or_404(AAFormuUserer, section_id=section_id)
    return render(request, "detalhes.html", {"formulario_user": formulario_user})

@login_required
def finalizados(request, section_id):
    formulario_user = get_object_or_404(AAFormuUserer, section_id=section_id)
    return render(request, "finalizados.html", {"formulario_user": formulario_user})

def pendencias(request, section_id):
    if request.method == "GET":
        formulario_user = get_object_or_404(AAFormuUserer, section_id=section_id)
        pendencia = Pendencia.objects.filter(username=formulario_user.username).first() 

        # Passa apenas a pendência específica no contexto
        context = {"pendencia": pendencia, "formulario_user": formulario_user}
        return render(request, "pendencias.html", context)

def adic_cartoes(request):
    return render(request, "adic-cartoes.html")

def avatar(request):
    avatars = range(1, 31)  # Gera uma lista de números de 1 a 30
    return render(request, "painel_agencia.html", {"avatars": avatars})


############################ 
def coresgenerica(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'coresgenerica.html', {'agencia': agencia})
    
def coressaude(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'coressaudehtml', {'agencia': agencia})

def corescasa(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'corescasa.html', {'agencia': agencia})
    

def coreseletronicos(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'coreseletronicos.html', {'agencia': agencia})

def coreskids(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'coreskids.html', {'agencia': agencia})

def corespets(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'corespets.html', {'agencia': agencia})
    

def coresfitness(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'coresfitness.html', {'agencia': agencia})

def coresfeminino(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'coresfeminino.html', {'agencia': agencia})


def coresmasculino(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'coresmasculino.html', {'agencia': agencia})


def coresmoda(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'coresmoda.html', {'agencia': agencia})


def coresrelogios(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'coresrelogios.html', {'agencia': agencia})


def coresoculos(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'coresoculos.html', {'agencia': agencia})


############################ 
def paislatan(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'agencias/paislatan.html', {'agencia': agencia})

def paisbras(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'agencias/paisbras.html', {'agencia': agencia})

def paiseua(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'agencias/paiseua.html', {'agencia': agencia})

def paisespan(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'agencias/paisespan.html', {'agencia': agencia})

def paisport(request, agencia_id):
    agencia = get_object_or_404(Agencia, id=agencia_id)
    return render(request, 'agencias/paisport.html', {'agencia': agencia})


############################
############################################################
                                                #########################   FUNÇÕES DE  PRODUÇÃO  
############################################################


def produz(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            username = data.get("username")

            # Armazenar public_ids na cache por 1 hora
            cache.set(f"public_ids_{username}", data.get("public_ids", []), timeout=3600)

            result = processar_usuario.delay(username)

            return JsonResponse(
                {"status": "success", "message": "Processamento iniciado!", "task_id": result.id}
            )
        except Exception as e:
            return JsonResponse(
                {"status": "error", "message": f"Erro ao iniciar o processamento: {str(e)}"},
                status=500,
            )

    return JsonResponse({"status": "error", "message": "Método não permitido."}, status=405)

# def produz(request):
#     if request.method == "POST":
#         try:
#             data = json.loads(request.body)
#             username = data.get("username")
#             print (data)

#             # Disparar a tarefa em segundo plano
#             result = processar_usuario.delay(username)
            
#             return JsonResponse(
#                 {
#                     "status": "success",
#                     "message": "Processamento iniciado!",
#                     "task_id": result.id,
#                 }
#             )

#         except Exception as e:
#             return JsonResponse(
#                 {
#                     "status": "error",
#                     "message": f"Erro ao iniciar o processamento: {str(e)}",
#                 },
#                 status=500,
#             )

#     return JsonResponse(
#         {"status": "error", "message": "Método não permitido."}, status=405
#     )


def produzir(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            print("Dados recebidos:", data)
            username = data.get("username")

            # Busca o usuário com base no username
            formulario_user = AAFormuUserer.objects.filter(username=username).first()

            if formulario_user:
                # Atualiza o status no banco de dados
                formulario_user.status = "iniciado"
                formulario_user.save()  # Salva a alteração no banco de dados
                print("Status atualizado para 'iniciado'")

                SHOP_URL = formulario_user.url_loja
                PRIVATE_APP_PASSWORD = formulario_user.token_senha
                API_KEY = formulario_user.chave_de_api
                API_SECRET = formulario_user.chave_secreta
                estilo = formulario_user.nicho
                pais = formulario_user.pais
                tema_base = formulario_user.tema_base
                paleta = formulario_user.cor
                telefone = formulario_user.telefone
                email_suporte = formulario_user.email_suporte
                empresa = formulario_user.empresa
                business_hours = formulario_user.business_hours
                whatsapp = str(telefone)

                API_VERSION = "2024-07"
                cod_pais = {
                    "EUA": "+1",
                    "Argentina": "+54",
                    "Bolívia": "+591",
                    "Brasil": "+55",
                    "Chile": "+56",
                    "Colômbia": "+57",
                    "Equador": "+593",
                    "Paraguai": "+595",
                    "Peru": "+51",
                    "Uruguai": "+598",
                    "Venezuela": "+58",
                    "Paquistão": "+92",
                    "Índia": "+91",
                    "Portugal": "+351",
                    "México": "+52",
                    "Espanha": "+34",
                }
                if pais in cod_pais:
                    tel_whats = cod_pais[pais] + str(telefone)
                    whatsapp_completo = cod_pais[pais] + whatsapp

                ##############################################################
                #### CORES ######
                if paleta == "Paleta de Cor 1":  # LARANJA
                    # barra de anúncio
                    cor1 = "#BB6A09"
                    cor2 = "#FE8B01"
                    textcolor = "#FFFFFF"

                    # Cabeçalho/barra colorida
                    degrade1 = "#BB6A09"
                    degrade2 = "#BB6A09"
                    degrade3 = "#BB6A09"
                    degrade4 = "#BB6A09"

                    # Cores gerais
                    # Geral
                    acent = "#0ABF58"
                    link_color = "#0ABF58"
                    primary_button_background = "#0ABF58"
                    secondary_button_background = "#0ABF58"

                    # Rodapé
                    footer_background = "#FE8B01"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#FE8B01"
                    header_text_color = "#F7F7F7"
                    header_light_text_color = "#F7F7F7"
                    accent_color = "#0ABF58"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                    product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # footer
                    degradefooter1 = "#BB6A09"
                    degradefooter2 = "#BB6A09"
                    degradefooter3 = "#BB6A09"
                    degradefooter4 = "#BB6A09"
                    #############################

                elif paleta == "Paleta de Cor 2":  # VERDE
                    # barra de anúncio
                    cor1 = "#60AD3F"
                    cor2 = "#0097B2"
                    textcolor = "#FFFFFF"

                    # Cabeçalho/barra colorida
                    degrade1 = "#7ED957"
                    degrade2 = "#0097B2"
                    degrade3 = "#0097B2"
                    degrade4 = "#7ED957"

                    # Cores gerais
                    # Geral
                    acent = "#0097B2"
                    link_color = "#860D7E"
                    primary_button_background = "#860D7E"
                    secondary_button_background = "#860D7E"

                    # Rodapé
                    footer_background = "7ED957"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#7ED957"
                    header_text_color = "#F7F7F7"
                    header_light_text_color = "#F7F7F7"
                    accent_color = "#0097B2"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                    product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # footer
                    degradefooter1 = "#7ED957"
                    degradefooter2 = "#0097B2"
                    degradefooter3 = "#0097B2"
                    degradefooter4 = "#7ED957"
                    #############################

                elif paleta == "Paleta de Cor 3":  # ROXO
                    # barra de anúncio
                    cor1 = "#860D7E"
                    cor2 = "#600D5A"
                    textcolor = "#FFFFFF"

                    # Cabeçalho/barra colorida
                    degrade1 = "#D600FF"
                    degrade2 = "#8610D8"
                    degrade3 = "#D600FF"
                    degrade4 = "#8610D8"

                    # Cores gerais
                    # Geral
                    acent = "#860D7E"
                    link_color = "#860D7E"
                    primary_button_background = "#860D7E"
                    secondary_button_background = "#860D7E"

                    # Rodapé
                    footer_background = "600D5A "
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#600D5A"
                    header_text_color = "#F7F7F7"
                    header_light_text_color = "#F7F7F7"
                    accent_color = "#860D7E"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                    product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # footer
                    degradefooter1 = "#D600FF"
                    degradefooter2 = "#8610D8"
                    degradefooter3 = "#D600FF"
                    degradefooter4 = "#8610D8"
                    #############################

                elif paleta == "Paleta de Cor 4":  # PRETO
                    # barra de anúncio
                    cor1 = "#393939"
                    cor2 = "#000000"
                    textcolor = "#FFFFFF"

                    # Cabeçalho/barra colorida
                    degrade1 = "#393939"
                    degrade2 = "#393939"
                    degrade3 = "#393939"
                    degrade4 = "#393939"

                    # Cores gerais
                    # Geral
                    acent = "#393939"
                    link_color = "#393939"
                    primary_button_background = "#393939"
                    secondary_button_background = "#393939"

                    # Rodapé
                    footer_background = "#000000"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#000000"
                    header_text_color = "#F7F7F7"
                    header_light_text_color = "#F7F7F7"
                    accent_color = "#393939"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                    product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # footer
                    degradefooter1 = "#393939"
                    degradefooter2 = "#393939"
                    degradefooter3 = "#393939"
                    degradefooter4 = "#393939"
                    #############################

                elif paleta == "Paleta de Cor 5":  # AZUL
                    # barra de anúncio
                    cor1 = "#5DE0E6"
                    cor2 = "#073D86"
                    textcolor = "#FFFFFF"

                    # Cabeçalho/barra colorida
                    degrade1 = "#5DE0E6"
                    degrade2 = "#004AAD"
                    degrade3 = "#5DE0E6"
                    degrade4 = "#004AAD"

                    # Cores gerais
                    # Geral
                    acent = "#5DE0E6"
                    link_color = "#5DE0E6"
                    primary_button_background = "#5DE0E6"
                    secondary_button_background = "#5DE0E6"

                    # Rodapé
                    footer_background = "#004AAD"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#004AAD"
                    header_text_color = "#F7F7F7"
                    header_light_text_color = "#F7F7F7"
                    accent_color = "#5DE0E6"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                    product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # footer
                    degradefooter1 = "#5DE0E6"
                    degradefooter2 = "#004AAD"
                    degradefooter3 = "#5DE0E6"
                    degradefooter4 = "#004AAD"
                    #############################

                elif paleta == "Paleta de Cor 6":  # VERMELHO
                    # barra de anúncio
                    cor1 = "#B42929"
                    cor2 = "#740606"
                    textcolor = "#FFFFFF"

                    # Cabeçalho/barra colorida
                    degrade1 = "#B42929"
                    degrade2 = "#B42929"
                    degrade3 = "#740606"
                    degrade4 = "#740606"

                    # Cores gerais
                    # Geral
                    acent = "#740606"
                    link_color = "#740606"
                    primary_button_background = "#B42929"
                    secondary_button_background = "#B42929"

                    # Rodapé
                    footer_background = "#000000"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#FF3131"
                    header_text_color = "#F7F7F7"
                    header_light_text_color = "#F7F7F7"
                    accent_color = "#740606"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                    product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # footer
                    degradefooter1 = "#B42929"
                    degradefooter2 = "#B42929"
                    degradefooter3 = "#740606"
                    degradefooter4 = "#740606"

                elif paleta == "Paleta de Cor 7":  # AZUL E AMARELO
                    # barra de anúncio
                    cor1 = "#0F0F33"
                    cor2 = "#0F0F33"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#D3991F"
                    degrade2 = "#D3991F"
                    degrade3 = "#D3991F"
                    degrade4 = "#D3991F"

                    # Cores gerais
                    # Geral
                    acent = "#0F0F33"
                    link_color = "#E7DCDC"
                    # Botão primário
                    primary_button_background = "#000000"
                    # Botão secundário
                    secondary_button_background = "#0F0F33"

                    # Rodapé
                    footer_background = "#0F0F33"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#F2F2F2"
                    header_text_color = "#0F0F33"
                    header_light_text_color = "#0F0F33"
                    accent_color = "#D3991F"

                    # Produtos
                    product_on_sale_accent = "#0F0F33"  # Barra economize
                    product_cor_do_preco_semdesc = "#0F0F33"  # Cor do preço
                    product_cor_do_preco = "#0F0F33"  # Cor do preço com desconto
                    product_cor_dos_titles = "#4D4D4B"  # Títulos das informações
                    product_in_stock_color = "#4D4D4B"  # Em estoque

                    # Footer
                    degradefooter1 = "#D3991F"
                    degradefooter2 = "#D3991F"
                    degradefooter3 = "#D3991F"
                    degradefooter4 = "#D3991F"

                elif paleta == "Paleta de Cor 8":  # CINZA ESCURO CINZA CLARO
                    # barra de anúncio
                    cor1 = "#4D4D4B"
                    cor2 = "#4D4D4B"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#4D4D4B"
                    degrade2 = "#4D4D4B"
                    degrade3 = "#4D4D4B"
                    degrade4 = "#4D4D4B"

                    # Cores gerais
                    # Geral
                    acent = "#4D4D4B"
                    link_color = "#4D4D4B"
                    # Botão primário
                    primary_button_background = "#4D4D4B"
                    # Botão secundário
                    secondary_button_background = "#4D4D4B"

                    # Rodapé
                    footer_background = "#4D4D4B"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#F2F2F2"
                    header_text_color = "#4D4D4B"
                    header_light_text_color = "#4D4D4B"
                    accent_color = "#AFAFAF"

                    # Produtos
                    product_on_sale_accent = "#4D4D4B"  # Barra economize
                    product_cor_do_preco_semdesc = "#4D4D4B"  # Cor do preço
                    product_cor_do_preco = "#4D4D4B"  # Cor do preço com desconto
                    product_cor_dos_titles = "#4D4D4B"  # Títulos das informações
                    product_in_stock_color = "#4D4D4B"  # Em estoque

                    # Footer
                    degradefooter1 = "#AFAFAF"
                    degradefooter2 = "#AFAFAF"
                    degradefooter3 = "#AFAFAF"
                    degradefooter4 = "#AFAFAF"

                elif paleta == "Paleta de Cor 9":  # MARROM E AZUL
                    # barra de anúncio
                    cor1 = "#042138"
                    cor2 = "#042138"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#042138"
                    degrade2 = "#042138"
                    degrade3 = "#042138"
                    degrade4 = "#042138"

                    # Cores gerais
                    # Geral
                    acent = "#042138"
                    link_color = "#042138"
                    # Botão primário
                    primary_button_background = "#042138"
                    # Botão secundário
                    secondary_button_background = "#042138"

                    # Rodapé
                    footer_background = "#042138"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#F2F2F2"
                    header_text_color = "#042138"
                    header_light_text_color = "#042138"
                    accent_color = "#9D7638"

                    # Produtos
                    product_on_sale_accent = "#042138"  # Barra economize
                    product_cor_do_preco_semdesc = "#042138"  # Cor do preço
                    product_cor_do_preco = "#042138"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#042138"  # Em estoque

                    # Footer
                    degradefooter1 = "#9D7638"
                    degradefooter2 = "#9D7638"
                    degradefooter3 = "#9D7638"
                    degradefooter4 = "#9D7638"

                elif paleta == "Paleta de Cor 10":  # MARROM ESCURO E BEGE
                    # barra de anúncio
                    cor1 = "#2D251C"
                    cor2 = "#2D251C"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#AD9D95"
                    degrade2 = "#AD9D95"
                    degrade3 = "#AD9D95"
                    degrade4 = "#AD9D95"

                    # Cores gerais
                    # Geral
                    acent = "#2D251C"
                    link_color = "#2D251C"
                    # Botão primário
                    primary_button_background = "#2D251C"
                    # Botão secundário
                    secondary_button_background = "#2D251C"

                    # Rodapé
                    footer_background = "#2D251C"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#F2F2F2"
                    header_text_color = "#2D251C"
                    header_light_text_color = "#2D251C"
                    accent_color = "#AD9D95"

                    # Produtos
                    product_on_sale_accent = "#042138"  # Barra economize
                    product_cor_do_preco_semdesc = "#042138"  # Cor do preço
                    product_cor_do_preco = "#042138"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#042138"  # Em estoque

                    # Footer
                    degradefooter1 = "#AD9D95"
                    degradefooter2 = "#AD9D95"
                    degradefooter3 = "#AD9D95"
                    degradefooter4 = "#AD9D95"

                elif paleta == "Paleta de Cor 11":  # MARROM e Cinza
                    # barra de anúncio
                    cor1 = "#BE8F59"
                    cor2 = "#BE8F59"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#BE8F59"
                    degrade2 = "#BE8F59"
                    degrade3 = "#BE8F59"
                    degrade4 = "#BE8F59"

                    # Cores gerais
                    # Geral
                    acent = "#BE8F59"
                    link_color = "#BE8F59"
                    # Botão primário
                    primary_button_background = "#BE8F59"
                    # Botão secundário
                    secondary_button_background = "#BE8F59"

                    # Rodapé
                    footer_background = "#555555"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#F2F2F2"
                    header_text_color = "#BE8F59"
                    header_light_text_color = "#BE8F59"
                    accent_color = "#CDCDCD"

                    # Produtos
                    product_on_sale_accent = "#BE8F59"  # Barra economize
                    product_cor_do_preco_semdesc = "#BE8F59"  # Cor do preço
                    product_cor_do_preco = "#BE8F59"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # Footer
                    degradefooter1 = "#BE8F59"
                    degradefooter2 = "#BE8F59"
                    degradefooter3 = "#BE8F59"
                    degradefooter4 = "#BE8F59"

                elif paleta == "Paleta de Cor 12":  # PRETO E BRANCO
                    # barra de anúncio
                    cor1 = "#000000"
                    cor2 = "#000000"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#000000"
                    degrade2 = "#000000"
                    degrade3 = "#000000"
                    degrade4 = "#000000"

                    # Cores gerais
                    # Geral
                    acent = "#000000"
                    link_color = "#E7DCDC"
                    # Botão primário
                    primary_button_background = "#000000"
                    # Botão secundário
                    secondary_button_background = "#000000"

                    # Rodapé
                    footer_background = "#000000"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#F2F2F2"
                    header_text_color = "#4D4D4B"
                    header_light_text_color = "#000000"
                    accent_color = "#000000"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#000000"  # Cor do preço
                    product_cor_do_preco = "#000000"  # Cor do preço com desconto
                    product_cor_dos_titles = "#4D4D4B"  # Títulos das informações
                    product_in_stock_color = "#4D4D4B"  # Em estoque

                    # Footer
                    degradefooter1 = "#AFAFAF"
                    degradefooter2 = "#AFAFAF"
                    degradefooter3 = "#AFAFAF"
                    degradefooter4 = "#AFAFAF"

                elif paleta == "Paleta de Cor 13":  # verde e azul
                    # barra de anúncio
                    cor1 = "#516075"
                    cor2 = "#516075"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#67A896"
                    degrade2 = "#67A896"
                    degrade3 = "#67A896"
                    degrade4 = "#67A896"

                    # Cores gerais
                    # Geral
                    acent = "#67A896"
                    link_color = "#67A896"
                    # Botão primário
                    primary_button_background = "#67A896"
                    # Botão secundário
                    secondary_button_background = "#67A896"

                    # Rodapé
                    footer_background = "#516075"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#FFFFFF"
                    header_text_color = "#516075"
                    header_light_text_color = "#516075"
                    accent_color = "#740606"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                    product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # Footer
                    degradefooter1 = "#67A896"
                    degradefooter2 = "#67A896"
                    degradefooter3 = "#67A896"
                    degradefooter4 = "#67A896"

                elif paleta == "Paleta de Cor 14":  # Verde Escuro e Bege
                    # barra de anúncio
                    cor1 = "#17453F"
                    cor2 = "#17453F"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#E8E3D7"
                    degrade2 = "#E8E3D7"
                    degrade3 = "#E8E3D7"
                    degrade4 = "#E8E3D7"

                    # Cores gerais
                    # Geral
                    acent = "#17453F"
                    link_color = "#17453F"
                    # Botão primário
                    primary_button_background = "#17453F"
                    # Botão secundário
                    secondary_button_background = "#17453F"

                    # Rodapé
                    footer_background = "#17453F"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#FFFFFF"
                    header_text_color = "#17453F"
                    header_light_text_color = "#17453F"
                    accent_color = "#E8E3D7"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                    product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # Footer
                    degradefooter1 = "#67A896"
                    degradefooter2 = "#67A896"
                    degradefooter3 = "#67A896"
                    degradefooter4 = "#67A896"

                elif paleta == "Paleta de Cor 15":  # Verde e Amarelo
                    # barra de anúncio
                    cor1 = "#5EB64E"
                    cor2 = "#5EB64E"
                    textcolor = "#FFFFFF"

                    # Cabeçalho/barra colorida
                    degrade1 = "#E8E3D7"
                    degrade2 = "#E8E3D7"
                    degrade3 = "#E8E3D7"
                    degrade4 = "#E8E3D7"

                    # Cores gerais
                    # Geral
                    acent = "#5EB64E"
                    link_color = "#5EB64E"
                    # Botão primário
                    primary_button_background = "#5EB64E"
                    # Botão secundário
                    secondary_button_background = "#5EB64E"

                    # Rodapé
                    footer_background = "#5EB64E"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#FFFFFF"
                    header_text_color = "#5EB64E"
                    header_light_text_color = "#5EB64E"
                    accent_color = "#E5C92C"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                    product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # Footer
                    degradefooter1 = "#E5C92C"
                    degradefooter2 = "#E5C92C"
                    degradefooter3 = "#E5C92C"
                    degradefooter4 = "#E5C92C"

                elif paleta == "Paleta de Cor 16":  # PRETO E BRANCO MODA
                    # barra de anúncio
                    cor1 = "#000000"
                    cor2 = "#000000"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#000000"
                    degrade2 = "#000000"
                    degrade3 = "#000000"
                    degrade4 = "#000000"

                    # Cores gerais
                    # Geral
                    acent = "#000000"
                    link_color = "#000000"
                    # Botão primário
                    primary_button_background = "#000000"
                    # Botão secundário
                    secondary_button_background = "#000000"

                    # Rodapé
                    footer_background = "#F2F2F2"
                    footer_text_color = "#000000"

                    # Cabeçalho
                    header_background = "#FFFFFF"
                    header_text_color = "#000000"
                    header_light_text_color = "#000000"
                    accent_color = "#000000"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#000000"  # Cor do preço
                    product_cor_do_preco = "#000000"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # Footer
                    degradefooter1 = "#000000"
                    degradefooter2 = "#000000"
                    degradefooter3 = "#000000"
                    degradefooter4 = "#000000"

                elif paleta == "Paleta de Cor 17":  #  Azul escuro e cinza
                    # barra de anúncio
                    cor1 = "#001438"
                    cor2 = "#001438"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#9197AA"
                    degrade2 = "#9197AA"
                    degrade3 = "#9197AA"
                    degrade4 = "#9197AA"

                    # Cores gerais
                    # Geral
                    acent = "#001438"
                    link_color = "#001438"
                    # Botão primário
                    primary_button_background = "#001438"
                    # Botão secundário
                    secondary_button_background = "#001438"

                    # Rodapé
                    footer_background = "#F2F2F2"
                    footer_text_color = "#001438"

                    # Cabeçalho
                    header_background = "#FFFFFF"
                    header_text_color = "#001438"
                    header_light_text_color = "#001438"
                    accent_color = "#9197AA"

                    # Produtos
                    product_on_sale_accent = "#001438"  # Barra economize
                    product_cor_do_preco_semdesc = "#001438"  # Cor do preço
                    product_cor_do_preco = "#001438"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # Footer
                    degradefooter1 = "#9197AA"
                    degradefooter2 = "#9197AA"
                    degradefooter3 = "#9197AA"
                    degradefooter4 = "#9197AA"

                elif paleta == "Paleta de Cor 18":  #  Bege e Marrom
                    # barra de anúncio
                    cor1 = "#9A5F43"
                    cor2 = "#9A5F43"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#D1B1A3"
                    degrade2 = "#D1B1A3"
                    degrade3 = "#D1B1A3"
                    degrade4 = "#D1B1A3"

                    # Cores gerais
                    # Geral
                    acent = "#9A5F43"
                    link_color = "#9A5F43"
                    # Botão primário
                    primary_button_background = "#9A5F43"
                    # Botão secundário
                    secondary_button_background = "#9A5F43"

                    # Rodapé
                    footer_background = "#F2F2F2"
                    footer_text_color = "#9A5F43"

                    # Cabeçalho
                    header_background = "#FFFFFF"
                    header_text_color = "#9A5F43"
                    header_light_text_color = "#9A5F43"
                    accent_color = "#D1B1A3"

                    # Produtos
                    product_on_sale_accent = "#9A5F43"  # Barra economize
                    product_cor_do_preco_semdesc = "#9A5F43"  # Cor do preço
                    product_cor_do_preco = "#9A5F43"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # Footer
                    degradefooter1 = "#D1B1A3"
                    degradefooter2 = "#D1B1A3"
                    degradefooter3 = "#D1B1A3"
                    degradefooter4 = "#D1B1A3"

                elif paleta == "Paleta de Cor 19":  # Cinza e Marrom
                    # barra de anúncio
                    cor1 = "#82746C"
                    cor2 = "#82746C"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#82746C"
                    degrade2 = "#82746C"
                    degrade3 = "#82746C"
                    degrade4 = "#82746C"

                    # Cores gerais
                    # Geral
                    acent = "#82746C"
                    link_color = "#82746C"
                    # Botão primário
                    primary_button_background = "#82746C"
                    # Botão secundário
                    secondary_button_background = "#82746C"
                    # Rodapé
                    footer_background = "#82746C"
                    footer_text_color = "#FFFFFF"

                    # Cabeçalho
                    header_background = "#FFFFFF"
                    header_text_color = "#5E5E5E"
                    header_light_text_color = "#5E5E5E"
                    accent_color = "#CFCAC4"

                    # Produtos
                    product_on_sale_accent = "#82746C"  # Barra economize
                    product_cor_do_preco_semdesc = "#82746C"  # Cor do preço
                    product_cor_do_preco = "#82746C"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # Footer
                    degradefooter1 = "#CFCAC4"
                    degradefooter2 = "#CFCAC4"
                    degradefooter3 = "#CFCAC4"
                    degradefooter4 = "#CFCAC4"

                elif paleta == "Paleta de Cor 20":  # ROSA E CINZA
                    # barra de anúncio
                    cor1 = "#5A5E61"
                    cor2 = "#5A5E61"
                    textcolor = "#FFFFFF"
                    # Cabeçalho/barra colorida
                    degrade1 = "#5A5E61"
                    degrade2 = "#5A5E61"
                    degrade3 = "#5A5E61"
                    degrade4 = "#5A5E61"

                    # Cores gerais
                    # Geral
                    acent = "#5A5E61"
                    link_color = "#5A5E61"
                    # Botão primário
                    primary_button_background = "#5A5E61"
                    # Botão secundário
                    secondary_button_background = "#5A5E61"
                    # Rodapé
                    footer_background = "#F2F2F2"
                    footer_text_color = "#5A5E61"

                    # Cabeçalho
                    header_background = "#FFFFFF"
                    header_text_color = "#5A5E61"
                    header_light_text_color = "#5A5E61"
                    accent_color = "#CF97A0"

                    # Produtos
                    product_on_sale_accent = "#000000"  # Barra economize
                    product_cor_do_preco_semdesc = "#000000"  # Cor do preço
                    product_cor_do_preco = "#000000"  # Cor do preço com desconto
                    product_cor_dos_titles = "#656866"  # Títulos das informações
                    product_in_stock_color = "#00D864"  # Em estoque

                    # Footer
                    degradefooter1 = "#CF97A0"
                    degradefooter2 = "#CF97A0"
                    degradefooter3 = "#CF97A0"
                    degradefooter4 = "#CF97A0"
                
                ##############################################################
                ##############################################################
                colecoes = []
                produtos = None
                ##########################                                                          
                ##########################
                if pais == "Brasil":
                    if business_hours == "hora1":
                        antend = "Segunda a Sexta das 9:00 as 18:00"
                    elif business_hours == "hora2": 
                        antend = "Segunda a Sexta das 9:00 as 22:00"
                    elif business_hours == "hora3":
                        antend = "Segunda a Sábado das 9:00 as 18:00"
                    elif business_hours == "hora4":
                        antend = "Segunda a Sábado das 9:00 as 22:00"
                    elif business_hours == "hora5":
                        antend = "Segunda a Sexta 9:00 as 18:00, Sábado das 9:00 as 12:00"
                    elif business_hours == "hora6":
                        antend = "Segunda a Sexta 9:00 as 22:00, Sábado das 9:00 as 12:00"
                    elif business_hours == "hora7":
                        antend = "Todos os dias das 9:00 as 18:00"
                    elif business_hours == "hora8":
                        antend = "Todos os dias das 9:00 as 22:00"
                    if tema_base == "Dropmeta":
                        Tema = "https://github.com/ronildob/TEMAS2025/raw/refs/heads/main/tema_pt_br%2001-2025.zip"
                        #Tema = "https://github.com/ronildob/temas/raw/refs/heads/main/portuguesBRoficial.zip"
                        footer = f"""<p>
                                        <strong>E-mail:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a>
                                    </p>
                                    <p>
                                        <strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{cod_pais[pais]} {whatsapp}</a>
                                    </p>
                                    <p>
                                        <strong>Hor\u00e1rio de atendimento: </strong>{antend}
                                    </p>"""
                        titlezap = f"Fale Conosco \n<span>\nO melhor suporte é com a {empresa}!\n</span>"
                    elif tema_base == "Evolution":
                        Tema = "https://github.com/ronildob/Themes/raw/refs/heads/main/TEMA_GL_NOVO.zip"
                        footer = f"""<p></p>
                                            <p>
                                                <strong>SERVIÇO DE SAC DISPONÍVEL</strong>
                                            </p>
                                            <p>
                                                <strong>E-mail:</strong> <a href="mailto:{email_suporte}">{email_suporte}</a>
                                            </p>
                                            <p>
                                                <strong>Atendimento:</strong><br/>\u00a0<a href="https://wa.me/{whatsapp_completo}">{cod_pais[pais]} {whatsapp}</a>
                                            </p>
                                    <p></p>"""
                    paginas = {
                            "Rastrear Pedido": "static/home/politicas/Rastrearpedido.txt",
                            "Contrato de E-Commerce": "static/home/politicas/Contrato de E-Commerce.docx",
                            "Termos de Uso": "static/home/politicas/Termos de Uso.docx",
                            "Prazo de Entrega": "static/home/politicas/Prazo de Entrega.docx",
                            "Política de Privacidade": "static/home/politicas/Política de Privacidade.docx",
                            "Trocas ou Devoluções": "static/home/politicas/Trocas ou Devolução.docx",
                            "Sobre Nós": "static/home/politicas/Sobre Nós.docx",
                            "Pagamento Seguro": "static/home/politicas/Pagamento  Seguro.docx",
                        }                                    
                    politicas_para_atualizar = {
                        "Legal notice": "static/home/politicas/Contrato de E-Commerce.docx",
                        "Terms of service": "static/home/politicas/Termos de Uso.docx",
                        "Shipping policy": "static/home/politicas/Prazo de Entrega.docx",
                        "Refund policy": "static/home/politicas/Trocas ou Devolução.docx",
                        #"Privacy policy": "static/home/politicas/Política de Privacidade.docx",
                    }
                                                        
                    if estilo == "Genérica":   
                        df_path = "static/home/csv/PT_BR/30GENERICOS.csv"
                        pri = "Casa"
                        seg = "Eletrônicos"
                        ter = "Pets"
                        qua = "Fitness"
                        qui = "Kids"
                        sex = "Mais Vendidos"

                        primeiro = "Casa"
                        segundo = "Eletrônicos"
                        terceiro = "Pets"
                        quarto = "Fitness"
                        quinto = "Kids"
                        sexto = "Mais-Vendidos"
                        colecoes = [
                            "Casa", 
                            "Eletrônicos", 
                            "Pets", 
                            "Fitness", 
                            "Kids", 
                            "Mais Vendidos"
                            ]
                        
                        # urls = {
                        #     "desktop1": "",
                        #     "desktop2": "",
                        #     "desktop3": "",
                        #     "mobile1": "",
                        #     "mobile2": "",
                        #     "mobile3": "",
                        #     "capa1": "",
                        #     "capa2": "",
                        #     "capa3": "",
                        #     "capa4": "",
                        #     "capa5": "",
                        #     "capa6": "",
                        # }

                    elif estilo == "Eletrônicos":
                        df_path = "static/home/csv/PT_BR/30ELETRONICOS.csv"
                        pri = "Fones"
                        seg = "Drones"
                        ter = "Smartwatch"
                        qua = "Acessórios"
                        qui = "Produtos Gamer"
                        sex = "Mais Vendidos"

                        primeiro = "Fones"
                        segundo = "Drones"
                        terceiro = "Smartwatch"
                        quarto = "Acessórios"
                        quinto = "Produtos-Gamer"
                        sexto = "Mais-Vendidos"
                        colecoes = ["Fones", "Drones", "Smartwatch", "Acessórios", "Produtos Gamer", "Mais Vendidos"]

                        urls = {
                            "desktop1": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736255612/banner1_rmpi15.svg",
                            "desktop2": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736255719/banner2_1_f4tkw0.svg",
                            "desktop3": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736255617/banner3_f0aate.svg",
                            "mobile1": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275595/mobile1_fwcsuu.svg",
                            "mobile2": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275593/mobile2_uc1t6y.svg",
                            "mobile3": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275595/mobile1_fwcsuu.svg",
                            "capa1": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275591/capa1_vseevh.svg",
                            "capa2": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275590/capa2_jws3hn.svg",
                            "capa3": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275591/capa3_jryrkt.svg",
                            "capa4": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275591/capa4_pqm3qy.svg",
                            "capa5": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275590/capa5_mj8wbe.svg",
                            "capa6": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275590/capa6_qpbjav.svg",

                           }

                    elif estilo == "Kids":
                        df_path = "static/home/csv/PT_BR/30KIDS.csv"
                        pri = "Moda"
                        seg = "Acessórios"
                        ter = "Brinquedos"
                        qua = "Cuidados"
                        qui = "Maternidade"
                        sex = "Mais Vendidos"

                        primeiro = "Moda"
                        segundo = "Acessórios"
                        terceiro = "Brinquedos"
                        quarto = "Cuidados"
                        quinto = "Maternidade"
                        sexto = "Mais-Vendidos"
                        colecoes = [
                            "Moda", 
                            "Acessórios", 
                            "Brinquedos", 
                            "Cuidados", 
                            "Maternidade", 
                            "Mais Vendidos"
                            ]
                        # urls = {
                        #     "desktop1": "",
                        #     "desktop2": "",
                        #     "desktop3": "",
                        #     "mobile1": "",
                        #     "mobile2": "",
                        #     "mobile3": "",
                        #     "capa1": "",
                        #     "capa2": "",
                        #     "capa3": "",
                        #     "capa4": "",
                        #     "capa5": "",
                        #     "capa6": "",
                        # }
                    
                    elif estilo == "Casa":
                        df_path = "static/home/csv/PORTUGUES_PT/CASA_PORTUGAL.csv"

                        pri = "Acessórios"
                        seg = "Iluminação"
                        ter = "Decoração de Ambientes"
                        qua = "Organização e Armazenamento"
                        qui = "Aromaterapia e Bem-estar"
                        sex = "Produtos Sazonais"

                        primeiro = "Acessórios"
                        segundo = "Iluminação"
                        terceiro = "Decoração-de-Ambientes"
                        quarto = "Organização-e-Armazenamento"
                        quinto = "Aromaterapia-e-Bem-estar"
                        sexto = "Produtos-Sazonais"
                        colecoes = [
                            "Acessórios",
                            "Iluminação",
                            "Decoração de Ambientes",
                            "Organização e Armazenamento",
                            "Aromaterapia e Bem-estar",
                            "Produtos Sazonais",
                        ]                    
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                                          
                    elif estilo == "Pet":
                        df_path = "static/home/csv/PT_BR/30PETS.csv" 

                        pri = "Camas"
                        seg = "Acessórios"
                        ter = "Comedouros"
                        qua = "Brinquedos"
                        qui = "Roupas"
                        sex = "Mais Vendidos"

                        primeiro = "Camas"
                        segundo = "Acessórios"
                        terceiro = "Comedouros"
                        quarto = "Brinquedos"
                        quinto = "Roupas"
                        sexto = "Mais-Vendidos"
                        colecoes = [
                            "Camas", 
                            "Acessórios", 
                            "Comedouros", 
                            "Brinquedos", 
                            "Roupas", 
                            "Mais Vendidos"
                            ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Fitness":
                        df_path = "static/home/csv/PT_BR/30FITNESS.csv"
                        pri = "Moda"
                        seg = "Acessórios"
                        ter = "Calçados"
                        qua = "Corretores"
                        qui = "Produtos Esportivos"
                        sex = "Mais Vendidos"

                        primeiro = "Moda"
                        segundo = "Acessórios"
                        terceiro = "Calçados"
                        quarto = "Corretores"
                        quinto = "Produtos-Esportivos"
                        sexto = "Mais-Vendidos"
                        colecoes = [
                            "Moda", 
                            "Acessórios", 
                            "Calçados", 
                            "Corretores", 
                            "Produtos Esportivos", 
                            "Mais Vendidos"
                            ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Masculino":
                        df_path = "static/home/csv/PT_BR/30MASCULINO.csv"

                        pri = "Acessórios"
                        seg = "Ferramentas"
                        ter = "Pesca"
                        qua = "Vestuário"
                        qui = "Automotivos"
                        sex = "Mais Vendidos"

                        primeiro = "Acessórios"
                        segundo = "Ferramentas"
                        terceiro = "Pesca"
                        quarto = "Vestuário"
                        quinto = "Automotivos"
                        sexto = "Mais-Vendidos"
                        colecoes = [
                            "Acessórios", 
                            "Ferramentas", 
                            "Pesca", 
                            "Vestuário", 
                            "Automotivos", 
                            "Mais Vendidos"
                            ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                         
                    elif estilo == "Feminino":
                        df_path = "static/home/csv/PT_BR/30FEMININO.csv"

                        pri = "Moda"
                        seg = "Acessórios"
                        ter = "Make"
                        qua = "Escovas Alisadoras"
                        qui = "Saúde"
                        sex = "Mais Vendidos"

                        primeiro = "Moda"
                        segundo = "Acessórios"
                        terceiro = "Make"
                        quarto = "Escovas-Alisadoras"
                        quinto = "Saúde"
                        sexto = "Mais-Vendidos"
                        colecoes = [
                            "Moda", 
                            "Acessórios", 
                            "Make", 
                            "Escovas Alisadoras", 
                            "Saúde", 
                            "Mais Vendidos"
                            ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                                            
                    elif estilo == "Moda":
                        df_path = "static/home/csv/PT_BR/MODA.csv"
                        pri = "Acessórios"
                        seg = "Feminino"
                        ter = "Masculino"
                        qua = "Infantil"
                        qui = "Jóias"
                        sex = "Promoções"

                        primeiro = "Acessórios"
                        segundo = "Feminino"
                        terceiro = "Masculino"
                        quarto = "Infantil"
                        quinto = "Jóias"
                        sexto = "Promoções"
                        colecoes = [
                            "Acessórios",
                            "Feminino",
                            "Masculino",
                            "Infantil",
                            "Jóias",
                            "Promoções",
                        ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                        
                    elif estilo == "Saúde e Beleza":
                        df_path = "static/home/csv/PT_BR/SAUDEEBELEZA.csv"
                        pri = "Fitness e Exercício"
                        seg = "Beleza"
                        ter = "Mais Vendidos"
                        qua = "Relaxamento"
                        qui = "Cuidados Pessoais"
                        sex = "Monitoramento da Saúde"

                        primeiro = "Fitness-e-Exercício"
                        segundo = "Beleza"
                        terceiro = "Mais-Vendidos"
                        quarto = "Relaxamento"
                        quinto = "Cuidados-Pessoais"
                        sexto = "Monitoramento-da-Saúde"
                        colecoes = [
                            "Fitness e Exercício",
                            "Beleza",
                            "Mais Vendidos",
                            "Relaxamento",
                            "Cuidados Pessoais",
                            "Monitoramento da Saúde",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                ##########################                                                          
                ########################## 
                elif pais == "Portugal":
                    if business_hours == "hora1":
                        antend = "Segunda a Sexta das 9:00 as 18:00"
                    elif business_hours == "hora2":
                        antend = "Segunda a Sexta das 9:00 as 22:00"
                    elif business_hours == "hora3":
                        antend = "Segunda a Sábado das 9:00 as 18:00"
                    elif business_hours == "hora4":
                        antend = "Segunda a Sábado das 9:00 as 22:00"
                    elif business_hours == "hora5":
                        antend = (
                            "Segunda a Sexta 9:00 as 18:00, Sábado das 9:00 as 12:00"
                        )
                    elif business_hours == "hora6":
                        antend = (
                            "Segunda a Sexta 9:00 as 22:00, Sábado das 9:00 as 12:00"
                        )
                    elif business_hours == "hora7":
                        antend = "Todos os dias das 9:00 as 18:00"
                    elif business_hours == "hora8":
                        antend = "Todos os dias das 9:00 as 22:00"
                    ##############################
                    footer = f"""<p>
                                    <strong>E-mail:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a>
                                </p>
                                <p>
                                    <strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{cod_pais[pais]} {whatsapp}</a>
                                </p>
                                <p>
                                    <strong>Hor\u00e1rio de atendimento: </strong>{antend}
                                </p>"""
                    titlezap = f"Fale Conosco \n<span>\nO melhor suporte é com a {empresa}!\n</span>"
                    ##############################
                    Tema = "https://github.com/ronildob/temas/raw/refs/heads/main/portuguesPToficial.zip"
                    paginas = {
                        "Rastrear Pedido": "static/home/politicas/Rastrearpedido.txt",
                        "Contrato de E-Commerce": "static/home/politicas/Contrato de E-Commerce.docx",
                        "Termos de Uso": "static/home/politicas/Termos de Uso.docx",
                        "Prazo de Entrega": "static/home/politicas/Prazo de Entrega.docx",
                        "Política de Privacidade": "static/home/politicas/Política de Privacidade.docx",
                        "Trocas ou Devoluções": "static/home/politicas/Trocas ou Devolução.docx",
                        "Sobre Nós": "static/home/politicas/Sobre Nós.docx",
                        "Pagamento Seguro": "static/home/politicas/Pagamento  Seguro.docx",
                    }
                    politicas_para_atualizar = {
                        "Legal notice": "static/home/politicas/Contrato de E-Commerce.docx",
                        "Terms of service": "static/home/politicas/Termos de Uso.docx",
                        "Shipping policy": "static/home/politicas/Prazo de Entrega.docx",
                        # "Privacy policy": "static/home/politicas/Política de Privacidade.docx",
                        "Refund policy": "static/home/politicas/Trocas ou Devolução.docx",
                    }
                    ##############################

                    if estilo == "Genérica":
                        df_path = "static/home/csv/PORTUGUES_PT/30GENERICOSPT.csv"

                        pri = "Casa"
                        seg = "Eletrónica"
                        ter = "Pets"
                        qua = "Fitness"
                        qui = "Kids"
                        sex = "Más Vendidos"

                        primeiro = "Casa"
                        segundo = "Eletrónica"
                        terceiro = "Pets"
                        quarto = "Fitness"
                        quinto = "Kids"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Casa",
                            "Eletrónica",
                            "Pets",
                            "Fitness",
                            "Kids",
                            "Más Vendidos",
                        ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                         
                    elif estilo == "Eletrônicos":
                        df_path = "static/home/csv/PORTUGUES_PT/ELETRONICOPT.csv"

                        pri = "Fones"
                        seg = "Drones"
                        ter = "Smartwatch"
                        qua = "Acessórios"
                        qui = "Produtos Gamer"
                        sex = "Más Vendidos"

                        primeiro = "Fones"
                        segundo = "Drones"
                        terceiro = "Smartwatch"
                        quarto = "Acessórios"
                        quinto = "Produtos-Gamer"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Fones",
                            "Drones",
                            "Smartwatch",
                            "Acessórios",
                            "Produtos Gamer",
                            "Más Vendidos",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Kids":
                        df_path = "static/home/csv/PORTUGUES_PT/KIDSPORTUGAL.csv"

                        pri = "Moda"
                        seg = "Acessórios"
                        ter = "Brinquedos"
                        qua = "Cuidados"
                        qui = "Maternidader"
                        sex = "Más Vendidos"

                        primeiro = "Moda"
                        segundo = "Acessórios"
                        terceiro = "Brinquedos"
                        quarto = "Cuidados"
                        quinto = "Maternidade"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Moda",
                            "Acessórios",
                            "Brinquedos",
                            "Cuidados",
                            "Maternidade",
                            "Más Vendidos",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Casa":
                        df_path = "static/home/csv/PORTUGUES_PT/CASA_PORTUGAL.csv"
                        
                        pri = "Acessórios"
                        seg = "Iluminação"
                        ter = "Decoração de Ambientes"
                        qua = "Organização e Armazenamento"
                        qui = "Aromaterapia e Bem-estar"
                        sex = "Produtos Sazonais"

                        primeiro = "Acessórios"
                        segundo = "Iluminação"
                        terceiro = "Decoração-de-Ambientes"
                        quarto = "Organização-e-Armazenamento"
                        quinto = "Aromaterapia-e-Bem-estar"
                        sexto = "Produtos-Sazonais"
                        colecoes = [
                            "Acessórios",
                            "Iluminação",
                            "Decoração de Ambientes",
                            "Organização e Armazenamento",
                            "Aromaterapia e Bem-estar",
                            "Produtos Sazonais",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Pet":
                        df_path = "static/home/csv/PORTUGUES_PT/PETPORTUGAL.csv"

                        pri = "Camas"
                        seg = "Acessórios"
                        ter = "Comedouros"
                        qua = "Brinquedos"
                        qui = "Roupas"
                        sex = "Más Vendidos"

                        primeiro = "Camas"
                        segundo = "Acessórios"
                        terceiro = "Comedouros"
                        quarto = "Brinquedos"
                        quinto = "Roupas"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Camas",
                            "Acessórios",
                            "Comedouros",
                            "Brinquedos",
                            "Roupas",
                            "Más Vendidos",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Fitness":
                        df_path = "static/home/csv/PORTUGUES_PT/FITNESSPORTUGAL.csv"

                        pri = "Moda"
                        seg = "Acessórios"
                        ter = "Calçados"
                        qua = "Corretores"
                        qui = "Produtos Esportivos"
                        sex = "Más Vendidos"

                        primeiro = "Moda"
                        segundo = "Acessórios"
                        terceiro = "Calçados"
                        quarto = "Corretores"
                        quinto = "Produtos-Esportivos"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Moda",
                            "Acessórios",
                            "Calçados",
                            "Corretores",
                            "Produtos Esportivos",
                            "Más Vendidos",
                        ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                    
                    elif estilo == "Masculino":
                        df_path = "static/home/csv/PT_BR/30MASCULINO.csv"

                        pri = "Acessórios"
                        seg = "Ferramentas"
                        ter = "Pesca"
                        qua = "Vestuário"
                        qui = "Automotivos"
                        sex = "Más Vendidos"

                        primeiro = "Acessórios"
                        segundo = "Ferramentas"
                        terceiro = "Pesca"
                        quarto = "Vestuário"
                        quinto = "Automotivos"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Acessórios", 
                            "Ferramentas", 
                            "Pesca", 
                            "Vestuário", 
                            "Automotivos", 
                            "Más Vendidos"
                            ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                         
                    elif estilo == "Feminino":
                        df_path = "static/home/csv/PT_BR/30FEMININO.csv"

                        pri = "Moda"
                        seg = "Acessórios"
                        ter = "Make"
                        qua = "Escovas Alisadoras"
                        qui = "Saúde"
                        sex = "Más Vendidos"

                        primeiro = "Moda"
                        segundo = "Acessórios"
                        terceiro = "Make"
                        quarto = "Escovas-Alisadoras"
                        quinto = "Saúde"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Moda", 
                            "Acessórios", 
                            "Make", 
                            "Escovas Alisadoras", 
                            "Saúde", 
                            "Más Vendidos"
                            ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                              
                    elif estilo == "Moda":
                        df_path = "static/home/csv/PT_BR/MODA.csv"
                        pri = "Acessórios"
                        seg = "Feminino"
                        ter = "Masculino"
                        qua = "Infantil"
                        qui = "Jóias"
                        sex = "Promoções"

                        primeiro = "Acessórios"
                        segundo = "Feminino"
                        terceiro = "Masculino"
                        quarto = "Infantil"
                        quinto = "Jóias"
                        sexto = "Promoções"
                        colecoes = [
                            "Acessórios",
                            "Feminino",
                            "Masculino",
                            "Infantil",
                            "Jóias",
                            "Promoções",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Saúde e Beleza":
                        df_path = "static/home/csv/PORTUGUES_PT/SAUDE_BEM_ESTAR_PT.csv"
                        pri = "Fitness e Exercício"
                        seg = "Beleza"
                        ter = "Alimentação"
                        qua = "Relaxamento"
                        qui = "Cuidados Pessoais"
                        sex = "Monitoramento da Saúde"

                        primeiro = "Fitness-e-Exercício"
                        segundo = "Beleza"
                        terceiro = "Alimentação"
                        quarto = "Relaxamento"
                        quinto = "Cuidados-Pessoais"
                        sexto = "Monitoramento-da-Saúde"
                        colecoes = [
                            "Fitness e Exercício",
                            "Beleza",
                            "Alimentação",
                            "Relaxamento",
                            "Cuidados Pessoais",
                            "Monitoramento da Saúde",
                        ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                
                ##########################
                ########################## 
                elif pais == "EUA":
                    
                    Tema = "https://github.com/ronildob/temas/raw/refs/heads/main/inglesoficial.zip"
                    paginas = {
                        "Track Order": "static/home/politicas/Track Order.txt",
                        "E-Commerce Agreement": "static/home/politicas/E-Commerce Agreement.docx",
                        "Terms of Use": "static/home/politicas/Terms of Use.docx",
                        "Deadline": "static/home/politicas/Deadline.docx",
                        "Privacy Policy": "static/home/politicas/Privacy Policy.docx",
                        "Exchanges or Return": "static/home/politicas/Exchanges or Return.docx",
                        "About Us": "static/home/politicas/About Us.docx",
                        "Secure Payment": "static/home/politicas/Secure Payment.docx",
                    }
                    politicas_para_atualizar = {
                        "Legal notice": "static/home/politicas/E-Commerce Agreement.docx",
                        "Terms of service": "static/home/politicas/Terms of Use.docx",
                        "Shipping policy": "static/home/politicas/Deadline.docx",
                        # "Privacy policy": "static/home/politicas/Privacy Policy.docx",
                        "Refund policy": "static/home/politicas/Exchanges or Return.docx",
                    }
                    ##############################
                    
                    if business_hours == "hora1":
                        antend = "Monday to Friday from 9:00 AM to 6:00 PM"
                    elif business_hours == "hora2":
                        antend = "Monday to Friday from 9:00 AM to 10:00 PM"
                    elif business_hours == "hora3":
                        antend = "Monday to Saturday from 9:00 AM to 6:00 PM"
                    elif business_hours == "hora4":
                        antend = "Monday to Saturday from 9:00 AM to 10:00 PM"
                    elif business_hours == "hora5":
                        antend = "Monday to Friday from 9:00 AM to 6:00 PM - Saturday from 9:00 AM to 12:00 PM"
                    elif business_hours == "hora6":
                        antend = "Monday to Friday from 9:00 AM to 10:00 PM - Saturday from 9:00 AM to 12:00 PM"
                    elif business_hours == "hora7":
                        antend = "Every day from 9:00 AM to 6:00 PM"
                    elif business_hours == "hora8":
                        antend = "Every day from 9:00 AM to 10:00 PM"
                    ##############################
                    footer = f"""<p><strong>Email:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a></p><p><strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{
                        cod_pais[pais]} {whatsapp}</a></p><p><strong>Business Hours: </strong>{antend}</p>"""
                    titlezap = f"Contact Us \n<span>\nThe best support is with {empresa}!\n</span>"
                    ##############################
                    if estilo == "Genérica":
                        df_path = "static/home/csv/INGLES/30GEN_EUA.csv"

                        pri = "House"
                        seg = "Electronics"
                        ter = "Pets"
                        qua = "Fitness"
                        qui = "Kids"
                        sex = "Best Sellers"
                    
                        primeiro = "House"
                        segundo = "Electronics"
                        terceiro = "Pets"
                        quarto = "Fitness"
                        quinto = "Kids"
                        sexto = "Best-Sellers"
                        colecoes = [
                            "House",
                            "Electronics",
                            "Pets",
                            "Fitness",
                            "Kids",
                            "Best Sellers",
                        ]
                        ##############################

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Eletrônicos":
                        df_path = "static/home/csv/INGLES/ELETRONICOPT.csv"

                        pri = "Headphones"
                        seg = "Drones"
                        ter = "Smartwatch"
                        qua = "Accessories"
                        qui = "Gamer Products"
                        sex = "Best Sellers"

                        primeiro = "Headphones"
                        segundo = "Drones"
                        terceiro = "Smartwatch"
                        quarto = "Accessories"
                        quinto = "Gamer-Products"
                        sexto = "Best-Sellers"
                        colecoes = [
                            "Headphones",
                            "Drones",
                            "Smartwatch",
                            "Accessories",
                            "Gamer Products",
                            "Best Sellers",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Kids":
                        df_path = "static/home/csv/INGLES/KIDSPORTUGAL.csv"

                        pri = "Fashion"
                        seg = "Accessories"
                        ter = "Toys"
                        qua = "Care"
                        qui = "Maternity"
                        sex = "Best Sellers"

                        primeiro = "Fashion"
                        segundo = "Accessories"
                        terceiro = "Toys"
                        quarto = "Care"
                        quinto = "Maternity"
                        sexto = "Best-Sellers"
                        colecoes = [
                            "Fashion",
                            "Accessories",
                            "Toys",
                            "Care",
                            "Maternity",
                            "Best Sellers",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Casa":
                        df_path = "static/home/csv/INGLES/CASA_PORTUGAL.csv"
                        
                        pri = "Accessories"
                        seg = "Lighting"
                        ter = "Environment Decoration"
                        qua = "Organization and Storage"
                        qui = "Aromatherapy and Wellbeing"
                        sex = "Seasonal Products"

                        primeiro = "Accessories"
                        segundo = "Lighting"
                        terceiro = "Environment-Decoration"
                        quarto = "Organization-and-Storage"
                        quinto = "Aromatherapy-and-Wellbeing"
                        sexto = "Seasonal-Products"
                        colecoes = [
                            "Accessories",
                            "Lighting",
                            "Environment Decoration",
                            "Organization and Storage",
                            "Aromatherapy and Wellbeing",
                            "Seasonal Products",
                        ]
                        
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Pet":
                        df_path = "static/home/csv/INGLES/PETPORTUGAL.csv"

                        pri = "Beds"
                        seg = "Accessories"
                        ter = "Feeders"
                        qua = "Toys"
                        qui = "Clothes"
                        sex = "Best Sellers"

                        primeiro = "Beds"
                        segundo = "Accessories"
                        terceiro = "Feeders"
                        quarto = "Toys"
                        quinto = "Clothes"
                        sexto = "Best-Sellers"
                        colecoes = [
                            "Beds",
                            "Accessories",
                            "Feeders",
                            "Toys",
                            "Clothes",
                            "Best Sellers",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Fitness":
                        df_path = "static/home/csv/INGLES/FITNESSPORTUGAL.csv"

                        pri = "Fashion"
                        seg = "Accessories"
                        ter = "Shoes"
                        qua = "Brokers"
                        qui = "Sports Products"
                        sex = "Best Sellers"

                        primeiro = "Fashion"
                        segundo = "Accessories"
                        terceiro = "Shoes"
                        quarto = "Brokers"
                        quinto = "Sports-Products"
                        sexto = "Best-Sellers"
                        colecoes = [
                            "Fashion",
                            "Accessories",
                            "Shoes",
                            "Brokers",
                            "Sports Products",
                            "Best Sellers",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Masculino":
                        df_path = "static/home/csv/PT_BR/30MASCULINO.csv"

                        pri = "Accessories"
                        seg = "Tools"
                        ter = "Fishing"
                        qua = "Clothing"
                        qui = "Automotive"
                        sex = "Best Sellers"

                        primeiro = "Accessories"
                        segundo = "Tools"
                        terceiro = "Fishing"
                        quarto = "Clothing"
                        quinto = "Automotive"
                        sexto = "Best-Sellers"
                        colecoes = [
                            "Accessories", 
                            "Tools", 
                            "Fishing", 
                            "Clothing", 
                            "Automotive", 
                            "Best Sellers"
                            ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                         
                    elif estilo == "Feminino":
                        df_path = "static/home/csv/PT_BR/30FEMININO.csv"

                        pri = "Fashion"
                        seg = "Accessories"
                        ter = "Makeup"
                        qua = "Straightening Brushes"
                        qui = "Health"
                        sex = "Best Sellers"

                        primeiro = "Fashion"
                        segundo = "Accessories"
                        terceiro = "Makeup"
                        quarto = "Straightening-Brushes"
                        quinto = "Health"
                        sexto = "Best-Sellers"
                        colecoes = [
                            "Fashion", 
                            "Accessories", 
                            "Makeup", 
                            "Straightening Brushes", 
                            "Health", 
                            "Best Sellers"
                            ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                              
                    elif estilo == "Moda":
                        df_path = "static/home/csv/INGLES/MODAPORTUGAL.csv"
                        pri = "Accessories"
                        seg = "Feminine"
                        ter = "Masculine"
                        qua = "Children"
                        qui = "Jewelry"
                        sex = "Promotions"

                        primeiro = "Accessories"
                        segundo = "Feminine"
                        terceiro = "Masculine"
                        quarto = "Children"
                        quinto = "Jewelry"
                        sexto = "Promotions"
                        colecoes = [
                            "Accessories",
                            "Feminine",
                            "Masculine",
                            "Children",
                            "Jewelry",
                            "Promotions",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                    
                    elif estilo == "Saúde e Beleza":
                        df_path = "static/home/csv/INGLES/SAUDEBLEZPORT.csv"
                        pri = "Fitness and Exercise"
                        seg = "Beauty"
                        ter = "Best Sellers"
                        qua = "Relaxation"
                        qui = "Personal Care"
                        sex = "Health Monitoring"

                        primeiro = "Fitness-and-Exercise"
                        segundo = "Beauty"
                        terceiro = "Best-Sellers"
                        quarto = "Relaxation"
                        quinto = "Personal-Care"
                        sexto = "Health-Monitoring"
                        colecoes = [
                            "Fitness and Exercise",
                            "Beauty",
                            "Best Sellers",
                            "Relaxation",
                            "Personal Care",
                            "Health Monitoring",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                ##########################
                ########################## 
                elif pais == "Espanha":
                    Tema = "https://github.com/ronildob/TEMAS2025/raw/refs/heads/main/tema_latan%2001-2025.zip"
                    #Tema = "https://github.com/ronildob/temas/raw/refs/heads/main/espanholoficial.zip"
                    paginas = {
                        "Rastrear Pedidos": "static/home/politicas/seguimiento-del-pedido.txt",
                        "Contrato de Comercio Electrónico": "static/home/politicas/Contrato de Comercio Electrónico.docx",
                        "Términos de Uso": "static/home/politicas/Términos de Uso.docx",
                        "Plazo de Entrega": "static/home/politicas/Plazo de Entrega.docx",
                        "Política de Privacidad": "static/home/politicas/Política de Privacidad.docx",
                        "Cambios o Devoluciones": "static/home/politicas/Cambios o Devoluciones.docx",
                        "Sobre Nosotros": "static/home/politicas/Sobre Nosotros.docx",
                        "Pago Seguro": "static/home/politicas/Pago Seguro.docx",
                    }
                    politicas_para_atualizar = {
                        "Legal notice": "static/home/politicas/Contrato de Comercio Electrónico.docx",
                        "Terms of service": "static/home/politicas/Términos de Uso.docx",
                        "Shipping policy": "static/home/politicas/Plazo de Entrega.docx",
                        # "Privacy policy": "static/home/politicas/Política de Privacidad.docx",
                        "Refund policy": "static/home/politicas/Cambios o Devoluciones.docx",
                    }
                    ##############################
                    
                    if business_hours == "hora1":
                        antend = "Lunes a viernes de 9:00 a 18:00"
                    elif business_hours == "hora2":
                        antend = "Lunes a viernes de 9:00 a 22:00"
                    elif business_hours == "hora3":
                        antend = "Lunes a Sábado de  9:00 a 18"
                    elif business_hours == "hora4":
                        antend = "Lunes a Sábado de  9:00 a 22"
                    elif business_hours == "hora5":
                        antend = "Lunes a Viernes 9:00 a 18:00, Sábado 9:00 a 12:00"
                    elif business_hours == "hora6":
                        antend = (
                            "Lunes a viernes de 9:00 a 22:00 - Sábado de 9:00 a 12:00"
                        )
                    elif business_hours == "hora7":
                        antend = "Todos los días de 9:00 a 18:00"
                    elif business_hours == "hora8":
                        antend = "Todos los días de 9:00 a 22:00"
                    ##############################
                    footer = f"""<p><strong>Correo electrónico:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a></p><p><strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{
                        cod_pais[pais]} {whatsapp}</a></p><p><strong>Horario de atención: </strong>{antend}</p>"""
                    titlezap = f"Contáctanos \n<span>\n¡El mejor soporte es con {empresa}!\n</span>"
                    ##############################
                    if estilo == "Genérica":
                        df_path = "static/home/csv/ESPANHOL/30GENER_ESPANHOL.csv"

                        pri = "Casa"
                        seg = "Electrónica"
                        ter = "Pets"
                        qua = "Fitness"
                        qui = "Kids"
                        sex = "Más Vendidos"

                        primeiro = "Casa"
                        segundo = "Electrónica"
                        terceiro = "Pets"
                        quarto = "Fitness"
                        quinto = "Kids"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Casa",
                            "Electrónica",
                            "Pets",
                            "Fitness",
                            "Kids",
                            "Más Vendidos",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                        
                    elif estilo == "Eletrônicos":
                        df_path = "static/home/csv/ESPANHOL/30ELETRONICOS.csv"
                        pri = "Auriculares"
                        seg = "Drones"
                        ter = "Smartwatch"
                        qua = "Accesorios"
                        qui = "Productos para Jugadores"
                        sex = "Más Vendidos"

                        primeiro = "Auriculares"
                        segundo = "Drones"
                        terceiro = "Smartwatch"
                        quarto = "Accesorios"
                        quinto = "Productos-para-Jugadores"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Auriculares", 
                            "Drones", 
                            "Smartwatch", 
                            "Accesorios", 
                            "Productos para Jugadores", 
                            "Más Vendidos"
                            ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                        
                    elif estilo == "Kids":
                        df_path = "static/home/csv/ESPANHOL/30KIDS.csv"
                        pri = "Moda"
                        seg = "Accesorios"
                        ter = "Juguetes"
                        qua = "Cuidado"
                        qui = "Maternidad"
                        sex = "Más Vendidos"

                        primeiro = "Moda"
                        segundo = "Accesorios"
                        terceiro = "Juguetes"
                        quarto = "Cuidado"
                        quinto = "Maternidad"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Moda", 
                            "Accesorios", 
                            "Juguetes", 
                            "Cuidado", 
                            "Maternidad", 
                            "Más Vendidos"]
                        
                       # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Casa":
                        df_path = "static/home/csv/ESPANHOL/"
                        
                        pri = "Accesorios"
                        seg = "Iluminación"
                        ter = "Decoración del ambiente"
                        qua = "Organización y almacenamiento"
                        qui = "Aromaterapia y Bienestar"
                        sex = "Productos de Temporada"

                        primeiro = "Accesorios"
                        segundo = "Iluminación"
                        terceiro = "Decoración-del-ambiente"
                        quarto = "Organización-y-almacenamiento"
                        quinto = "Aromaterapia-y-Bienestar"
                        sexto = "Productos-de-Temporada"
                        colecoes = [
                            "Accesorios",
                            "Iluminación",
                            "Decoración del ambiente",
                            "Organización y almacenamiento",
                            "Aromaterapia y Bienestar",
                            "Productos de Temporada",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Pet":
                        df_path = "static/home/csv/ESPANHOL/PETPORTUGAL.csv"

                        pri = "Camas"
                        seg = "Accesorios"
                        ter = "Comederos"
                        qua = "Juguetes"
                        qui = "Ropa"
                        sex = "Más Vendidos"

                        primeiro = "Camas"
                        segundo = "Accesorios"
                        terceiro = "Comederos"
                        quarto = "Juguetes"
                        quinto = "Ropa"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Camas",
                            "Accesorios",
                            "Comederos",
                            "Juguetes",
                            "Ropa",
                            "Más Vendidos",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Fitness":
                        df_path = "static/home/csv/ESPANHOL/FITNESSPORTUGAL.csv"

                        pri = "Moda"
                        seg = "Accesorios"
                        ter = "Zapatos"
                        qua = "Correctores"
                        qui = "Productos Deportivos"
                        sex = "Más Vendidos"

                        primeiro = "Moda"
                        segundo = "Accesorios"
                        terceiro = "Zapatos"
                        quarto = "Correctores"
                        quinto = "Productos-Deportivos"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Moda",
                            "Accesorios",
                            "Zapatos",
                            "Correctores",
                            "Productos Deportivos",
                            "Más Vendidos",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Masculino":
                        df_path = "static/home/csv/PT_BR/30MASCULINO.csv"

                        pri = "Accesorios"
                        seg = "Herramientas"
                        ter = "Pesca"
                        qua = "Ropa"
                        qui = "Automotor"
                        sex = "Más Vendidos"

                        primeiro = "Accesorios"
                        segundo = "Herramientas"
                        terceiro = "Pesca"
                        quarto = "Ropa"
                        quinto = "Automotor"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Accesorios", 
                            "Herramientas", 
                            "Pesca", 
                            "Ropa", 
                            "Automotor", 
                            "Más Vendidos"
                            ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                         
                    elif estilo == "Feminino":
                        df_path = "static/home/csv/PT_BR/30FEMININO.csv"

                        pri = "Moda"
                        seg = "Accesorios"
                        ter = "Make"
                        qua = "Cepillos Alisadores"
                        qui = "Salud"
                        sex = "Más Vendidos"

                        primeiro = "Moda"
                        segundo = "Accesorios"
                        terceiro = "Make"
                        quarto = "Cepillos-Alisadores"
                        quinto = "Salud"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Moda", 
                            "Accesorios", 
                            "Make", 
                            "Cepillos Alisadores", 
                            "Salud", 
                            "Más Vendidos"
                            ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                        
                    elif estilo == "Moda":
                        df_path = "static/home/csv/ESPANHOL/"
                        pri = "Accesorios"
                        seg = "Femenino"
                        ter = "Masculino"
                        qua = "Niños"
                        qui = "Joyas"
                        sex = "Promociones"

                        primeiro = "Accesorios"
                        segundo = "Femenino"
                        terceiro = "Masculino"
                        quarto = "Niños"
                        quinto = "Joyas"
                        sexto = "Promociones"
                        colecoes = [
                            "Accesorios",
                            "Femenino",
                            "Masculino",
                            "Niños",
                            "Joyas",
                            "Promociones",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Saúde e Beleza":
                        df_path = "static/home/csv/ESPANHOL/"
                        
                        
                        pri = "Fitness y Ejercicio"
                        seg = "Belleza"
                        ter = "Más Vendidos"
                        qua = "Relajación"
                        qui = "Cuidado Personal"
                        sex = "Monitoreo de Salud"

                        primeiro = "Fitness-y-Ejercicio"
                        segundo = "Belleza"
                        terceiro = "Más-Vendidos"
                        quarto = "Relajación"
                        quinto = "Cuidado-Personal"
                        sexto = "Monitoreo-de-Salud"
                        colecoes = [
                            "Fitness y Ejercicio",
                            "Belleza",
                            "Más Vendidos",
                            "Relajación",
                            "Cuidado Personal",
                            "Monitoreo de Salud",
                        ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                ##########################
                ########################## 
                else:  
                    Tema = "https://github.com/ronildob/temas/raw/refs/heads/main/espanlatanoficial.zip"
                    paginas = {
                        "Rastrear Pedidos": "static/home/politicas/seguimiento-del-pedido.txt",
                        "Contrato de Comercio Electrónico": "static/home/politicas/Contrato de Comercio Electrónico.docx",
                        "Términos de Uso": "static/home/politicas/Términos de Uso.docx",
                        "Plazo de Entrega": "static/home/politicas/Plazo de Entrega.docx",
                        "Política de Privacidad": "static/home/politicas/Política de Privacidad.docx",
                        "Cambios o Devoluciones": "static/home/politicas/Cambios o Devoluciones.docx",
                        "Sobre Nosotros": "static/home/politicas/Sobre Nosotros.docx",
                        "Pago Seguro": "static/home/politicas/Pago Seguro.docx",
                    }
                    politicas_para_atualizar = {
                        "Legal notice": "static/home/politicas/Contrato de Comercio Electrónico.docx",
                        "Terms of service": "static/home/politicas/Términos de Uso.docx",
                        "Shipping policy": "static/home/politicas/Plazo de Entrega.docx",
                        # "Privacy policy": "static/home/politicas/Política de Privacidad.docx",
                        "Refund policy": "static/home/politicas/Cambios o Devoluciones.docx",
                    }
                    ##############################
                
                    if business_hours == "hora1":
                        antend = "Lunes a viernes de 9:00 a 18:00"
                    elif business_hours == "hora2":
                        antend = "Lunes a viernes de 9:00 a 22:00"
                    elif business_hours == "hora3":
                        antend = "Lunes a Sábado de  9:00 a 18"
                    elif business_hours == "hora4":
                        antend = "Lunes a Sábado de  9:00 a 22"
                    elif business_hours == "hora5":
                        antend = "Lunes a Viernes 9:00 a 18:00, Sábado 9:00 a 12:00"
                    elif business_hours == "hora6":
                        antend = (
                            "Lunes a viernes de 9:00 a 22:00 - Sábado de 9:00 a 12:00"
                        )
                    elif business_hours == "hora7":
                        antend = "Todos los días de 9:00 a 18:00"
                    elif business_hours == "hora8":
                        antend = "Todos los días de 9:00 a 22:00"
                    ##############################
                    footer = f"""<p><strong>Correo electrónico:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a></p><p><strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{
                        cod_pais[pais]} {whatsapp}</a></p><p><strong>Horario de atención: </strong>{antend}</p>"""
                    titlezap = f"Contáctanos \n<span>\n¡El mejor soporte es con {empresa}!\n</span>"
                    ##############################
                    ##############################
                    if estilo == "Genérica":
                        df_path = "static/home/csv/ESPANHOL/30GENER_ESPANHOL.csv"

                        pri = "Casa"
                        seg = "Electrónica"
                        ter = "Pets"
                        qua = "Fitness"
                        qui = "Kids"
                        sex = "Más Vendidos"

                        primeiro = "Casa"
                        segundo = "Electrónica"
                        terceiro = "Pets"
                        quarto = "Fitness"
                        quinto = "Kids"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Casa",
                            "Electrónica",
                            "Pets",
                            "Fitness",
                            "Kids",
                            "Más Vendidos",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                        
                    elif estilo == "Eletrônicos":
                        df_path = "static/home/csv/ESPANHOL/30ELETRONICOS.csv"
                        pri = "Auriculares"
                        seg = "Drones"
                        ter = "Smartwatch"
                        qua = "Accesorios"
                        qui = "Productos para Jugadores"
                        sex = "Más Vendidos"

                        primeiro = "Auriculares"
                        segundo = "Drones"
                        terceiro = "Smartwatch"
                        quarto = "Accesorios"
                        quinto = "Productos-para-Jugadores"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Auriculares", 
                            "Drones", 
                            "Smartwatch", 
                            "Accesorios", 
                            "Productos para Jugadores", 
                            "Más Vendidos"
                            ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                        
                    elif estilo == "Kids":
                        df_path = "static/home/csv/ESPANHOL/30KIDS.csv"
                        pri = "Moda"
                        seg = "Accesorios"
                        ter = "Juguetes"
                        qua = "Cuidado"
                        qui = "Maternidad"
                        sex = "Más Vendidos"

                        primeiro = "Moda"
                        segundo = "Accesorios"
                        terceiro = "Juguetes"
                        quarto = "Cuidado"
                        quinto = "Maternidad"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Moda", 
                            "Accesorios", 
                            "Juguetes", 
                            "Cuidado", 
                            "Maternidad", 
                            "Más Vendidos"]
                        
                       # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Casa":
                        df_path = "static/home/csv/ESPANHOL/"
                        
                        pri = "Accesorios"
                        seg = "Iluminación"
                        ter = "Decoración del ambiente"
                        qua = "Organización y almacenamiento"
                        qui = "Aromaterapia y Bienestar"
                        sex = "Productos de Temporada"

                        primeiro = "Accesorios"
                        segundo = "Iluminación"
                        terceiro = "Decoración-del-ambiente"
                        quarto = "Organización-y-almacenamiento"
                        quinto = "Aromaterapia-y-Bienestar"
                        sexto = "Productos-de-Temporada"
                        colecoes = [
                            "Accesorios",
                            "Iluminación",
                            "Decoración del ambiente",
                            "Organización y almacenamiento",
                            "Aromaterapia y Bienestar",
                            "Productos de Temporada",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Pet":
                        df_path = "static/home/csv/ESPANHOL/PETPORTUGAL.csv"

                        pri = "Camas"
                        seg = "Accesorios"
                        ter = "Comederos"
                        qua = "Juguetes"
                        qui = "Ropa"
                        sex = "Más Vendidos"

                        primeiro = "Camas"
                        segundo = "Accesorios"
                        terceiro = "Comederos"
                        quarto = "Juguetes"
                        quinto = "Ropa"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Camas",
                            "Accesorios",
                            "Comederos",
                            "Juguetes",
                            "Ropa",
                            "Más Vendidos",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Fitness":
                        df_path = "static/home/csv/ESPANHOL/FITNESSPORTUGAL.csv"

                        pri = "Moda"
                        seg = "Accesorios"
                        ter = "Zapatos"
                        qua = "Correctores"
                        qui = "Productos Deportivos"
                        sex = "Más Vendidos"

                        primeiro = "Moda"
                        segundo = "Accesorios"
                        terceiro = "Zapatos"
                        quarto = "Correctores"
                        quinto = "Productos-Deportivos"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Moda",
                            "Accesorios",
                            "Zapatos",
                            "Correctores",
                            "Productos Deportivos",
                            "Más Vendidos",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Masculino":
                        df_path = "static/home/csv/PT_BR/30MASCULINO.csv"

                        pri = "Accesorios"
                        seg = "Herramientas"
                        ter = "Pesca"
                        qua = "Ropa"
                        qui = "Automotor"
                        sex = "Más Vendidos"

                        primeiro = "Accesorios"
                        segundo = "Herramientas"
                        terceiro = "Pesca"
                        quarto = "Ropa"
                        quinto = "Automotor"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Accesorios", 
                            "Herramientas", 
                            "Pesca", 
                            "Ropa", 
                            "Automotor", 
                            "Más Vendidos"
                            ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                         
                    elif estilo == "Feminino":
                        df_path = "static/home/csv/PT_BR/30FEMININO.csv"

                        pri = "Moda"
                        seg = "Accesorios"
                        ter = "Make"
                        qua = "Cepillos Alisadores"
                        qui = "Salud"
                        sex = "Más Vendidos"

                        primeiro = "Moda"
                        segundo = "Accesorios"
                        terceiro = "Make"
                        quarto = "Cepillos-Alisadores"
                        quinto = "Salud"
                        sexto = "Más-Vendidos"
                        colecoes = [
                            "Moda", 
                            "Accesorios", 
                            "Make", 
                            "Cepillos Alisadores", 
                            "Salud", 
                            "Más Vendidos"
                            ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }
                        
                    elif estilo == "Moda":
                        df_path = "static/home/csv/ESPANHOL/"
                        pri = "Accesorios"
                        seg = "Femenino"
                        ter = "Masculino"
                        qua = "Niños"
                        qui = "Joyas"
                        sex = "Promociones"

                        primeiro = "Accesorios"
                        segundo = "Femenino"
                        terceiro = "Masculino"
                        quarto = "Niños"
                        quinto = "Joyas"
                        sexto = "Promociones"
                        colecoes = [
                            "Accesorios",
                            "Femenino",
                            "Masculino",
                            "Niños",
                            "Joyas",
                            "Promociones",
                        ]
                        
                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                    elif estilo == "Saúde e Beleza":
                        df_path = "static/home/csv/ESPANHOL/"
                        
                        
                        pri = "Fitness y Ejercicio"
                        seg = "Belleza"
                        ter = "Más Vendidos"
                        qua = "Relajación"
                        qui = "Cuidado Personal"
                        sex = "Monitoreo de Salud"

                        primeiro = "Fitness-y-Ejercicio"
                        segundo = "Belleza"
                        terceiro = "Más-Vendidos"
                        quarto = "Relajación"
                        quinto = "Cuidado-Personal"
                        sexto = "Monitoreo-de-Salud"
                        colecoes = [
                            "Fitness y Ejercicio",
                            "Belleza",
                            "Más Vendidos",
                            "Relajación",
                            "Cuidado Personal",
                            "Monitoreo de Salud",
                        ]

                        # urls = {
                            # "desktop1": "",
                            # "desktop2": "",
                            # "desktop3": "",
                            # "mobile1": "",
                            # "mobile2": "",
                            # "mobile3": "",
                            # "capa1": "",
                            # "capa2": "",
                            # "capa3": "",
                            # "capa4": "",
                            # "capa5": "",
                            # "capa6": "",
                        # }

                ##########################
                
                ##########################
                ##########################

                shopify.Session.setup(api_key=API_KEY, secret=API_SECRET)
                session = shopify.Session(SHOP_URL, API_VERSION, PRIVATE_APP_PASSWORD)
                shopify.ShopifyResource.activate_session(session)

                df = pd.read_csv(df_path)

                def importar_produtos(df, pais):
                    try:
                        erros = []  # Lista para acumular erros de cada produto

                        # Dicionário com as taxas de câmbio para cada país em
                        # relação ao Brasil
                        taxa_cambio = {
                            "Brasil": 1.0,
                            "Argentina": 0.04,
                            "Bolívia": 0.14,
                            "Chile": 0.14,
                            "Colômbia": 0.72,
                            "Equador": 0.27,
                            "Paraguai": 0.14,
                            "Peru": 0.24,
                            "Uruguai": 0.25,
                            "Venezuela": 0.11,
                            "EUA": 5.0,
                            "Portugal": 5.5,
                            "Espanha": 5.5,
                        }

                        # Dicionário de formatação de moeda para cada país
                        formato_moeda = {
                            "Brasil": "{:.2f}",
                            "Argentina": "{:.2f}",
                            "Bolívia": "{:.2f}",
                            "Chile": "{:.2f}",
                            "Colômbia": "{:.2f}",
                            "Equador": "{:.2f}",
                            "Paraguai": "{:.2f}",
                            "Peru": "S/. {:.2f}",
                            "Uruguai": "{:.2f}",
                            "Venezuela": "{:.2f}",
                            "EUA": "{:.2f}",
                            "Portugal": "{:.2f}",
                            "Espanha": "{:.2f}",
                        }

                        for index, row in df.iterrows():
                            try:
                                handle = row["Handle"]
                                title = (
                                    row["Title"]
                                    if pd.notna(row["Title"]) and row["Title"].strip()
                                    else "Título padrão"
                                )
                                image_src = row["Image Src"]

                                # Obter a taxa de câmbio com base no país
                                taxa = taxa_cambio.get(pais, 1.0)

                                # Ajuste de preços dependendo do país
                                if pais in ["EUA", "Portugal", "Espanha"]:
                                    variant_price = (
                                        format(row["Variant Price"] / taxa, ".2f")
                                        if pd.notna(row["Variant Price"])
                                        else "0.00"
                                    )
                                    variant_compare_at_price = (
                                        format(
                                            row["Variant Compare At Price"] / taxa,
                                            ".2f",
                                        )
                                        if pd.notna(row["Variant Compare At Price"])
                                        else "0.00"
                                    )
                                else:
                                    variant_price = (
                                        format(row["Variant Price"] * taxa, ".2f")
                                        if pd.notna(row["Variant Price"])
                                        else "0.00"
                                    )
                                    variant_compare_at_price = (
                                        format(
                                            row["Variant Compare At Price"] * taxa,
                                            ".2f",
                                        )
                                        if pd.notna(row["Variant Compare At Price"])
                                        else "0.00"
                                    )

                                # Aplicar formatação de moeda
                                variant_price = formato_moeda.get(
                                    pais, "${:.2f}"
                                ).format(float(variant_price))
                                variant_compare_at_price = formato_moeda.get(
                                    pais, "${:.2f}"
                                ).format(float(variant_compare_at_price))

                                # Criação do produto
                                product = shopify.Product(
                                    {
                                        "title": title,
                                        "body_html": (
                                            row["Body (HTML)"]
                                            if pd.notna(row["Body (HTML)"])
                                            else ""
                                        ),
                                        "vendor": (
                                            row["Vendor"]
                                            if pd.notna(row["Vendor"])
                                            else ""
                                        ),
                                        "product_type": (
                                            row["Type"] if pd.notna(row["Type"]) else ""
                                        ),
                                        "tags": (
                                            unidecode.unidecode(row["Tags"])
                                            .lower()
                                            .replace(" ", "-")
                                            if pd.notna(row["Tags"])
                                            else ""
                                        ),
                                        "published": (
                                            row["Published"]
                                            if pd.notna(row["Published"])
                                            else True
                                        ),
                                    }
                                )

                                if not product.save():
                                    erros.append(
                                        f"Erro ao salvar produto '{title}': {product.errors.full_messages()}")
                                    continue

                                # Adicionar variante ao produto
                                variant = shopify.Variant(
                                    {
                                        "taxable": "false",
                                        "price": variant_price,
                                        "compare_at_price": variant_compare_at_price,
                                        "inventory_quantity": 1000,
                                        "inventory_management": "shopify",
                                        "inventory_policy": "continue",
                                    }
                                )
                                product.variants = [variant]

                                if not product.save():
                                    erros.append(
                                        f"Erro ao salvar variante para o produto '{title}': {product.errors.full_messages()}")
                                    continue

                                # Adicionar imagem ao produto
                                if pd.notna(image_src):
                                    image = shopify.Image({"src": image_src})
                                    image.product_id = product.id
                                    if not image.save():
                                        erros.append(
                                            f"Erro ao salvar imagem para o produto '{title}': {image.errors.full_messages()}")

                                print(f"Produto '{title}' carregado com sucesso!")

                            except Exception as e:
                                erros.append(
                                    f"Erro ao importar produto '{title}': {str(e)}")

                        # Retorno consolidado com a lista de erros ou sucesso
                        if erros:
                            return "Erro: " + " | ".join(erros)
                        else:
                            return (
                                "Sucesso: Todos os produtos foram importados com êxito."
                            )

                    except Exception as e:
                        return f"Erro: Exceção geral ao importar produtos - {str(e)}"

                def importar_tema(Tema):
                    try:
                        # Configuração do novo tema na Shopify
                        new_theme = shopify.Theme()
                        new_theme.name = "ShopBotfy"
                        new_theme.src = Tema
                        new_theme.role = "main"

                        # Tentativa de salvar o tema e verificação de erros
                        if new_theme.save():
                            return f"Sucesso: Tema '{new_theme.name}' criado com êxito. ID: {new_theme.id}"
                        else:
                            erros = new_theme.errors.full_messages()
                            return f"Erro: Falha ao criar o tema. Detalhes: {' | '.join(erros)}"

                    except Exception as e:
                        return f"Erro: Exceção ao importar tema - {str(e)}"

        
                # imagens
                def download_svg(url):
                    response = requests.get(url)
                    if response.status_code == 200:
                        local_path = url.split("/")[-1]
                        with open(local_path, "wb") as file:
                            file.write(response.content)
                        return local_path
                    else:
                        raise Exception(f"Erro ao baixar o SVG: {response.status_code}")

                def edit_svg_colors(svg_path, color_mapping):
                    try:
                        tree = ET.parse(svg_path)
                        root = tree.getroot()
                        namespace = {"svg": "http://www.w3.org/2000/svg"}

                        for element in root.findall(".//svg:*", namespaces=namespace):
                            fill = element.attrib.get("fill")
                            if fill and fill in color_mapping:
                                element.set("fill", color_mapping[fill])

                        edited_path = "edited_" + svg_path
                        tree.write(edited_path)
                        return edited_path
                    except Exception as e:
                        raise Exception(f"Erro ao editar o SVG: {e}")

                def validate_svg(svg_path):
                    try:
                        tree = ET.parse(svg_path)
                        root = tree.getroot()
                        return root.tag == "{http://www.w3.org/2000/svg}svg"
                    except Exception:
                        return False

                def convert_svg_to_png(svg_path):
                    try:
                        response = cloudinary.uploader.upload(svg_path, resource_type="image", format="png", folder="banners_e_capas")
                        return response["secure_url"], response["public_id"]
                    except Exception as e:
                        raise Exception(f"Erro ao converter o SVG para PNG: {e}")

                def delete_image_from_cloudinary(public_id):
                    try:
                        cloudinary.uploader.destroy(public_id)
                        print(f"Imagem {public_id} deletada com sucesso no Cloudinary.")
                    except Exception as e:
                        print(f"Erro ao deletar imagem no Cloudinary: {e}")

                def upload_png_to_shopify(png_url, file_name):
                    mutation = """
                        mutation fileCreate($files: [FileCreateInput!]!) {
                            fileCreate(files: $files) {
                                files {
                                    id
                                    alt
                                }
                                userErrors {
                                    message
                                }
                            }
                        }
                    """
                    file_data = {
                        "files": [
                            {"alt": file_name, "originalSource": png_url}
                        ]
                    }

                    response = requests.post(
                        f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json",
                        json={"query": mutation, "variables": file_data},
                        headers={"X-Shopify-Access-Token": PRIVATE_APP_PASSWORD},
                    )
                    if response.status_code != 200:
                        raise Exception(f"Erro ao enviar para Shopify: {response.text}")

                if "public_ids_to_delete" not in request.session:
                    request.session["public_ids_to_delete"] = []

                urls_shopify = {}

                for name, url in urls.items():
                    try:
                        svg_path = download_svg(url)
                        edited_svg_path = edit_svg_colors(svg_path, {"#000000": cor1})

                        if validate_svg(edited_svg_path):
                            png_url, public_id = convert_svg_to_png(edited_svg_path)
                            
                            upload_png_to_shopify(png_url, name)

                            urls_shopify[name] = png_url
                            request.session["public_ids_to_delete"].append(public_id)
                            
                        else:
                            raise Exception(f"SVG inválido: {edited_svg_path}")
                    except Exception as e:
                        print(f"Erro no processamento de {name}: {e}")

                def criar_colecoes(colecoes):
                    try:
                        erros = []  # Lista para acumular mensagens de erro

                        # Iteração sobre cada coleção a ser criada
                        for nome in colecoes:
                            try:
                                nova_colecao = shopify.CustomCollection()
                                nova_colecao.title = nome

                                # Tentativa de salvar a nova coleção
                                if not nova_colecao.save():
                                    erros.append(
                                        f"Erro ao salvar a coleção '{nome}': {nova_colecao.errors.full_messages()}"
                                    )
                                else:
                                    print(f"Coleção '{nome}' criada com sucesso!")

                            except Exception as e:
                                erros.append(
                                    f"Erro ao criar a coleção '{nome}': {str(e)}"
                                )

                        # Retorna uma mensagem final com base no resultado
                        if erros:
                            return "Erro: " + " | ".join(erros)
                        else:
                            return "Sucesso: Todas as coleções foram criadas com êxito."

                    except Exception as e:
                        return f"Erro: Exceção geral ao criar coleções - {str(e)}"

                def criar_paginas(paginas):
                    try:
                        erros = []  # Lista para acumular erros de cada página

                        # Iteração sobre as páginas a serem criadas
                        for titulo, caminho_arquivo in paginas.items():
                            try:
                                # Leitura do conteúdo com base no tipo de
                                # arquivo
                                if caminho_arquivo.endswith(".docx"):
                                    doc = Document(caminho_arquivo)
                                    conteudo = "\n".join(
                                        [p.text for p in doc.paragraphs]
                                    )
                                elif caminho_arquivo.endswith(".txt"):
                                    with open(caminho_arquivo, "r") as f:
                                        conteudo = f.read()
                                else:
                                    erros.append(
                                        f"Formato de arquivo não suportado: {caminho_arquivo}"
                                    )
                                    continue

                                # Criação da nova página na Shopify
                                nova_pagina = shopify.Page()
                                nova_pagina.title = titulo
                                nova_pagina.body_html = conteudo

                                # Salvando a página e verificando sucesso
                                if not nova_pagina.save():
                                    erros.append(
                                        f"Erro ao salvar a página '{titulo}': {nova_pagina.errors.full_messages()}"
                                    )

                            except Exception as e:
                                erros.append(
                                    f"Erro ao criar a página '{titulo}': {str(e)}"
                                )

                        # Verificação final para consolidar a resposta
                        if erros:
                            return "Erro: " + " | ".join(erros)
                        else:
                            return "Sucesso: Todas as páginas foram criadas com êxito."

                    except Exception as e:
                        return f"Erro: Exceção geral ao criar páginas - {str(e)}"

                def ler_documento(caminho):
                    document = Document(caminho)
                    texto = ""
                    for paragraph in document.paragraphs:
                        texto += paragraph.text + "\n"
                    return texto

                def atualizar_politicas():
                    try:
                        # Endpoint para atualizar as políticas
                        endpoint = (
                            f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json"
                        )

                        # Cabeçalhos da requisição com autenticação
                        headers = {
                            "X-Shopify-Access-Token": PRIVATE_APP_PASSWORD,
                            "Content-Type": "application/json",
                        }

                        # Mutation para atualizar a política
                        mutation = """
                                mutation shopPolicyUpdate($shopPolicy: ShopPolicyInput!) {
                                    shopPolicyUpdate(shopPolicy: $shopPolicy) {
                                        shopPolicy {
                                            id
                                        }
                                        userErrors {
                                            field
                                            message
                                        }
                                    }
                                }
                                """

                        # Lista para coletar erros (se houver) durante o
                        # processo
                        erros = []

                        # Iteração para cada política e tentativa de
                        # atualização
                        for (
                            titulo,
                            caminho_documento,
                        ) in politicas_para_atualizar.items():
                            corpo = ler_documento(caminho_documento)
                            politica_para_atualizar = {
                                "body": corpo,
                                "type": titulo.upper().replace(" ", "_"),
                            }

                            variables = {"shopPolicy": politica_para_atualizar}

                            # Envia a requisição POST com a mutação
                            response = requests.post(
                                endpoint,
                                json={"query": mutation, "variables": variables},
                                headers=headers,
                            )

                            # Verifica a resposta da requisição
                            if response.status_code == 200:
                                data = response.json()
                                if "errors" in data:
                                    for error in data["errors"]:
                                        erros.append(
                                            f"Erro ao atualizar a política '{titulo}': {error['message']}"
                                        )
                                elif (
                                    data.get("data", {})
                                    .get("shopPolicyUpdate", {})
                                    .get("userErrors")
                                ):
                                    # Verifica os erros de usuário
                                    for user_error in data["data"]["shopPolicyUpdate"][
                                        "userErrors"
                                    ]:
                                        erros.append(
                                            f"Erro ao atualizar a política '{titulo}': {user_error['message']}"
                                        )
                                else:
                                    print(
                                        f"Política '{titulo}' atualizada com sucesso!"
                                    )
                            else:
                                erros.append(
                                    f"Erro ao enviar a requisição para a política '{titulo}': Status Code - {response.status_code}, Response Text - {response.text}"
                                )

                        # Se houver erros, retorna todos como uma string única
                        if erros:
                            return "Erro: " + " | ".join(erros)
                        else:
                            return "Sucesso: Todas as políticas foram atualizadas com êxito."

                    except Exception as e:
                        return f"Erro: Exceção durante a atualização das políticas - {str(e)}"

                def adicionar_produtos_colecoes(colecoes, produtos):
                    try:
                        colecoes = shopify.CustomCollection.find()
                        produtos = shopify.Product.find()
                        headers = {
                            "X-Shopify-Access-Token": PRIVATE_APP_PASSWORD,
                            "Content-Type": "application/json",
                        }
                        mutation = """
                                    mutation collectionAddProductsV2($id: ID!, $productIds: [ID!]!) {
                                        collectionAddProductsV2(id: $id, productIds: $productIds) {
                                            job {
                                                done
                                                id
                                            }
                                            userErrors {
                                                field
                                                message
                                            }
                                        }
                                    }
                                    """

                        # Itera sobre cada coleção e produto
                        for colecao in colecoes:
                            collection_id = colecao.id
                            collection_handle = colecao.handle

                            # Filtra os produtos que correspondem à coleção
                            # pelo handle
                            ids_produtos_correspondentes = [
                                produto.id
                                for produto in produtos
                                if collection_handle in produto.tags
                            ]

                            # Adiciona todos os produtos correspondentes à
                            # coleção
                            for product_id in ids_produtos_correspondentes:
                                # Criação das variáveis da mutação
                                variables = {
                                    "id": f"gid://shopify/Collection/{collection_id}",
                                    "productIds": [
                                        f"gid://shopify/Product/{product_id}"
                                    ],
                                }

                                # Criação do corpo da requisição
                                body = {
                                    "query": mutation,
                                    "variables": variables,
                                }
                                # Realização da requisição à API da Shopify
                                response = requests.post(
                                    f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json",headers=headers,data=json.dumps(body),
                                )

                                # Verificação da resposta para cada produto
                                if response.status_code != 200:
                                    raise Exception(
                                        f"Erro ao adicionar produto {product_id} à coleção {collection_id}: {response.content}"
                                    )

                        return "Sucesso: Todos os produtos foram adicionados às coleções com êxito."

                    except Exception as e:
                        return f"Erro: {str(e)}"

                
                def editar_tema(urls_shopify, max_tentativas=5):
                    tentativa = 0
                    while tentativa < max_tentativas:
                        try:
                            temas = shopify.Theme.find()
                            tema_publicado = [tema for tema in temas if tema.role == "main"][0]
                            id_tema_publicado = tema_publicado.id
                            theme = shopify.Theme.find(id_tema_publicado)
                            settings_asset = shopify.Asset.find(
                                "config/settings_data.json", theme_id=theme.id
                            )

                            print("Arquivo de configurações encontrado.")

                            settings_data = settings_asset.value

                            # Verifica se settings_data é uma string e tenta carregar como JSON
                            if isinstance(settings_data, str):
                                settings_data = json.loads(settings_data)

                            print("Configurações carregadas com sucesso.")


                            # Armazenando os valores em variáveis
                            desktop1 = f"shopify://shop_images/{urls_shopify['desktop1'].split('/')[-1]}"
                            desktop2 = f"shopify://shop_images/{urls_shopify['desktop2'].split('/')[-1]}"
                            desktop3 = f"shopify://shop_images/{urls_shopify['desktop3'].split('/')[-1]}"
                            mobile1 = f"shopify://shop_images/{urls_shopify['mobile1'].split('/')[-1]}"
                            mobile2 = f"shopify://shop_images/{urls_shopify['mobile2'].split('/')[-1]}"
                            mobile3 = f"shopify://shop_images/{urls_shopify['mobile3'].split('/')[-1]}"
                            capa1 = f"shopify://shop_images/{urls_shopify['capa1'].split('/')[-1]}"
                            capa2 = f"shopify://shop_images/{urls_shopify['capa2'].split('/')[-1]}"
                            capa3 = f"shopify://shop_images/{urls_shopify['capa3'].split('/')[-1]}"
                            capa4 = f"shopify://shop_images/{urls_shopify['capa4'].split('/')[-1]}"
                            capa5 = f"shopify://shop_images/{urls_shopify['capa5'].split('/')[-1]}"
                            capa6 = f"shopify://shop_images/{urls_shopify['capa6'].split('/')[-1]}"

                            if tema_base == "Evolution":
                                #######################################################
                                # CORES DO TEMA ###########
                                #######################################################

                                ###### HEADER #########
                                settings_data["current"]["header_accent_color"] = acent
                                settings_data["current"]["header_background"] = header_background
                                settings_data["current"]["header_text_color"] = header_text_color
                                settings_data["current"]["accent_color"] = accent_color
                                settings_data["current"]["link_color"] = link_color
                                settings_data["current"]["header_light_text_color"] = header_light_text_color
                                settings_data["current"]["primary_button_background"] = primary_button_background
                                settings_data["current"]["secondary_button_background"] = secondary_button_background
                                settings_data["current"]["product_on_sale_accent"] = product_on_sale_accent
                                settings_data["current"]["product_cor_do_preco_semdesc"] = product_cor_do_preco_semdesc
                                settings_data["current"]["product_cor_do_preco"] = product_cor_do_preco
                                settings_data["current"]["product_cor_dos_titles"] = product_cor_dos_titles
                                settings_data["current"]["product_in_stock_color"] = product_in_stock_color

                                settings_data["current"]["sections"]["header"]["settings"]["background1"] = degrade1
                                settings_data["current"]["sections"]["header"]["settings"]["background2"] = degrade2
                                settings_data["current"]["sections"]["header"]["settings"]["background3"] = degrade3
                                settings_data["current"]["sections"]["header"]["settings"]["background4"] = degrade4
                                settings_data["current"]["sections"]["header"]["settings"]["backgroundlocal1"] = cor1
                                settings_data["current"]["sections"]["header"]["settings"]["backgroundlocal2"] = cor2
                                settings_data["current"]["sections"]["header"]["settings"]["colortextlocal"] = textcolor
                                settings_data["current"]["sections"]["header"]["settings"]["navigation_phone_number"] = whatsapp
                                settings_data["current"]["sections"]["header"]["settings"]["navigation_email"] = email_suporte

                                ###### FOOTER #########

                                #######################################################
                                ####    DADOS DA LOJA   ###########
                                #######################################################
                                ### CABEÇALHO ######
                                settings_data["current"]["sections"]["header"]["settings"]["navigation_phone_number"] = whatsapp
                                settings_data["current"]["sections"]["header"]["settings"]["navigation_email"] = email_suporte

                                """#### WHATSAPP #######
                                settings_data['current']['addzap'] = True
                                settings_data['current']['numberzap'] = tel_whats
                                settings_data['current']['titlezap'] = titlezap
                                """

                                ##### EDITA FOOTER ######
                                settings_data["current"]["sections"]["footer"]["blocks"]["text_image_t6X8qj"]["settings"]["content"] = footer
                                settings_data["current"]["footer_background_color"] = footer_background
                                settings_data["current"]["footer_heading_text_color"] = footer_text_color
                                settings_data["current"]["footer_body_text_color"] = footer_text_color

                                ##### INSERIR BANNERS ##########
                                settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-1"]["settings"]["image"] = desktop1
                                settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-1"]["settings"]["mobile_image"] = mobile1
                                settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-2"]["settings"]["image"] = desktop2
                                settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-2"]["settings"]["mobile_image"] = mobile2
                                settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-3"]["settings"]["image"] = desktop3
                                settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-3"]["settings"]["mobile_image"] = mobile3

                                #### COLEÇÔES E CAPAS  ###########
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_6EMmhW"]["settings"]["collection"] = primeiro
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_6EMmhW"]["settings"]["image"] = capa1
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_Anzt8E"]["settings"]["collection"] = segundo
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_Anzt8E"]["settings"]["image"] = capa2
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_NqzGkT"]["settings"]["collection"] = terceiro
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_NqzGkT"]["settings"]["image"] = capa3
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_X6edXj"]["settings"]["collection"] = quarto
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_X6edXj"]["settings"]["image"] = capa4
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_pPAM48"]["settings"]["collection"] = quinto
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_pPAM48"]["settings"]["image"] = capa5
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_j4N3Rp"]["settings"]["collection"] = sexto
                                settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_j4N3Rp"]["settings"]["image"] = capa6

                                #### FEATURED COLECC #####################
                                settings_data["current"]["sections"]["featured_collection_EXHCLn"]["settings"]["collection"] = primeiro
                                settings_data["current"]["sections"]["featured_collection_EXHCLn"]["settings"]["title"] = primeiro
                                settings_data["current"]["sections"]["featured_collection_jLrtgd"]["settings"]["collection"] = segundo
                                settings_data["current"]["sections"]["featured_collection_jLrtgd"]["settings"]["title"] = segundo
                                settings_data["current"]["sections"]["featured_collection_YAcpkY"]["settings"]["collection"] = terceiro
                                settings_data["current"]["sections"]["featured_collection_YAcpkY"]["settings"]["title"] = terceiro
                                settings_data["current"]["sections"]["featured_collection_NRCkDA"]["settings"]["collection"] = quarto
                                settings_data["current"]["sections"]["featured_collection_NRCkDA"]["settings"]["title"] = quarto
                                settings_data["current"]["sections"]["featured_collection_YhPALQ"]["settings"]["collection"] = quinto
                                settings_data["current"]["sections"]["featured_collection_YhPALQ"]["settings"]["title"] = quinto
                                settings_data["current"]["sections"]["collection_with_image_6pAgRL"]["settings"]["collection"] = sexto
                                settings_data["current"]["sections"]["collection_with_image_6pAgRL"]["settings"]["title"] = sexto
                                ##### COLEÇÃO image_6pAgRL#########
                                settings_data["current"]["sections"]["collection_with_image_6pAgRL"]["settings"]["background"] = cor1
                                settings_data["current"]["sections"]["collection_with_image_6pAgRL"]["settings"]["text_color"] = textcolor

                                settings_asset.value = json.dumps(settings_data)
                                settings_asset.save()

                                return "Sucesso: Estilização aplicada com êxito."
                            
                            elif tema_base == "Dropmeta":

                                #######################################################
                                ### CORES DO TEMA ###########
                                #######################################################
                                ###### BARRA DE ANUNCIOS #########
                                
                                settings_data["current"]["sections"]["announcement-bar"]["settings"]["background1"] = cor1
                                settings_data["current"]["sections"]["announcement-bar"]["settings"]["background2"] = cor2
                                settings_data["current"]["sections"]["announcement-bar"]["settings"]["text_color"] = textcolor
                               
                                
                                # Header
                                settings_data["current"]["header_accent_color"] = acent
                                settings_data["current"]["accent_color"] = accent_color                       
                                settings_data["current"]["link_color"] = link_color
                                settings_data["current"]["header_background"] = header_background
                                settings_data["current"]["header_text_color"] = header_text_color
                                settings_data["current"]["header_light_text_color"] = header_light_text_color
                                settings_data["current"]["primary_button_background"] = primary_button_background
                                settings_data["current"]["secondary_button_background"] = secondary_button_background
                                settings_data["current"]["product_on_sale_accent"] = product_on_sale_accent
                                settings_data["current"]["product_cor_do_preco_semdesc"] = product_cor_do_preco_semdesc
                                settings_data["current"]["product_cor_do_preco"] = product_cor_do_preco
                                settings_data["current"]["product_cor_dos_titles"] = product_cor_dos_titles
                                settings_data["current"]["product_in_stock_color"] = product_in_stock_color
                                settings_data["current"]["sections"]["header"]["settings"]["background1"] = degrade1
                                settings_data["current"]["sections"]["header"]["settings"]["background2"] = degrade2
                                settings_data["current"]["sections"]["header"]["settings"]["background3"] = degrade3 
                                settings_data["current"]["sections"]["header"]["settings"]["background4"] = degrade4
                                
                                    
                                ###### FOOTER #########
                                settings_data["current"]["sections"]["footer"]["settings"]["background1"] = degradefooter1
                                settings_data["current"]["sections"]["footer"]["settings"]["background2"] = degradefooter2
                                settings_data["current"]["sections"]["footer"]["settings"]["background3"] = degradefooter3
                                settings_data["current"]["sections"]["footer"]["settings"]["background4"] = degradefooter4
                                settings_data["current"]["footer_background"] = footer_background
                                settings_data["current"]["footer_text_color"] = footer_text_color

                                ##### COLEÇÃO #########
                                
                                settings_data["current"]["sections"]["1649913264519d35a7"]["settings"]["background1"] = cor1
                                settings_data["current"]["sections"]["1649913264519d35a7"]["settings"]["background2"] = cor2
                                settings_data["current"]["sections"]["1649913264519d35a7"]["settings"]["text_color"] = textcolor
                                
                                    
                                #######################################################
                                ####    DADOS DA LOJA   ###########
                                #######################################################
                                ### CABEÇALHO ######
                                
                                settings_data["current"]["sections"]["header"]["settings"]["navigation_phone_number"] = whatsapp
                                settings_data["current"]["sections"]["header"]["settings"]["navigation_email"] = email_suporte

                                #### WHATSAPP #######
                                settings_data["current"]["addzap"] = True
                                settings_data["current"]["numberzap"] = tel_whats
                                settings_data["current"]["titlezap"] = titlezap
                              
                                ##### EDITA FOOTER ######
                                settings_data["current"]["sections"]["footer"]["blocks"]["c9a6c378-573f-4c54-9fc6-5ca6e279f3f2"]["settings"]["content"] = footer

                                ##### INSERIR BANNERS ##########
                                # banner 1
                                settings_data["current"]["sections"]["slideshow"]["blocks"]["4665d2ed-db3b-479e-8984-d272fdfab8d8"]["settings"]["image"] = desktop1
                                settings_data["current"]["sections"]["slideshow"]["blocks"]["4665d2ed-db3b-479e-8984-d272fdfab8d8"]["settings"]["mobile_image"] = mobile1
                                # banner 2
                                settings_data["current"]["sections"]["slideshow"]["blocks"]["bbea030d-de91-4721-9ea1-9660f647bbd7"]["settings"]["image"] = desktop2
                                settings_data["current"]["sections"]["slideshow"]["blocks"]["bbea030d-de91-4721-9ea1-9660f647bbd7"]["settings"]["mobile_image"] = mobile2
                                # banner 3
                                if pais == "Espanha":
                                    settings_data["current"]["sections"]["slideshow_7n83kx"]["blocks"]["image_H6WbRH"]["settings"]["image"] = desktop3
                                    settings_data["current"]["sections"]["slideshow_7n83kx"]["blocks"]["image_H6WbRH"]["settings"]["mobile_image"] = mobile3
                                    
                                elif pais == "EUA":
                                    settings_data["current"]["sections"]["slideshow_37CJTG"]["blocks"]["image_CQGpKw"]["settings"]["image"] = desktop3
                                    settings_data["current"]["sections"]["slideshow_37CJTG"]["blocks"]["image_CQGpKw"]["settings"]["mobile_image"] = mobile3


                                elif pais == "Portugal":
                                    settings_data["current"]["sections"]["slideshow_DViiTn"]["blocks"]["image_WiCjtP"]["settings"]["image"] = desktop3
                                    settings_data["current"]["sections"]["slideshow_DViiTn"]["blocks"]["image_WiCjtP"]["settings"]["mobile_image"] = mobile3


                                elif pais == "Brasil":
                                    settings_data["current"]["sections"]["slideshow_NcxY8c"]["blocks"]["image_8RNyne"]["settings"]["image"] = desktop3
                                    settings_data["current"]["sections"]["slideshow_NcxY8c"]["blocks"]["image_8RNyne"]["settings"]["mobile_image"] = mobile3

                                else:
                                    settings_data["current"]["sections"]["slideshow_LR4JhV"]["blocks"]["image_7Y9Pdh"]["settings"]["image"] = desktop3
                                    settings_data["current"]["sections"]["slideshow_LR4JhV"]["blocks"]["image_7Y9Pdh"]["settings"]["mobile_image"] = mobile3
                                
                                #### COLEÇÔES E CAPAS  ###########
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-1"]["settings"]["collection"] = primeiro
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-1"]["settings"]["image"] = capa1
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-2"]["settings"]["collection"] = segundo
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-2"]["settings"]["image"] = capa2
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-3"]["settings"]["collection"] = terceiro
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-3"]["settings"]["image"] = capa3
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-5"]["settings"]["collection"] = quarto
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-5"]["settings"]["image"] = capa4
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["d5a5159b-779d-45d2-871b-4a5e622bb3b2"]["settings"]["collection"] = quinto
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["d5a5159b-779d-45d2-871b-4a5e622bb3b2"]["settings"]["image"] = capa5
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["0c56a93b-d66f-4823-ae91-bd01bb634a00"]["settings"]["collection"] = sexto
                                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["0c56a93b-d66f-4823-ae91-bd01bb634a00"]["settings"]["image"] = capa6
                                
                                #### FEATURED COLECC #####################
                                settings_data["current"]["sections"]["16153539675ef431a7"]["settings"]["collection"] = sexto
                                settings_data["current"]["sections"]["16153539675ef431a7"]["settings"]["title"] = sex
                                settings_data["current"]["sections"]["16153541158fc73818"]["settings"]["collection"] = primeiro
                                settings_data["current"]["sections"]["16153541158fc73818"]["settings"]["title"] = pri
                                settings_data["current"]["sections"]["1649913264519d35a7"]["settings"]["collection"] = segundo      
                                settings_data["current"]["sections"]["1649913264519d35a7"]["settings"]["title"] = seg
                                settings_data["current"]["sections"]["7196b9c9-7ec2-46f0-87be-f84dcc1d14b5"]["settings"]["collection"] = terceiro
                                settings_data["current"]["sections"]["7196b9c9-7ec2-46f0-87be-f84dcc1d14b5"]["settings"]["title"] = ter
                                settings_data["current"]["sections"]["9dd41389-f9fe-453e-b65e-1cc7383b8c1f"]["settings"]["collection"] = quarto
                                settings_data["current"]["sections"]["9dd41389-f9fe-453e-b65e-1cc7383b8c1f"]["settings"]["title"] = qua
                                settings_data["current"]["sections"]["329bc2be-34eb-46d8-8601-39d8eda721d9"]["settings"]["collection"] = quinto
                                settings_data["current"]["sections"]["329bc2be-34eb-46d8-8601-39d8eda721d9"]["settings"]["title"] = qui
                                
                                
                                
                                settings_asset.value = json.dumps(settings_data)
                                settings_asset.save()

                                return "Sucesso: Estilização aplicada com êxito."

                        except Exception as e:
                            tentativa += 1
                            print(f"Tentativa {tentativa} falhou. Erro: {str(e)}")
                            time.sleep(2)  # Aguarda 2 segundos antes de tentar novamente
                    
                    return "Erro: Número máximo de tentativas atingido."
            

                # Chamar funções
                if not formulario_user.produtos:
                    resposta = importar_produtos(df, pais)
                    formulario_user.produtos = True
                    formulario_user.resposta_produtos = resposta
                    formulario_user.save()
                if not formulario_user.tema:
                    resposta = importar_tema(Tema)
                    formulario_user.tema = True
                    formulario_user.resposta_tema = resposta
                    formulario_user.save()
                
                # if not formulario_user.banners:
                #     resposta = subir_banners(urls)
                #     formulario_user.banners = True
                #     formulario_user.resposta_banners = resposta
                #     formulario_user.save()
                    
                if not formulario_user.colecoes:
                    resposta = criar_colecoes(colecoes)
                    formulario_user.colecoes = True
                    formulario_user.resposta_colecoes = resposta
                    formulario_user.save()
                if not formulario_user.paginas:
                    resposta = criar_paginas(paginas)
                    formulario_user.paginas = True
                    formulario_user.resposta_paginas = resposta
                    formulario_user.save()
                if not formulario_user.politicas:
                    resposta = atualizar_politicas()
                    formulario_user.politicas = True
                    formulario_user.resposta_politicas = resposta
                    formulario_user.save()
                if not formulario_user.prod_col:
                    resposta = adicionar_produtos_colecoes(colecoes, produtos)
                    formulario_user.prod_col = True
                    formulario_user.resposta_prod_col = resposta
                    formulario_user.save()
                if not formulario_user.estilizacao:
                    resposta = editar_tema(urls_shopify)
                    formulario_user.estilizacao = True
                    formulario_user.lojaproduzida = True
                    formulario_user.resposta_estilizacao = (
                        resposta  # Campo novo no modelo para armazenar a resposta
                    )
                    formulario_user.save()

                

                # Após a edição do tema, deletar imagens no Cloudinary
                for public_id in request.session["public_ids_to_delete"]:
                    delete_image_from_cloudinary(public_id)

                request.session["public_ids_to_delete"] = []
                shopify.ShopifyResource.clear_session()

                # Atualiza o status para 'completed'
                formulario_user.status = "finalizado"  # Ao final de tudo
                formulario_user.save()
                print(f"loja de {empresa} concluida com sucesso!")

                return JsonResponse(
                    {"status": "success", "message": "Tarefas executadas com sucesso!"}
                )
            else:
                # Caso o usuário não seja encontrado
                task_status[username] = "user_not_found"
                return JsonResponse(
                    {"status": "error", "message": "Usuário não encontrado."},
                    status=404,
                )

        except Exception as e:
            # Em caso de erro
            task_status[username] = "error"
            formulario_user.status = "erro"
            formulario_user.save()
            return JsonResponse(
                {
                    "status": "error",
                    "message": f"Erro ao executar as tarefas: {str(e)}",
                },
                status=500,
            )

    # Se não for um POST, retornar erro de método não permitido
    return JsonResponse(
        {"status": "error", "message": "Método não permitido."}, status=405
    )

def testar(request):
    if request.method == "POST":
        SHOP_URL = "imagensbanners.myshopify.com"
        PRIVATE_APP_PASSWORD = "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"
        API_KEY = "a8512d5f934bffb3258a3c2f866e77a9"
        API_SECRET = "c7d3fb82de5bf0d9a1f748459e1a03cb"
        API_VERSION = "2024-07"

        # Configuração da sessão Shopify
        shopify.Session.setup(api_key=API_KEY, secret=API_SECRET)
        session = shopify.Session(SHOP_URL, API_VERSION, PRIVATE_APP_PASSWORD)
        shopify.ShopifyResource.activate_session(session)

        cor1 = "#17453F"

        # Lista de URLs para diferentes categorias
        urls = {
            "desktop1": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "desktop2": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "desktop3": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "mobile1": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "mobile2": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "mobile3": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "capa1": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "capa2": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "capa3": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "capa4": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "capa5": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
            "capa6": "https://res.cloudinary.com/hnrqudfke/image/upload/v1734034004/desktop1_2_nfbkpr_w5uq2b.svg",
        }

        def download_svg(url):
            response = requests.get(url)
            if response.status_code == 200:
                local_path = url.split("/")[-1]
                with open(local_path, "wb") as file:
                    file.write(response.content)
                return local_path
            else:
                raise Exception(f"Erro ao baixar o SVG: {response.status_code}")

        def edit_svg_colors(svg_path, color_mapping):
            try:
                tree = ET.parse(svg_path)
                root = tree.getroot()
                namespace = {"svg": "http://www.w3.org/2000/svg"}

                for element in root.findall(".//svg:*", namespaces=namespace):
                    fill = element.attrib.get("fill")
                    if fill and fill in color_mapping:
                        element.set("fill", color_mapping[fill])

                edited_path = "edited_" + svg_path
                tree.write(edited_path)
                return edited_path
            except Exception as e:
                raise Exception(f"Erro ao editar o SVG: {e}")

        def validate_svg(svg_path):
            try:
                tree = ET.parse(svg_path)
                root = tree.getroot()
                return root.tag == "{http://www.w3.org/2000/svg}svg"
            except Exception:
                return False

        def convert_svg_to_png(svg_path):
            try:
                response = cloudinary.uploader.upload(svg_path, resource_type="image", format="png", folder="banners_e_capas")
                return response["secure_url"], response["public_id"]
            except Exception as e:
                raise Exception(f"Erro ao converter o SVG para PNG: {e}")

        def delete_image_from_cloudinary(public_id):
            try:
                cloudinary.uploader.destroy(public_id)
                print(f"Imagem {public_id} deletada com sucesso no Cloudinary.")
            except Exception as e:
                print(f"Erro ao deletar imagem no Cloudinary: {e}")

        def upload_png_to_shopify(png_url, file_name):
            mutation = """
                mutation fileCreate($files: [FileCreateInput!]!) {
                    fileCreate(files: $files) {
                        files {
                            id
                            alt
                        }
                        userErrors {
                            message
                        }
                    }
                }
            """
            file_data = {
                "files": [
                    {"alt": file_name, "originalSource": png_url}
                ]
            }

            response = requests.post(
                f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json",
                json={"query": mutation, "variables": file_data},
                headers={"X-Shopify-Access-Token": PRIVATE_APP_PASSWORD},
            )
            if response.status_code != 200:
                raise Exception(f"Erro ao enviar para Shopify: {response.text}")

        if "public_ids_to_delete" not in request.session:
            request.session["public_ids_to_delete"] = []

        urls_shopify = {}

        for name, url in urls.items():
            try:
                svg_path = download_svg(url)
                edited_svg_path = edit_svg_colors(svg_path, {"#000000": cor1})

                if validate_svg(edited_svg_path):
                    png_url, public_id = convert_svg_to_png(edited_svg_path)
                    upload_png_to_shopify(png_url, name)
                    urls_shopify[name] = png_url
                    request.session["public_ids_to_delete"].append(public_id)
                else:
                    raise Exception(f"SVG inválido: {edited_svg_path}")
            except Exception as e:
                print(f"Erro no processamento de {name}: {e}")


        def editar_tema(urls_shopify):
            try:
                temas = shopify.Theme.find()
                tema_publicado = next((tema for tema in temas if tema.role == "main"), None)
                if not tema_publicado:
                    raise Exception("Tema principal não encontrado")

                settings_asset = shopify.Asset.find("config/settings_data.json", theme_id=tema_publicado.id)
                settings_data = json.loads(settings_asset.value)

                # Armazenando os valores em variáveis
                desktop1 = f"shopify://shop_images/{urls_shopify['desktop1'].split('/')[-1]}"
                desktop2 = f"shopify://shop_images/{urls_shopify['desktop2'].split('/')[-1]}"
                desktop3 = f"shopify://shop_images/{urls_shopify['desktop3'].split('/')[-1]}"
                mobile1 = f"shopify://shop_images/{urls_shopify['mobile1'].split('/')[-1]}"
                mobile2 = f"shopify://shop_images/{urls_shopify['mobile2'].split('/')[-1]}"
                mobile3 = f"shopify://shop_images/{urls_shopify['mobile3'].split('/')[-1]}"
                capa1 = f"shopify://shop_images/{urls_shopify['capa1'].split('/')[-1]}"
                capa2 = f"shopify://shop_images/{urls_shopify['capa2'].split('/')[-1]}"
                capa3 = f"shopify://shop_images/{urls_shopify['capa3'].split('/')[-1]}"
                capa4 = f"shopify://shop_images/{urls_shopify['capa4'].split('/')[-1]}"
                capa5 = f"shopify://shop_images/{urls_shopify['capa5'].split('/')[-1]}"
                capa6 = f"shopify://shop_images/{urls_shopify['capa6'].split('/')[-1]}"

                # Atualizando o JSON com as variáveis
                settings_data["current"]["sections"]["slideshow"]["blocks"]["4665d2ed-db3b-479e-8984-d272fdfab8d8"]["settings"]["image"] = desktop1
                settings_data["current"]["sections"]["slideshow"]["blocks"]["4665d2ed-db3b-479e-8984-d272fdfab8d8"]["settings"]["mobile_image"] = mobile1

                settings_data["current"]["sections"]["slideshow"]["blocks"]["bbea030d-de91-4721-9ea1-9660f647bbd7"]["settings"]["image"] = desktop2
                settings_data["current"]["sections"]["slideshow"]["blocks"]["bbea030d-de91-4721-9ea1-9660f647bbd7"]["settings"]["mobile_image"] = mobile2

                settings_data["current"]["sections"]["slideshow"]["blocks"]["fa2eac7c-9fbf-4ea9-a172-eca738a9012d"]["settings"]["image"] = desktop3
                settings_data["current"]["sections"]["slideshow"]["blocks"]["fa2eac7c-9fbf-4ea9-a172-eca738a9012d"]["settings"]["mobile_image"] = mobile3

                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-1"]["settings"]["image"] = capa1
                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-2"]["settings"]["image"] = capa2
                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-3"]["settings"]["image"] = capa3
                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-5"]["settings"]["image"] = capa4
                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["d5a5159b-779d-45d2-871b-4a5e622bb3b2"]["settings"]["image"] = capa5
                settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["0c56a93b-d66f-4823-ae91-bd01bb634a00"]["settings"]["image"] = capa6

                settings_asset.value = json.dumps(settings_data)
                settings_asset.save()
            except Exception as e:
                raise Exception(f"Erro ao editar tema: {e}")

       
       
        editar_tema(urls_shopify)

        # Após a edição do tema, deletar imagens no Cloudinary
        for public_id in request.session["public_ids_to_delete"]:
            delete_image_from_cloudinary(public_id)

        request.session["public_ids_to_delete"] = []
        shopify.ShopifyResource.clear_session()

        return render(request, "testar.html", {"result": "Sucesso!"})

    return render(request, "testar.html", {"result": ""})
        


