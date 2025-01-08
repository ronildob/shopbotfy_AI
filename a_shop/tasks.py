from datetime import timedelta
from django.utils import timezone
import time
import requests
from shopbotfy.models import Agencia
from shop_api.models import AAFormuUserer, Pendencia, FilaDeEspera
from .helpers import handle_produz_request
from celery.exceptions import MaxRetriesExceededError
from celery import shared_task
import logging
import pandas as pd
import shopify 
import unidecode
from docx import Document
import json
import xml.etree.ElementTree as ET
import cloudinary
from cloudinary.uploader import upload
from cloudinary.utils import cloudinary_url

from django.core.cache import cache

logger = logging.getLogger(__name__)


@shared_task
def verificar_planos_status():
    now = timezone.now()
    agencias = Agencia.objects.all()
    for agencia in agencias:
        data_referencia = agencia.data_renovacao or agencia.data_criacao

        if agencia.plano == "teste" and now > agencia.data_criacao + timedelta(
            hours=72
        ):
            agencia.plano = "vencido"
            agencia.save()

        elif agencia.plano.startswith("atraso"):
            dias_atraso = int(agencia.plano.replace("atraso", ""))
            if now > data_referencia + timedelta(days=dias_atraso + 1):
                proximo_atraso = f"atraso{dias_atraso + 1}"
                if proximo_atraso in dict(Agencia.PLANO_CHOICES):
                    agencia.plano = proximo_atraso
                else:
                    agencia.plano = "vencido"
                agencia.save()

        elif agencia.plano == "ativo" and now > data_referencia + timedelta(days=30):
            agencia.plano = "atraso1"
            agencia.save()


@shared_task
def verificar_planos():
    print("Verificando planos...")
    pendencias = Pendencia.objects.all()
    
    for pendencia in pendencias:
        username = pendencia.username
        print(f"Processando usuário: {username}")
        
        # Incrementa o contador de verificações
        pendencia.verificacoes += 1
        pendencia.save()  # Salva o contador atualizado no banco de dados
        
        # Busca o usuário correspondente na tabela AAFormuUserer
        formulario_user = AAFormuUserer.objects.filter(username=username).first()

        if formulario_user:
            SHOP_URL = formulario_user.url_loja
            PRIVATE_APP_PASSWORD = formulario_user.token_senha
            API_KEY = formulario_user.chave_de_api
            API_VERSION = '2024-07'
            
            # Conecta com a API Shopify para verificar o plano
            endpoint = f"https://{SHOP_URL}/admin/api/{API_VERSION}/shop.json"
            headers = {'Content-Type': 'application/json'}
            auth = (API_KEY, PRIVATE_APP_PASSWORD)
            try:
                response = requests.get(endpoint, headers=headers, auth=auth)
                if response.status_code == 200:
                    shop_info = response.json()
                    plan_name = shop_info.get('shop', {}).get('plan_name')
                    print(f"Plano encontrado para {username}: {plan_name}")

                    if plan_name in ['affiliate', 'basic', 'Shopify', 'Advanced', "unlimited"]:
                        formulario_user.status = 'iniciado'
                        formulario_user.save()
                        pendencia.delete()
                        handle_produz_request(username)
                       
                    elif plan_name in ['starter', 'trial']:
                        formulario_user.status = 'pendencia'
                        formulario_user.save()
                        print(f"Usuário {username} com plano {plan_name}, status mudado para 'pendencia'.")
                    else:
                        formulario_user.status = 'pendencia'
                        formulario_user.save()
                        print(f"Usuário {username} com plano desconhecido, status mudado para 'pendencia'.")
                else:
                    print(f"Erro na resposta da API Shopify para {username}: {response.status_code}")
            except requests.exceptions.RequestException as e:
                print(f"Erro ao conectar à API da Shopify para {username}: {e}")
 

@shared_task(bind=True, max_retries=3, default_retry_delay=15)  # Tenta novamente até 3 vezes, com intervalo de 60 segundos

def processar_usuario(self, username):
    try:
        formulario_user = AAFormuUserer.objects.filter(username=username).first()

        if formulario_user:
            # Atualiza o status para 'iniciado'
            formulario_user.status = 'iniciado'
            formulario_user.save()
                
            SHOP_URL = formulario_user.url_loja
            PRIVATE_APP_PASSWORD = formulario_user.token_senha
            API_KEY = formulario_user.chave_de_api
            API_SECRET = formulario_user.chave_secreta
            estilo = formulario_user.nicho
            pais = formulario_user.pais
            tema_base = formulario_user.tema_base
            paleta = formulario_user.cor
            telefone = formulario_user.telefone
            email_suporte = formulario_user.email_suporte
            empresa = formulario_user.empresa
            business_hours = formulario_user.business_hours
            whatsapp = str(telefone)

            print(SHOP_URL, PRIVATE_APP_PASSWORD, API_KEY)
            API_VERSION = '2024-07'
            cod_pais = {
                "EUA": "+1",
                "Argentina": "+54",
                "Bolívia": "+591",
                "Brasil": "+55",
                "Chile": "+56",
                "Colômbia": "+57",
                "Equador": "+593",
                "Paraguai": "+595",
                "Peru": "+51",
                "Uruguai": "+598",
                "Venezuela": "+58",
                "Paquistão": "+92",
                "Índia": "+91",
                "Portugal": "+351",
                "México" : "+52",
                "Espanha" : "+34"
            }
            if pais in cod_pais:
                tel_whats = cod_pais[pais] + str(telefone)
                whatsapp_completo = cod_pais[pais] + whatsapp
            
            ##############################################################
            #### CORES ######
            if paleta == "Paleta de Cor 1":  # LARANJA
                # barra de anúncio
                cor1 = "#BB6A09"
                cor2 = "#FE8B01"
                textcolor = "#FFFFFF"

                # Cabeçalho/barra colorida
                degrade1 = "#BB6A09"
                degrade2 = "#BB6A09"
                degrade3 = "#BB6A09"
                degrade4 = "#BB6A09"

                # Cores gerais
                # Geral
                acent = "#0ABF58"
                link_color = "#0ABF58"
                primary_button_background = "#0ABF58"
                secondary_button_background = "#0ABF58"

                # Rodapé
                footer_background = "#FE8B01"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#FE8B01"
                header_text_color = "#F7F7F7"
                header_light_text_color = "#F7F7F7"
                accent_color = "#0ABF58"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # footer
                degradefooter1 = "#BB6A09"
                degradefooter2 = "#BB6A09"
                degradefooter3 = "#BB6A09"
                degradefooter4 = "#BB6A09"
                #############################

            elif paleta == "Paleta de Cor 2":  # VERDE
                # barra de anúncio
                cor1 = "#60AD3F"
                cor2 = "#0097B2"
                textcolor = "#FFFFFF"

                # Cabeçalho/barra colorida
                degrade1 = "#7ED957"
                degrade2 = "#0097B2"
                degrade3 = "#0097B2"
                degrade4 = "#7ED957"

                # Cores gerais
                # Geral
                acent = "#0097B2"
                link_color = "#860D7E"
                primary_button_background = "#860D7E"
                secondary_button_background = "#860D7E"

                # Rodapé
                footer_background = "7ED957"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#7ED957"
                header_text_color = "#F7F7F7"
                header_light_text_color = "#F7F7F7"
                accent_color = "#0097B2"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # footer
                degradefooter1 = "#7ED957"
                degradefooter2 = "#0097B2"
                degradefooter3 = "#0097B2"
                degradefooter4 = "#7ED957"
                #############################

            elif paleta == "Paleta de Cor 3":  # ROXO
                # barra de anúncio
                cor1 = "#860D7E"
                cor2 = "#600D5A"
                textcolor = "#FFFFFF"

                # Cabeçalho/barra colorida
                degrade1 = "#D600FF"
                degrade2 = "#8610D8"
                degrade3 = "#D600FF"
                degrade4 = "#8610D8"

                # Cores gerais
                # Geral
                acent = "#860D7E"
                link_color = "#860D7E"
                primary_button_background = "#860D7E"
                secondary_button_background = "#860D7E"

                # Rodapé
                footer_background = "600D5A "
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#600D5A"
                header_text_color = "#F7F7F7"
                header_light_text_color = "#F7F7F7"
                accent_color = "#860D7E"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # footer
                degradefooter1 = "#D600FF"
                degradefooter2 = "#8610D8"
                degradefooter3 = "#D600FF"
                degradefooter4 = "#8610D8"
                #############################

            elif paleta == "Paleta de Cor 4":  # PRETO
                # barra de anúncio
                cor1 = "#393939"
                cor2 = "#000000"
                textcolor = "#FFFFFF"

                # Cabeçalho/barra colorida
                degrade1 = "#393939"
                degrade2 = "#393939"
                degrade3 = "#393939"
                degrade4 = "#393939"

                # Cores gerais
                # Geral
                acent = "#393939"
                link_color = "#393939"
                primary_button_background = "#393939"
                secondary_button_background = "#393939"

                # Rodapé
                footer_background = "#000000"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#000000"
                header_text_color = "#F7F7F7"
                header_light_text_color = "#F7F7F7"
                accent_color = "#393939"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # footer
                degradefooter1 = "#393939"
                degradefooter2 = "#393939"
                degradefooter3 = "#393939"
                degradefooter4 = "#393939"
                #############################

            elif paleta == "Paleta de Cor 5":  # AZUL
                # barra de anúncio
                cor1 = "#5DE0E6"
                cor2 = "#073D86"
                textcolor = "#FFFFFF"

                # Cabeçalho/barra colorida
                degrade1 = "#5DE0E6"
                degrade2 = "#004AAD"
                degrade3 = "#5DE0E6"
                degrade4 = "#004AAD"

                # Cores gerais
                # Geral
                acent = "#5DE0E6"
                link_color = "#5DE0E6"
                primary_button_background = "#5DE0E6"
                secondary_button_background = "#5DE0E6"

                # Rodapé
                footer_background = "#004AAD"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#004AAD"
                header_text_color = "#F7F7F7"
                header_light_text_color = "#F7F7F7"
                accent_color = "#5DE0E6"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # footer
                degradefooter1 = "#5DE0E6"
                degradefooter2 = "#004AAD"
                degradefooter3 = "#5DE0E6"
                degradefooter4 = "#004AAD"
                #############################

            elif paleta == "Paleta de Cor 6":  # VERMELHO
                # barra de anúncio
                cor1 = "#B42929"
                cor2 = "#740606"
                textcolor = "#FFFFFF"

                # Cabeçalho/barra colorida
                degrade1 = "#B42929"
                degrade2 = "#B42929"
                degrade3 = "#740606"
                degrade4 = "#740606"

                # Cores gerais
                # Geral
                acent = "#740606"
                link_color = "#740606"
                primary_button_background = "#B42929"
                secondary_button_background = "#B42929"

                # Rodapé
                footer_background = "#000000"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#FF3131"
                header_text_color = "#F7F7F7"
                header_light_text_color = "#F7F7F7"
                accent_color = "#740606"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # footer
                degradefooter1 = "#B42929"
                degradefooter2 = "#B42929"
                degradefooter3 = "#740606"
                degradefooter4 = "#740606"

            elif paleta == "Paleta de Cor 7":  # AZUL E AMARELO
                # barra de anúncio
                cor1 = "#0F0F33"
                cor2 = "#0F0F33"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#D3991F"
                degrade2 = "#D3991F"
                degrade3 = "#D3991F"
                degrade4 = "#D3991F"

                # Cores gerais
                # Geral
                acent = "#0F0F33"
                link_color = "#E7DCDC"
                # Botão primário
                primary_button_background = "#000000"
                # Botão secundário
                secondary_button_background = "#0F0F33"

                # Rodapé
                footer_background = "#0F0F33"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#F2F2F2"
                header_text_color = "#0F0F33"
                header_light_text_color = "#0F0F33"
                accent_color = "#D3991F"

                # Produtos
                product_on_sale_accent = "#0F0F33"  # Barra economize
                product_cor_do_preco_semdesc = "#0F0F33"  # Cor do preço
                product_cor_do_preco = "#0F0F33"  # Cor do preço com desconto
                product_cor_dos_titles = "#4D4D4B"  # Títulos das informações
                product_in_stock_color = "#4D4D4B"  # Em estoque

                # Footer
                degradefooter1 = "#D3991F"
                degradefooter2 = "#D3991F"
                degradefooter3 = "#D3991F"
                degradefooter4 = "#D3991F"

            elif paleta == "Paleta de Cor 8":  # CINZA ESCURO CINZA CLARO
                # barra de anúncio
                cor1 = "#4D4D4B"
                cor2 = "#4D4D4B"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#4D4D4B"
                degrade2 = "#4D4D4B"
                degrade3 = "#4D4D4B"
                degrade4 = "#4D4D4B"

                # Cores gerais
                # Geral
                acent = "#4D4D4B"
                link_color = "#4D4D4B"
                # Botão primário
                primary_button_background = "#4D4D4B"
                # Botão secundário
                secondary_button_background = "#4D4D4B"

                # Rodapé
                footer_background = "#4D4D4B"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#F2F2F2"
                header_text_color = "#4D4D4B"
                header_light_text_color = "#4D4D4B"
                accent_color = "#AFAFAF"

                # Produtos
                product_on_sale_accent = "#4D4D4B"  # Barra economize
                product_cor_do_preco_semdesc = "#4D4D4B"  # Cor do preço
                product_cor_do_preco = "#4D4D4B"  # Cor do preço com desconto
                product_cor_dos_titles = "#4D4D4B"  # Títulos das informações
                product_in_stock_color = "#4D4D4B"  # Em estoque

                # Footer
                degradefooter1 = "#AFAFAF"
                degradefooter2 = "#AFAFAF"
                degradefooter3 = "#AFAFAF"
                degradefooter4 = "#AFAFAF"

            elif paleta == "Paleta de Cor 9":  # MARROM E AZUL
                # barra de anúncio
                cor1 = "#042138"
                cor2 = "#042138"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#042138"
                degrade2 = "#042138"
                degrade3 = "#042138"
                degrade4 = "#042138"

                # Cores gerais
                # Geral
                acent = "#042138"
                link_color = "#042138"
                # Botão primário
                primary_button_background = "#042138"
                # Botão secundário
                secondary_button_background = "#042138"

                # Rodapé
                footer_background = "#042138"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#F2F2F2"
                header_text_color = "#042138"
                header_light_text_color = "#042138"
                accent_color = "#9D7638"

                # Produtos
                product_on_sale_accent = "#042138"  # Barra economize
                product_cor_do_preco_semdesc = "#042138"  # Cor do preço
                product_cor_do_preco = "#042138"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#042138"  # Em estoque

                # Footer
                degradefooter1 = "#9D7638"
                degradefooter2 = "#9D7638"
                degradefooter3 = "#9D7638"
                degradefooter4 = "#9D7638"

            elif paleta == "Paleta de Cor 10":  # MARROM ESCURO E BEGE
                # barra de anúncio
                cor1 = "#2D251C"
                cor2 = "#2D251C"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#AD9D95"
                degrade2 = "#AD9D95"
                degrade3 = "#AD9D95"
                degrade4 = "#AD9D95"

                # Cores gerais
                # Geral
                acent = "#2D251C"
                link_color = "#2D251C"
                # Botão primário
                primary_button_background = "#2D251C"
                # Botão secundário
                secondary_button_background = "#2D251C"

                # Rodapé
                footer_background = "#2D251C"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#F2F2F2"
                header_text_color = "#2D251C"
                header_light_text_color = "#2D251C"
                accent_color = "#AD9D95"

                # Produtos
                product_on_sale_accent = "#042138"  # Barra economize
                product_cor_do_preco_semdesc = "#042138"  # Cor do preço
                product_cor_do_preco = "#042138"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#042138"  # Em estoque

                # Footer
                degradefooter1 = "#AD9D95"
                degradefooter2 = "#AD9D95"
                degradefooter3 = "#AD9D95"
                degradefooter4 = "#AD9D95"

            elif paleta == "Paleta de Cor 11":  # MARROM e Cinza
                # barra de anúncio
                cor1 = "#BE8F59"
                cor2 = "#BE8F59"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#BE8F59"
                degrade2 = "#BE8F59"
                degrade3 = "#BE8F59"
                degrade4 = "#BE8F59"

                # Cores gerais
                # Geral
                acent = "#BE8F59"
                link_color = "#BE8F59"
                # Botão primário
                primary_button_background = "#BE8F59"
                # Botão secundário
                secondary_button_background = "#BE8F59"

                # Rodapé
                footer_background = "#555555"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#F2F2F2"
                header_text_color = "#BE8F59"
                header_light_text_color = "#BE8F59"
                accent_color = "#CDCDCD"

                # Produtos
                product_on_sale_accent = "#BE8F59"  # Barra economize
                product_cor_do_preco_semdesc = "#BE8F59"  # Cor do preço
                product_cor_do_preco = "#BE8F59"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # Footer
                degradefooter1 = "#BE8F59"
                degradefooter2 = "#BE8F59"
                degradefooter3 = "#BE8F59"
                degradefooter4 = "#BE8F59"

            elif paleta == "Paleta de Cor 12":  # PRETO E BRANCO
                # barra de anúncio
                cor1 = "#000000"
                cor2 = "#000000"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#000000"
                degrade2 = "#000000"
                degrade3 = "#000000"
                degrade4 = "#000000"

                # Cores gerais
                # Geral
                acent = "#000000"
                link_color = "#E7DCDC"
                # Botão primário
                primary_button_background = "#000000"
                # Botão secundário
                secondary_button_background = "#000000"

                # Rodapé
                footer_background = "#000000"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#F2F2F2"
                header_text_color = "#4D4D4B"
                header_light_text_color = "#000000"
                accent_color = "#000000"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#000000"  # Cor do preço
                product_cor_do_preco = "#000000"  # Cor do preço com desconto
                product_cor_dos_titles = "#4D4D4B"  # Títulos das informações
                product_in_stock_color = "#4D4D4B"  # Em estoque

                # Footer
                degradefooter1 = "#AFAFAF"
                degradefooter2 = "#AFAFAF"
                degradefooter3 = "#AFAFAF"
                degradefooter4 = "#AFAFAF"

            elif paleta == "Paleta de Cor 13":  # verde e azul
                # barra de anúncio
                cor1 = "#516075"
                cor2 = "#516075"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#67A896"
                degrade2 = "#67A896"
                degrade3 = "#67A896"
                degrade4 = "#67A896"

                # Cores gerais
                # Geral
                acent = "#67A896"
                link_color = "#67A896"
                # Botão primário
                primary_button_background = "#67A896"
                # Botão secundário
                secondary_button_background = "#67A896"

                # Rodapé
                footer_background = "#516075"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#FFFFFF"
                header_text_color = "#516075"
                header_light_text_color = "#516075"
                accent_color = "#740606"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # Footer
                degradefooter1 = "#67A896"
                degradefooter2 = "#67A896"
                degradefooter3 = "#67A896"
                degradefooter4 = "#67A896"

            elif paleta == "Paleta de Cor 14":  # Verde Escuro e Bege
                # barra de anúncio
                cor1 = "#17453F"
                cor2 = "#17453F"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#E8E3D7"
                degrade2 = "#E8E3D7"
                degrade3 = "#E8E3D7"
                degrade4 = "#E8E3D7"

                # Cores gerais
                # Geral
                acent = "#17453F"
                link_color = "#17453F"
                # Botão primário
                primary_button_background = "#17453F"
                # Botão secundário
                secondary_button_background = "#17453F"

                # Rodapé
                footer_background = "#17453F"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#FFFFFF"
                header_text_color = "#17453F"
                header_light_text_color = "#17453F"
                accent_color = "#E8E3D7"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # Footer
                degradefooter1 = "#67A896"
                degradefooter2 = "#67A896"
                degradefooter3 = "#67A896"
                degradefooter4 = "#67A896"

            elif paleta == "Paleta de Cor 15":  # Verde e Amarelo
                # barra de anúncio
                cor1 = "#5EB64E"
                cor2 = "#5EB64E"
                textcolor = "#FFFFFF"

                # Cabeçalho/barra colorida
                degrade1 = "#E8E3D7"
                degrade2 = "#E8E3D7"
                degrade3 = "#E8E3D7"
                degrade4 = "#E8E3D7"

                # Cores gerais
                # Geral
                acent = "#5EB64E"
                link_color = "#5EB64E"
                # Botão primário
                primary_button_background = "#5EB64E"
                # Botão secundário
                secondary_button_background = "#5EB64E"

                # Rodapé
                footer_background = "#5EB64E"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#FFFFFF"
                header_text_color = "#5EB64E"
                header_light_text_color = "#5EB64E"
                accent_color = "#E5C92C"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#00D864"  # Cor do preço
                product_cor_do_preco = "#00D864"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # Footer
                degradefooter1 = "#E5C92C"
                degradefooter2 = "#E5C92C"
                degradefooter3 = "#E5C92C"
                degradefooter4 = "#E5C92C"

            elif paleta == "Paleta de Cor 16":  # PRETO E BRANCO MODA
                # barra de anúncio
                cor1 = "#000000"
                cor2 = "#000000"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#000000"
                degrade2 = "#000000"
                degrade3 = "#000000"
                degrade4 = "#000000"

                # Cores gerais
                # Geral
                acent = "#000000"
                link_color = "#000000"
                # Botão primário
                primary_button_background = "#000000"
                # Botão secundário
                secondary_button_background = "#000000"

                # Rodapé
                footer_background = "#F2F2F2"
                footer_text_color = "#000000"

                # Cabeçalho
                header_background = "#FFFFFF"
                header_text_color = "#000000"
                header_light_text_color = "#000000"
                accent_color = "#000000"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#000000"  # Cor do preço
                product_cor_do_preco = "#000000"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # Footer
                degradefooter1 = "#000000"
                degradefooter2 = "#000000"
                degradefooter3 = "#000000"
                degradefooter4 = "#000000"

            elif paleta == "Paleta de Cor 17":  #  Azul escuro e cinza
                # barra de anúncio
                cor1 = "#001438"
                cor2 = "#001438"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#9197AA"
                degrade2 = "#9197AA"
                degrade3 = "#9197AA"
                degrade4 = "#9197AA"

                # Cores gerais
                # Geral
                acent = "#001438"
                link_color = "#001438"
                # Botão primário
                primary_button_background = "#001438"
                # Botão secundário
                secondary_button_background = "#001438"

                # Rodapé
                footer_background = "#F2F2F2"
                footer_text_color = "#001438"

                # Cabeçalho
                header_background = "#FFFFFF"
                header_text_color = "#001438"
                header_light_text_color = "#001438"
                accent_color = "#9197AA"

                # Produtos
                product_on_sale_accent = "#001438"  # Barra economize
                product_cor_do_preco_semdesc = "#001438"  # Cor do preço
                product_cor_do_preco = "#001438"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # Footer
                degradefooter1 = "#9197AA"
                degradefooter2 = "#9197AA"
                degradefooter3 = "#9197AA"
                degradefooter4 = "#9197AA"

            elif paleta == "Paleta de Cor 18":  #  Bege e Marrom
                # barra de anúncio
                cor1 = "#9A5F43"
                cor2 = "#9A5F43"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#D1B1A3"
                degrade2 = "#D1B1A3"
                degrade3 = "#D1B1A3"
                degrade4 = "#D1B1A3"

                # Cores gerais
                # Geral
                acent = "#9A5F43"
                link_color = "#9A5F43"
                # Botão primário
                primary_button_background = "#9A5F43"
                # Botão secundário
                secondary_button_background = "#9A5F43"

                # Rodapé
                footer_background = "#F2F2F2"
                footer_text_color = "#9A5F43"

                # Cabeçalho
                header_background = "#FFFFFF"
                header_text_color = "#9A5F43"
                header_light_text_color = "#9A5F43"
                accent_color = "#D1B1A3"

                # Produtos
                product_on_sale_accent = "#9A5F43"  # Barra economize
                product_cor_do_preco_semdesc = "#9A5F43"  # Cor do preço
                product_cor_do_preco = "#9A5F43"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # Footer
                degradefooter1 = "#D1B1A3"
                degradefooter2 = "#D1B1A3"
                degradefooter3 = "#D1B1A3"
                degradefooter4 = "#D1B1A3"

            elif paleta == "Paleta de Cor 19":  # Cinza e Marrom
                # barra de anúncio
                cor1 = "#82746C"
                cor2 = "#82746C"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#82746C"
                degrade2 = "#82746C"
                degrade3 = "#82746C"
                degrade4 = "#82746C"

                # Cores gerais
                # Geral
                acent = "#82746C"
                link_color = "#82746C"
                # Botão primário
                primary_button_background = "#82746C"
                # Botão secundário
                secondary_button_background = "#82746C"
                # Rodapé
                footer_background = "#82746C"
                footer_text_color = "#FFFFFF"

                # Cabeçalho
                header_background = "#FFFFFF"
                header_text_color = "#5E5E5E"
                header_light_text_color = "#5E5E5E"
                accent_color = "#CFCAC4"

                # Produtos
                product_on_sale_accent = "#82746C"  # Barra economize
                product_cor_do_preco_semdesc = "#82746C"  # Cor do preço
                product_cor_do_preco = "#82746C"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # Footer
                degradefooter1 = "#CFCAC4"
                degradefooter2 = "#CFCAC4"
                degradefooter3 = "#CFCAC4"
                degradefooter4 = "#CFCAC4"

            elif paleta == "Paleta de Cor 20":  # ROSA E CINZA
                # barra de anúncio
                cor1 = "#5A5E61"
                cor2 = "#5A5E61"
                textcolor = "#FFFFFF"
                # Cabeçalho/barra colorida
                degrade1 = "#5A5E61"
                degrade2 = "#5A5E61"
                degrade3 = "#5A5E61"
                degrade4 = "#5A5E61"

                # Cores gerais
                # Geral
                acent = "#5A5E61"
                link_color = "#5A5E61"
                # Botão primário
                primary_button_background = "#5A5E61"
                # Botão secundário
                secondary_button_background = "#5A5E61"
                # Rodapé
                footer_background = "#F2F2F2"
                footer_text_color = "#5A5E61"

                # Cabeçalho
                header_background = "#FFFFFF"
                header_text_color = "#5A5E61"
                header_light_text_color = "#5A5E61"
                accent_color = "#CF97A0"

                # Produtos
                product_on_sale_accent = "#000000"  # Barra economize
                product_cor_do_preco_semdesc = "#000000"  # Cor do preço
                product_cor_do_preco = "#000000"  # Cor do preço com desconto
                product_cor_dos_titles = "#656866"  # Títulos das informações
                product_in_stock_color = "#00D864"  # Em estoque

                # Footer
                degradefooter1 = "#CF97A0"
                degradefooter2 = "#CF97A0"
                degradefooter3 = "#CF97A0"
                degradefooter4 = "#CF97A0"
            
            ##############################################################
            ##############################################################
            colecoes = []
            produtos = None
            ##########################                                                          
            ##########################
            if pais == "Brasil":
                if business_hours == "hora1":
                    antend = "Segunda a Sexta das 9:00 as 18:00"
                elif business_hours == "hora2": 
                    antend = "Segunda a Sexta das 9:00 as 22:00"
                elif business_hours == "hora3":
                    antend = "Segunda a Sábado das 9:00 as 18:00"
                elif business_hours == "hora4":
                    antend = "Segunda a Sábado das 9:00 as 22:00"
                elif business_hours == "hora5":
                    antend = "Segunda a Sexta 9:00 as 18:00, Sábado das 9:00 as 12:00"
                elif business_hours == "hora6":
                    antend = "Segunda a Sexta 9:00 as 22:00, Sábado das 9:00 as 12:00"
                elif business_hours == "hora7":
                    antend = "Todos os dias das 9:00 as 18:00"
                elif business_hours == "hora8":
                    antend = "Todos os dias das 9:00 as 22:00"
                if tema_base == "Dropmeta":
                    Tema = "https://github.com/ronildob/TEMAS2025/raw/refs/heads/main/tema_pt_br%2001-2025.zip"
                    footer = f"""<p>
                                    <strong>E-mail:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a>
                                </p>
                                <p>
                                    <strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{cod_pais[pais]} {whatsapp}</a>
                                </p>
                                <p>
                                    <strong>Hor\u00e1rio de atendimento: </strong>{antend}
                                </p>"""
                    titlezap = f"Fale Conosco \n<span>\nO melhor suporte é com a {empresa}!\n</span>"
                elif tema_base == "Evolution":
                    Tema = "https://github.com/ronildob/Themes/raw/refs/heads/main/TEMA_GL_NOVO.zip"
                    footer = f"""<p></p>
                                        <p>
                                            <strong>SERVIÇO DE SAC DISPONÍVEL</strong>
                                        </p>
                                        <p>
                                            <strong>E-mail:</strong> <a href="mailto:{email_suporte}">{email_suporte}</a>
                                        </p>
                                        <p>
                                            <strong>Atendimento:</strong><br/>\u00a0<a href="https://wa.me/{whatsapp_completo}">{cod_pais[pais]} {whatsapp}</a>
                                        </p>
                                <p></p>"""
                paginas = {
                        "Rastrear Pedido": "static/home/politicas/Rastrearpedido.txt",
                        "Contrato de E-Commerce": "static/home/politicas/Contrato de E-Commerce.docx",
                        "Termos de Uso": "static/home/politicas/Termos de Uso.docx",
                        "Prazo de Entrega": "static/home/politicas/Prazo de Entrega.docx",
                        "Política de Privacidade": "static/home/politicas/Política de Privacidade.docx",
                        "Trocas ou Devoluções": "static/home/politicas/Trocas ou Devolução.docx",
                        "Sobre Nós": "static/home/politicas/Sobre Nós.docx",
                        "Pagamento Seguro": "static/home/politicas/Pagamento  Seguro.docx",
                    }                                    
                politicas_para_atualizar = {
                    "Legal notice": "static/home/politicas/Contrato de E-Commerce.docx",
                    "Terms of service": "static/home/politicas/Termos de Uso.docx",
                    "Shipping policy": "static/home/politicas/Prazo de Entrega.docx",
                    "Refund policy": "static/home/politicas/Trocas ou Devolução.docx",
                    #"Privacy policy": "static/home/politicas/Política de Privacidade.docx",
                }
                                                    
                if estilo == "Genérica":   
                    df_path = "static/home/csv/PT_BR/30GENERICOS.csv"
                    pri = "Casa"
                    seg = "Eletrônicos"
                    ter = "Pets"
                    qua = "Fitness"
                    qui = "Kids"
                    sex = "Mais Vendidos"

                    primeiro = "Casa"
                    segundo = "Eletrônicos"
                    terceiro = "Pets"
                    quarto = "Fitness"
                    quinto = "Kids"
                    sexto = "Mais-Vendidos"
                    colecoes = [
                        "Casa", 
                        "Eletrônicos", 
                        "Pets", 
                        "Fitness", 
                        "Kids", 
                        "Mais Vendidos"
                        ]
                    
                    # urls = {
                    #     "desktop1": "",
                    #     "desktop2": "",
                    #     "desktop3": "",
                    #     "mobile1": "",
                    #     "mobile2": "",
                    #     "mobile3": "",
                    #     "capa1": "",
                    #     "capa2": "",
                    #     "capa3": "",
                    #     "capa4": "",
                    #     "capa5": "",
                    #     "capa6": "",
                    # }

                elif estilo == "Eletrônicos":
                    df_path = "static/home/csv/PT_BR/30ELETRONICOS.csv"
                    pri = "Fones"
                    seg = "Drones"
                    ter = "Smartwatch"
                    qua = "Acessórios"
                    qui = "Produtos Gamer"
                    sex = "Mais Vendidos"

                    primeiro = "Fones"
                    segundo = "Drones"
                    terceiro = "Smartwatch"
                    quarto = "Acessórios"
                    quinto = "Produtos-Gamer"
                    sexto = "Mais-Vendidos"
                    colecoes = ["Fones", "Drones", "Smartwatch", "Acessórios", "Produtos Gamer", "Mais Vendidos"]

                    urls = {
                        "desktop1": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736255612/banner1_rmpi15.svg",
                        "desktop2": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736255719/banner2_1_f4tkw0.svg",
                        "desktop3": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736255617/banner3_f0aate.svg",
                        "mobile1": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275595/mobile1_fwcsuu.svg",
                        "mobile2": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275593/mobile2_uc1t6y.svg",
                        "mobile3": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275595/mobile1_fwcsuu.svg",
                        "capa1": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275591/capa1_vseevh.svg",
                        "capa2": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275590/capa2_jws3hn.svg",
                        "capa3": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275591/capa3_jryrkt.svg",
                        "capa4": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275591/capa4_pqm3qy.svg",
                        "capa5": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275590/capa5_mj8wbe.svg",
                        "capa6": "https://res.cloudinary.com/hnrqudfke/image/upload/v1736275590/capa6_qpbjav.svg",

                        }

                elif estilo == "Kids":
                    df_path = "static/home/csv/PT_BR/30KIDS.csv"
                    pri = "Moda"
                    seg = "Acessórios"
                    ter = "Brinquedos"
                    qua = "Cuidados"
                    qui = "Maternidade"
                    sex = "Mais Vendidos"

                    primeiro = "Moda"
                    segundo = "Acessórios"
                    terceiro = "Brinquedos"
                    quarto = "Cuidados"
                    quinto = "Maternidade"
                    sexto = "Mais-Vendidos"
                    colecoes = [
                        "Moda", 
                        "Acessórios", 
                        "Brinquedos", 
                        "Cuidados", 
                        "Maternidade", 
                        "Mais Vendidos"
                        ]
                    # urls = {
                    #     "desktop1": "",
                    #     "desktop2": "",
                    #     "desktop3": "",
                    #     "mobile1": "",
                    #     "mobile2": "",
                    #     "mobile3": "",
                    #     "capa1": "",
                    #     "capa2": "",
                    #     "capa3": "",
                    #     "capa4": "",
                    #     "capa5": "",
                    #     "capa6": "",
                    # }
                
                elif estilo == "Casa":
                    df_path = "static/home/csv/PORTUGUES_PT/CASA_PORTUGAL.csv"

                    pri = "Acessórios"
                    seg = "Iluminação"
                    ter = "Decoração de Ambientes"
                    qua = "Organização e Armazenamento"
                    qui = "Aromaterapia e Bem-estar"
                    sex = "Produtos Sazonais"

                    primeiro = "Acessórios"
                    segundo = "Iluminação"
                    terceiro = "Decoração-de-Ambientes"
                    quarto = "Organização-e-Armazenamento"
                    quinto = "Aromaterapia-e-Bem-estar"
                    sexto = "Produtos-Sazonais"
                    colecoes = [
                        "Acessórios",
                        "Iluminação",
                        "Decoração de Ambientes",
                        "Organização e Armazenamento",
                        "Aromaterapia e Bem-estar",
                        "Produtos Sazonais",
                    ]                    
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                                        
                elif estilo == "Pet":
                    df_path = "static/home/csv/PT_BR/30PETS.csv" 

                    pri = "Camas"
                    seg = "Acessórios"
                    ter = "Comedouros"
                    qua = "Brinquedos"
                    qui = "Roupas"
                    sex = "Mais Vendidos"

                    primeiro = "Camas"
                    segundo = "Acessórios"
                    terceiro = "Comedouros"
                    quarto = "Brinquedos"
                    quinto = "Roupas"
                    sexto = "Mais-Vendidos"
                    colecoes = [
                        "Camas", 
                        "Acessórios", 
                        "Comedouros", 
                        "Brinquedos", 
                        "Roupas", 
                        "Mais Vendidos"
                        ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Fitness":
                    df_path = "static/home/csv/PT_BR/30FITNESS.csv"
                    pri = "Moda"
                    seg = "Acessórios"
                    ter = "Calçados"
                    qua = "Corretores"
                    qui = "Produtos Esportivos"
                    sex = "Mais Vendidos"

                    primeiro = "Moda"
                    segundo = "Acessórios"
                    terceiro = "Calçados"
                    quarto = "Corretores"
                    quinto = "Produtos-Esportivos"
                    sexto = "Mais-Vendidos"
                    colecoes = [
                        "Moda", 
                        "Acessórios", 
                        "Calçados", 
                        "Corretores", 
                        "Produtos Esportivos", 
                        "Mais Vendidos"
                        ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Masculino":
                    df_path = "static/home/csv/PT_BR/30MASCULINO.csv"

                    pri = "Acessórios"
                    seg = "Ferramentas"
                    ter = "Pesca"
                    qua = "Vestuário"
                    qui = "Automotivos"
                    sex = "Mais Vendidos"

                    primeiro = "Acessórios"
                    segundo = "Ferramentas"
                    terceiro = "Pesca"
                    quarto = "Vestuário"
                    quinto = "Automotivos"
                    sexto = "Mais-Vendidos"
                    colecoes = [
                        "Acessórios", 
                        "Ferramentas", 
                        "Pesca", 
                        "Vestuário", 
                        "Automotivos", 
                        "Mais Vendidos"
                        ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                        
                elif estilo == "Feminino":
                    df_path = "static/home/csv/PT_BR/30FEMININO.csv"

                    pri = "Moda"
                    seg = "Acessórios"
                    ter = "Make"
                    qua = "Escovas Alisadoras"
                    qui = "Saúde"
                    sex = "Mais Vendidos"

                    primeiro = "Moda"
                    segundo = "Acessórios"
                    terceiro = "Make"
                    quarto = "Escovas-Alisadoras"
                    quinto = "Saúde"
                    sexto = "Mais-Vendidos"
                    colecoes = [
                        "Moda", 
                        "Acessórios", 
                        "Make", 
                        "Escovas Alisadoras", 
                        "Saúde", 
                        "Mais Vendidos"
                        ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                                        
                elif estilo == "Moda":
                    df_path = "static/home/csv/PT_BR/MODA.csv"
                    pri = "Acessórios"
                    seg = "Feminino"
                    ter = "Masculino"
                    qua = "Infantil"
                    qui = "Jóias"
                    sex = "Promoções"

                    primeiro = "Acessórios"
                    segundo = "Feminino"
                    terceiro = "Masculino"
                    quarto = "Infantil"
                    quinto = "Jóias"
                    sexto = "Promoções"
                    colecoes = [
                        "Acessórios",
                        "Feminino",
                        "Masculino",
                        "Infantil",
                        "Jóias",
                        "Promoções",
                    ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                    
                elif estilo == "Saúde e Beleza":
                    df_path = "static/home/csv/PT_BR/SAUDEEBELEZA.csv"
                    pri = "Fitness e Exercício"
                    seg = "Beleza"
                    ter = "Mais Vendidos"
                    qua = "Relaxamento"
                    qui = "Cuidados Pessoais"
                    sex = "Monitoramento da Saúde"

                    primeiro = "Fitness-e-Exercício"
                    segundo = "Beleza"
                    terceiro = "Mais-Vendidos"
                    quarto = "Relaxamento"
                    quinto = "Cuidados-Pessoais"
                    sexto = "Monitoramento-da-Saúde"
                    colecoes = [
                        "Fitness e Exercício",
                        "Beleza",
                        "Mais Vendidos",
                        "Relaxamento",
                        "Cuidados Pessoais",
                        "Monitoramento da Saúde",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

            ##########################                                                          
            ########################## 
            elif pais == "Portugal":
                if business_hours == "hora1":
                    antend = "Segunda a Sexta das 9:00 as 18:00"
                elif business_hours == "hora2":
                    antend = "Segunda a Sexta das 9:00 as 22:00"
                elif business_hours == "hora3":
                    antend = "Segunda a Sábado das 9:00 as 18:00"
                elif business_hours == "hora4":
                    antend = "Segunda a Sábado das 9:00 as 22:00"
                elif business_hours == "hora5":
                    antend = (
                        "Segunda a Sexta 9:00 as 18:00, Sábado das 9:00 as 12:00"
                    )
                elif business_hours == "hora6":
                    antend = (
                        "Segunda a Sexta 9:00 as 22:00, Sábado das 9:00 as 12:00"
                    )
                elif business_hours == "hora7":
                    antend = "Todos os dias das 9:00 as 18:00"
                elif business_hours == "hora8":
                    antend = "Todos os dias das 9:00 as 22:00"
                ##############################
                footer = f"""<p>
                                <strong>E-mail:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a>
                            </p>
                            <p>
                                <strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{cod_pais[pais]} {whatsapp}</a>
                            </p>
                            <p>
                                <strong>Hor\u00e1rio de atendimento: </strong>{antend}
                            </p>"""
                titlezap = f"Fale Conosco \n<span>\nO melhor suporte é com a {empresa}!\n</span>"
                ##############################
                Tema = "https://github.com/ronildob/temas/raw/refs/heads/main/portuguesPToficial.zip"
                paginas = {
                    "Rastrear Pedido": "static/home/politicas/Rastrearpedido.txt",
                    "Contrato de E-Commerce": "static/home/politicas/Contrato de E-Commerce.docx",
                    "Termos de Uso": "static/home/politicas/Termos de Uso.docx",
                    "Prazo de Entrega": "static/home/politicas/Prazo de Entrega.docx",
                    "Política de Privacidade": "static/home/politicas/Política de Privacidade.docx",
                    "Trocas ou Devoluções": "static/home/politicas/Trocas ou Devolução.docx",
                    "Sobre Nós": "static/home/politicas/Sobre Nós.docx",
                    "Pagamento Seguro": "static/home/politicas/Pagamento  Seguro.docx",
                }
                politicas_para_atualizar = {
                    "Legal notice": "static/home/politicas/Contrato de E-Commerce.docx",
                    "Terms of service": "static/home/politicas/Termos de Uso.docx",
                    "Shipping policy": "static/home/politicas/Prazo de Entrega.docx",
                    # "Privacy policy": "static/home/politicas/Política de Privacidade.docx",
                    "Refund policy": "static/home/politicas/Trocas ou Devolução.docx",
                }
                ##############################

                if estilo == "Genérica":
                    df_path = "static/home/csv/PORTUGUES_PT/30GENERICOSPT.csv"

                    pri = "Casa"
                    seg = "Eletrónica"
                    ter = "Pets"
                    qua = "Fitness"
                    qui = "Kids"
                    sex = "Más Vendidos"

                    primeiro = "Casa"
                    segundo = "Eletrónica"
                    terceiro = "Pets"
                    quarto = "Fitness"
                    quinto = "Kids"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Casa",
                        "Eletrónica",
                        "Pets",
                        "Fitness",
                        "Kids",
                        "Más Vendidos",
                    ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                        
                elif estilo == "Eletrônicos":
                    df_path = "static/home/csv/PORTUGUES_PT/ELETRONICOPT.csv"

                    pri = "Fones"
                    seg = "Drones"
                    ter = "Smartwatch"
                    qua = "Acessórios"
                    qui = "Produtos Gamer"
                    sex = "Más Vendidos"

                    primeiro = "Fones"
                    segundo = "Drones"
                    terceiro = "Smartwatch"
                    quarto = "Acessórios"
                    quinto = "Produtos-Gamer"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Fones",
                        "Drones",
                        "Smartwatch",
                        "Acessórios",
                        "Produtos Gamer",
                        "Más Vendidos",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Kids":
                    df_path = "static/home/csv/PORTUGUES_PT/KIDSPORTUGAL.csv"

                    pri = "Moda"
                    seg = "Acessórios"
                    ter = "Brinquedos"
                    qua = "Cuidados"
                    qui = "Maternidader"
                    sex = "Más Vendidos"

                    primeiro = "Moda"
                    segundo = "Acessórios"
                    terceiro = "Brinquedos"
                    quarto = "Cuidados"
                    quinto = "Maternidade"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Moda",
                        "Acessórios",
                        "Brinquedos",
                        "Cuidados",
                        "Maternidade",
                        "Más Vendidos",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Casa":
                    df_path = "static/home/csv/PORTUGUES_PT/CASA_PORTUGAL.csv"
                    
                    pri = "Acessórios"
                    seg = "Iluminação"
                    ter = "Decoração de Ambientes"
                    qua = "Organização e Armazenamento"
                    qui = "Aromaterapia e Bem-estar"
                    sex = "Produtos Sazonais"

                    primeiro = "Acessórios"
                    segundo = "Iluminação"
                    terceiro = "Decoração-de-Ambientes"
                    quarto = "Organização-e-Armazenamento"
                    quinto = "Aromaterapia-e-Bem-estar"
                    sexto = "Produtos-Sazonais"
                    colecoes = [
                        "Acessórios",
                        "Iluminação",
                        "Decoração de Ambientes",
                        "Organização e Armazenamento",
                        "Aromaterapia e Bem-estar",
                        "Produtos Sazonais",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Pet":
                    df_path = "static/home/csv/PORTUGUES_PT/PETPORTUGAL.csv"

                    pri = "Camas"
                    seg = "Acessórios"
                    ter = "Comedouros"
                    qua = "Brinquedos"
                    qui = "Roupas"
                    sex = "Más Vendidos"

                    primeiro = "Camas"
                    segundo = "Acessórios"
                    terceiro = "Comedouros"
                    quarto = "Brinquedos"
                    quinto = "Roupas"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Camas",
                        "Acessórios",
                        "Comedouros",
                        "Brinquedos",
                        "Roupas",
                        "Más Vendidos",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Fitness":
                    df_path = "static/home/csv/PORTUGUES_PT/FITNESSPORTUGAL.csv"

                    pri = "Moda"
                    seg = "Acessórios"
                    ter = "Calçados"
                    qua = "Corretores"
                    qui = "Produtos Esportivos"
                    sex = "Más Vendidos"

                    primeiro = "Moda"
                    segundo = "Acessórios"
                    terceiro = "Calçados"
                    quarto = "Corretores"
                    quinto = "Produtos-Esportivos"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Moda",
                        "Acessórios",
                        "Calçados",
                        "Corretores",
                        "Produtos Esportivos",
                        "Más Vendidos",
                    ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                
                elif estilo == "Masculino":
                    df_path = "static/home/csv/PT_BR/30MASCULINO.csv"

                    pri = "Acessórios"
                    seg = "Ferramentas"
                    ter = "Pesca"
                    qua = "Vestuário"
                    qui = "Automotivos"
                    sex = "Más Vendidos"

                    primeiro = "Acessórios"
                    segundo = "Ferramentas"
                    terceiro = "Pesca"
                    quarto = "Vestuário"
                    quinto = "Automotivos"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Acessórios", 
                        "Ferramentas", 
                        "Pesca", 
                        "Vestuário", 
                        "Automotivos", 
                        "Más Vendidos"
                        ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                        
                elif estilo == "Feminino":
                    df_path = "static/home/csv/PT_BR/30FEMININO.csv"

                    pri = "Moda"
                    seg = "Acessórios"
                    ter = "Make"
                    qua = "Escovas Alisadoras"
                    qui = "Saúde"
                    sex = "Más Vendidos"

                    primeiro = "Moda"
                    segundo = "Acessórios"
                    terceiro = "Make"
                    quarto = "Escovas-Alisadoras"
                    quinto = "Saúde"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Moda", 
                        "Acessórios", 
                        "Make", 
                        "Escovas Alisadoras", 
                        "Saúde", 
                        "Más Vendidos"
                        ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                            
                elif estilo == "Moda":
                    df_path = "static/home/csv/PT_BR/MODA.csv"
                    pri = "Acessórios"
                    seg = "Feminino"
                    ter = "Masculino"
                    qua = "Infantil"
                    qui = "Jóias"
                    sex = "Promoções"

                    primeiro = "Acessórios"
                    segundo = "Feminino"
                    terceiro = "Masculino"
                    quarto = "Infantil"
                    quinto = "Jóias"
                    sexto = "Promoções"
                    colecoes = [
                        "Acessórios",
                        "Feminino",
                        "Masculino",
                        "Infantil",
                        "Jóias",
                        "Promoções",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Saúde e Beleza":
                    df_path = "static/home/csv/PORTUGUES_PT/SAUDE_BEM_ESTAR_PT.csv"
                    pri = "Fitness e Exercício"
                    seg = "Beleza"
                    ter = "Alimentação"
                    qua = "Relaxamento"
                    qui = "Cuidados Pessoais"
                    sex = "Monitoramento da Saúde"

                    primeiro = "Fitness-e-Exercício"
                    segundo = "Beleza"
                    terceiro = "Alimentação"
                    quarto = "Relaxamento"
                    quinto = "Cuidados-Pessoais"
                    sexto = "Monitoramento-da-Saúde"
                    colecoes = [
                        "Fitness e Exercício",
                        "Beleza",
                        "Alimentação",
                        "Relaxamento",
                        "Cuidados Pessoais",
                        "Monitoramento da Saúde",
                    ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
            
            ##########################
            ########################## 
            elif pais == "EUA":
                
                Tema = "https://github.com/ronildob/temas/raw/refs/heads/main/inglesoficial.zip"
                paginas = {
                    "Track Order": "static/home/politicas/Track Order.txt",
                    "E-Commerce Agreement": "static/home/politicas/E-Commerce Agreement.docx",
                    "Terms of Use": "static/home/politicas/Terms of Use.docx",
                    "Deadline": "static/home/politicas/Deadline.docx",
                    "Privacy Policy": "static/home/politicas/Privacy Policy.docx",
                    "Exchanges or Return": "static/home/politicas/Exchanges or Return.docx",
                    "About Us": "static/home/politicas/About Us.docx",
                    "Secure Payment": "static/home/politicas/Secure Payment.docx",
                }
                politicas_para_atualizar = {
                    "Legal notice": "static/home/politicas/E-Commerce Agreement.docx",
                    "Terms of service": "static/home/politicas/Terms of Use.docx",
                    "Shipping policy": "static/home/politicas/Deadline.docx",
                    # "Privacy policy": "static/home/politicas/Privacy Policy.docx",
                    "Refund policy": "static/home/politicas/Exchanges or Return.docx",
                }
                ##############################
                
                if business_hours == "hora1":
                    antend = "Monday to Friday from 9:00 AM to 6:00 PM"
                elif business_hours == "hora2":
                    antend = "Monday to Friday from 9:00 AM to 10:00 PM"
                elif business_hours == "hora3":
                    antend = "Monday to Saturday from 9:00 AM to 6:00 PM"
                elif business_hours == "hora4":
                    antend = "Monday to Saturday from 9:00 AM to 10:00 PM"
                elif business_hours == "hora5":
                    antend = "Monday to Friday from 9:00 AM to 6:00 PM - Saturday from 9:00 AM to 12:00 PM"
                elif business_hours == "hora6":
                    antend = "Monday to Friday from 9:00 AM to 10:00 PM - Saturday from 9:00 AM to 12:00 PM"
                elif business_hours == "hora7":
                    antend = "Every day from 9:00 AM to 6:00 PM"
                elif business_hours == "hora8":
                    antend = "Every day from 9:00 AM to 10:00 PM"
                ##############################
                footer = f"""<p><strong>Email:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a></p><p><strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{
                    cod_pais[pais]} {whatsapp}</a></p><p><strong>Business Hours: </strong>{antend}</p>"""
                titlezap = f"Contact Us \n<span>\nThe best support is with {empresa}!\n</span>"
                ##############################
                if estilo == "Genérica":
                    df_path = "static/home/csv/INGLES/30GEN_EUA.csv"

                    pri = "House"
                    seg = "Electronics"
                    ter = "Pets"
                    qua = "Fitness"
                    qui = "Kids"
                    sex = "Best Sellers"
                
                    primeiro = "House"
                    segundo = "Electronics"
                    terceiro = "Pets"
                    quarto = "Fitness"
                    quinto = "Kids"
                    sexto = "Best-Sellers"
                    colecoes = [
                        "House",
                        "Electronics",
                        "Pets",
                        "Fitness",
                        "Kids",
                        "Best Sellers",
                    ]
                    ##############################

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Eletrônicos":
                    df_path = "static/home/csv/INGLES/ELETRONICOPT.csv"

                    pri = "Headphones"
                    seg = "Drones"
                    ter = "Smartwatch"
                    qua = "Accessories"
                    qui = "Gamer Products"
                    sex = "Best Sellers"

                    primeiro = "Headphones"
                    segundo = "Drones"
                    terceiro = "Smartwatch"
                    quarto = "Accessories"
                    quinto = "Gamer-Products"
                    sexto = "Best-Sellers"
                    colecoes = [
                        "Headphones",
                        "Drones",
                        "Smartwatch",
                        "Accessories",
                        "Gamer Products",
                        "Best Sellers",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Kids":
                    df_path = "static/home/csv/INGLES/KIDSPORTUGAL.csv"

                    pri = "Fashion"
                    seg = "Accessories"
                    ter = "Toys"
                    qua = "Care"
                    qui = "Maternity"
                    sex = "Best Sellers"

                    primeiro = "Fashion"
                    segundo = "Accessories"
                    terceiro = "Toys"
                    quarto = "Care"
                    quinto = "Maternity"
                    sexto = "Best-Sellers"
                    colecoes = [
                        "Fashion",
                        "Accessories",
                        "Toys",
                        "Care",
                        "Maternity",
                        "Best Sellers",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Casa":
                    df_path = "static/home/csv/INGLES/CASA_PORTUGAL.csv"
                    
                    pri = "Accessories"
                    seg = "Lighting"
                    ter = "Environment Decoration"
                    qua = "Organization and Storage"
                    qui = "Aromatherapy and Wellbeing"
                    sex = "Seasonal Products"

                    primeiro = "Accessories"
                    segundo = "Lighting"
                    terceiro = "Environment-Decoration"
                    quarto = "Organization-and-Storage"
                    quinto = "Aromatherapy-and-Wellbeing"
                    sexto = "Seasonal-Products"
                    colecoes = [
                        "Accessories",
                        "Lighting",
                        "Environment Decoration",
                        "Organization and Storage",
                        "Aromatherapy and Wellbeing",
                        "Seasonal Products",
                    ]
                    
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Pet":
                    df_path = "static/home/csv/INGLES/PETPORTUGAL.csv"

                    pri = "Beds"
                    seg = "Accessories"
                    ter = "Feeders"
                    qua = "Toys"
                    qui = "Clothes"
                    sex = "Best Sellers"

                    primeiro = "Beds"
                    segundo = "Accessories"
                    terceiro = "Feeders"
                    quarto = "Toys"
                    quinto = "Clothes"
                    sexto = "Best-Sellers"
                    colecoes = [
                        "Beds",
                        "Accessories",
                        "Feeders",
                        "Toys",
                        "Clothes",
                        "Best Sellers",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Fitness":
                    df_path = "static/home/csv/INGLES/FITNESSPORTUGAL.csv"

                    pri = "Fashion"
                    seg = "Accessories"
                    ter = "Shoes"
                    qua = "Brokers"
                    qui = "Sports Products"
                    sex = "Best Sellers"

                    primeiro = "Fashion"
                    segundo = "Accessories"
                    terceiro = "Shoes"
                    quarto = "Brokers"
                    quinto = "Sports-Products"
                    sexto = "Best-Sellers"
                    colecoes = [
                        "Fashion",
                        "Accessories",
                        "Shoes",
                        "Brokers",
                        "Sports Products",
                        "Best Sellers",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Masculino":
                    df_path = "static/home/csv/PT_BR/30MASCULINO.csv"

                    pri = "Accessories"
                    seg = "Tools"
                    ter = "Fishing"
                    qua = "Clothing"
                    qui = "Automotive"
                    sex = "Best Sellers"

                    primeiro = "Accessories"
                    segundo = "Tools"
                    terceiro = "Fishing"
                    quarto = "Clothing"
                    quinto = "Automotive"
                    sexto = "Best-Sellers"
                    colecoes = [
                        "Accessories", 
                        "Tools", 
                        "Fishing", 
                        "Clothing", 
                        "Automotive", 
                        "Best Sellers"
                        ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                        
                elif estilo == "Feminino":
                    df_path = "static/home/csv/PT_BR/30FEMININO.csv"

                    pri = "Fashion"
                    seg = "Accessories"
                    ter = "Makeup"
                    qua = "Straightening Brushes"
                    qui = "Health"
                    sex = "Best Sellers"

                    primeiro = "Fashion"
                    segundo = "Accessories"
                    terceiro = "Makeup"
                    quarto = "Straightening-Brushes"
                    quinto = "Health"
                    sexto = "Best-Sellers"
                    colecoes = [
                        "Fashion", 
                        "Accessories", 
                        "Makeup", 
                        "Straightening Brushes", 
                        "Health", 
                        "Best Sellers"
                        ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                            
                elif estilo == "Moda":
                    df_path = "static/home/csv/INGLES/MODAPORTUGAL.csv"
                    pri = "Accessories"
                    seg = "Feminine"
                    ter = "Masculine"
                    qua = "Children"
                    qui = "Jewelry"
                    sex = "Promotions"

                    primeiro = "Accessories"
                    segundo = "Feminine"
                    terceiro = "Masculine"
                    quarto = "Children"
                    quinto = "Jewelry"
                    sexto = "Promotions"
                    colecoes = [
                        "Accessories",
                        "Feminine",
                        "Masculine",
                        "Children",
                        "Jewelry",
                        "Promotions",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                
                elif estilo == "Saúde e Beleza":
                    df_path = "static/home/csv/INGLES/SAUDEBLEZPORT.csv"
                    pri = "Fitness and Exercise"
                    seg = "Beauty"
                    ter = "Best Sellers"
                    qua = "Relaxation"
                    qui = "Personal Care"
                    sex = "Health Monitoring"

                    primeiro = "Fitness-and-Exercise"
                    segundo = "Beauty"
                    terceiro = "Best-Sellers"
                    quarto = "Relaxation"
                    quinto = "Personal-Care"
                    sexto = "Health-Monitoring"
                    colecoes = [
                        "Fitness and Exercise",
                        "Beauty",
                        "Best Sellers",
                        "Relaxation",
                        "Personal Care",
                        "Health Monitoring",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

            ##########################
            ########################## 
            elif pais == "Espanha":
                
                Tema = "https://github.com/ronildob/temas/raw/refs/heads/main/espanholoficial.zip"
                paginas = {
                    "Rastrear Pedidos": "static/home/politicas/seguimiento-del-pedido.txt",
                    "Contrato de Comercio Electrónico": "static/home/politicas/Contrato de Comercio Electrónico.docx",
                    "Términos de Uso": "static/home/politicas/Términos de Uso.docx",
                    "Plazo de Entrega": "static/home/politicas/Plazo de Entrega.docx",
                    "Política de Privacidad": "static/home/politicas/Política de Privacidad.docx",
                    "Cambios o Devoluciones": "static/home/politicas/Cambios o Devoluciones.docx",
                    "Sobre Nosotros": "static/home/politicas/Sobre Nosotros.docx",
                    "Pago Seguro": "static/home/politicas/Pago Seguro.docx",
                }
                politicas_para_atualizar = {
                    "Legal notice": "static/home/politicas/Contrato de Comercio Electrónico.docx",
                    "Terms of service": "static/home/politicas/Términos de Uso.docx",
                    "Shipping policy": "static/home/politicas/Plazo de Entrega.docx",
                    # "Privacy policy": "static/home/politicas/Política de Privacidad.docx",
                    "Refund policy": "static/home/politicas/Cambios o Devoluciones.docx",
                }
                ##############################
                
                if business_hours == "hora1":
                    antend = "Lunes a viernes de 9:00 a 18:00"
                elif business_hours == "hora2":
                    antend = "Lunes a viernes de 9:00 a 22:00"
                elif business_hours == "hora3":
                    antend = "Lunes a Sábado de  9:00 a 18"
                elif business_hours == "hora4":
                    antend = "Lunes a Sábado de  9:00 a 22"
                elif business_hours == "hora5":
                    antend = "Lunes a Viernes 9:00 a 18:00, Sábado 9:00 a 12:00"
                elif business_hours == "hora6":
                    antend = (
                        "Lunes a viernes de 9:00 a 22:00 - Sábado de 9:00 a 12:00"
                    )
                elif business_hours == "hora7":
                    antend = "Todos los días de 9:00 a 18:00"
                elif business_hours == "hora8":
                    antend = "Todos los días de 9:00 a 22:00"
                ##############################
                footer = f"""<p><strong>Correo electrónico:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a></p><p><strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{
                    cod_pais[pais]} {whatsapp}</a></p><p><strong>Horario de atención: </strong>{antend}</p>"""
                titlezap = f"Contáctanos \n<span>\n¡El mejor soporte es con {empresa}!\n</span>"
                ##############################
                if estilo == "Genérica":
                    df_path = "static/home/csv/ESPANHOL/30GENER_ESPANHOL.csv"

                    pri = "Casa"
                    seg = "Electrónica"
                    ter = "Pets"
                    qua = "Fitness"
                    qui = "Kids"
                    sex = "Más Vendidos"

                    primeiro = "Casa"
                    segundo = "Electrónica"
                    terceiro = "Pets"
                    quarto = "Fitness"
                    quinto = "Kids"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Casa",
                        "Electrónica",
                        "Pets",
                        "Fitness",
                        "Kids",
                        "Más Vendidos",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                    
                elif estilo == "Eletrônicos":
                    df_path = "static/home/csv/ESPANHOL/30ELETRONICOS.csv"
                    pri = "Auriculares"
                    seg = "Drones"
                    ter = "Smartwatch"
                    qua = "Accesorios"
                    qui = "Productos para Jugadores"
                    sex = "Más Vendidos"

                    primeiro = "Auriculares"
                    segundo = "Drones"
                    terceiro = "Smartwatch"
                    quarto = "Accesorios"
                    quinto = "Productos-para-Jugadores"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Auriculares", 
                        "Drones", 
                        "Smartwatch", 
                        "Accesorios", 
                        "Productos para Jugadores", 
                        "Más Vendidos"
                        ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                    
                elif estilo == "Kids":
                    df_path = "static/home/csv/ESPANHOL/30KIDS.csv"
                    pri = "Moda"
                    seg = "Accesorios"
                    ter = "Juguetes"
                    qua = "Cuidado"
                    qui = "Maternidad"
                    sex = "Más Vendidos"

                    primeiro = "Moda"
                    segundo = "Accesorios"
                    terceiro = "Juguetes"
                    quarto = "Cuidado"
                    quinto = "Maternidad"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Moda", 
                        "Accesorios", 
                        "Juguetes", 
                        "Cuidado", 
                        "Maternidad", 
                        "Más Vendidos"]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Casa":
                    df_path = "static/home/csv/ESPANHOL/"
                    
                    pri = "Accesorios"
                    seg = "Iluminación"
                    ter = "Decoración del ambiente"
                    qua = "Organización y almacenamiento"
                    qui = "Aromaterapia y Bienestar"
                    sex = "Productos de Temporada"

                    primeiro = "Accesorios"
                    segundo = "Iluminación"
                    terceiro = "Decoración-del-ambiente"
                    quarto = "Organización-y-almacenamiento"
                    quinto = "Aromaterapia-y-Bienestar"
                    sexto = "Productos-de-Temporada"
                    colecoes = [
                        "Accesorios",
                        "Iluminación",
                        "Decoración del ambiente",
                        "Organización y almacenamiento",
                        "Aromaterapia y Bienestar",
                        "Productos de Temporada",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Pet":
                    df_path = "static/home/csv/ESPANHOL/PETPORTUGAL.csv"

                    pri = "Camas"
                    seg = "Accesorios"
                    ter = "Comederos"
                    qua = "Juguetes"
                    qui = "Ropa"
                    sex = "Más Vendidos"

                    primeiro = "Camas"
                    segundo = "Accesorios"
                    terceiro = "Comederos"
                    quarto = "Juguetes"
                    quinto = "Ropa"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Camas",
                        "Accesorios",
                        "Comederos",
                        "Juguetes",
                        "Ropa",
                        "Más Vendidos",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Fitness":
                    df_path = "static/home/csv/ESPANHOL/FITNESSPORTUGAL.csv"

                    pri = "Moda"
                    seg = "Accesorios"
                    ter = "Zapatos"
                    qua = "Correctores"
                    qui = "Productos Deportivos"
                    sex = "Más Vendidos"

                    primeiro = "Moda"
                    segundo = "Accesorios"
                    terceiro = "Zapatos"
                    quarto = "Correctores"
                    quinto = "Productos-Deportivos"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Moda",
                        "Accesorios",
                        "Zapatos",
                        "Correctores",
                        "Productos Deportivos",
                        "Más Vendidos",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Masculino":
                    df_path = "static/home/csv/PT_BR/30MASCULINO.csv"

                    pri = "Accesorios"
                    seg = "Herramientas"
                    ter = "Pesca"
                    qua = "Ropa"
                    qui = "Automotor"
                    sex = "Más Vendidos"

                    primeiro = "Accesorios"
                    segundo = "Herramientas"
                    terceiro = "Pesca"
                    quarto = "Ropa"
                    quinto = "Automotor"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Accesorios", 
                        "Herramientas", 
                        "Pesca", 
                        "Ropa", 
                        "Automotor", 
                        "Más Vendidos"
                        ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                        
                elif estilo == "Feminino":
                    df_path = "static/home/csv/PT_BR/30FEMININO.csv"

                    pri = "Moda"
                    seg = "Accesorios"
                    ter = "Make"
                    qua = "Cepillos Alisadores"
                    qui = "Salud"
                    sex = "Más Vendidos"

                    primeiro = "Moda"
                    segundo = "Accesorios"
                    terceiro = "Make"
                    quarto = "Cepillos-Alisadores"
                    quinto = "Salud"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Moda", 
                        "Accesorios", 
                        "Make", 
                        "Cepillos Alisadores", 
                        "Salud", 
                        "Más Vendidos"
                        ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                    
                elif estilo == "Moda":
                    df_path = "static/home/csv/ESPANHOL/"
                    pri = "Accesorios"
                    seg = "Femenino"
                    ter = "Masculino"
                    qua = "Niños"
                    qui = "Joyas"
                    sex = "Promociones"

                    primeiro = "Accesorios"
                    segundo = "Femenino"
                    terceiro = "Masculino"
                    quarto = "Niños"
                    quinto = "Joyas"
                    sexto = "Promociones"
                    colecoes = [
                        "Accesorios",
                        "Femenino",
                        "Masculino",
                        "Niños",
                        "Joyas",
                        "Promociones",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Saúde e Beleza":
                    df_path = "static/home/csv/ESPANHOL/"
                    
                    
                    pri = "Fitness y Ejercicio"
                    seg = "Belleza"
                    ter = "Más Vendidos"
                    qua = "Relajación"
                    qui = "Cuidado Personal"
                    sex = "Monitoreo de Salud"

                    primeiro = "Fitness-y-Ejercicio"
                    segundo = "Belleza"
                    terceiro = "Más-Vendidos"
                    quarto = "Relajación"
                    quinto = "Cuidado-Personal"
                    sexto = "Monitoreo-de-Salud"
                    colecoes = [
                        "Fitness y Ejercicio",
                        "Belleza",
                        "Más Vendidos",
                        "Relajación",
                        "Cuidado Personal",
                        "Monitoreo de Salud",
                    ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

            ##########################
            ########################## 
            else:  
                Tema = "https://github.com/ronildob/TEMAS2025/raw/refs/heads/main/tema_latan%2001-2025.zip"
                paginas = {
                    "Rastrear Pedidos": "static/home/politicas/seguimiento-del-pedido.txt",
                    "Contrato de Comercio Electrónico": "static/home/politicas/Contrato de Comercio Electrónico.docx",
                    "Términos de Uso": "static/home/politicas/Términos de Uso.docx",
                    "Plazo de Entrega": "static/home/politicas/Plazo de Entrega.docx",
                    "Política de Privacidad": "static/home/politicas/Política de Privacidad.docx",
                    "Cambios o Devoluciones": "static/home/politicas/Cambios o Devoluciones.docx",
                    "Sobre Nosotros": "static/home/politicas/Sobre Nosotros.docx",
                    "Pago Seguro": "static/home/politicas/Pago Seguro.docx",
                }
                politicas_para_atualizar = {
                    "Legal notice": "static/home/politicas/Contrato de Comercio Electrónico.docx",
                    "Terms of service": "static/home/politicas/Términos de Uso.docx",
                    "Shipping policy": "static/home/politicas/Plazo de Entrega.docx",
                    # "Privacy policy": "static/home/politicas/Política de Privacidad.docx",
                    "Refund policy": "static/home/politicas/Cambios o Devoluciones.docx",
                }
                ##############################
            
                if business_hours == "hora1":
                    antend = "Lunes a viernes de 9:00 a 18:00"
                elif business_hours == "hora2":
                    antend = "Lunes a viernes de 9:00 a 22:00"
                elif business_hours == "hora3":
                    antend = "Lunes a Sábado de  9:00 a 18"
                elif business_hours == "hora4":
                    antend = "Lunes a Sábado de  9:00 a 22"
                elif business_hours == "hora5":
                    antend = "Lunes a Viernes 9:00 a 18:00, Sábado 9:00 a 12:00"
                elif business_hours == "hora6":
                    antend = (
                        "Lunes a viernes de 9:00 a 22:00 - Sábado de 9:00 a 12:00"
                    )
                elif business_hours == "hora7":
                    antend = "Todos los días de 9:00 a 18:00"
                elif business_hours == "hora8":
                    antend = "Todos los días de 9:00 a 22:00"
                ##############################
                footer = f"""<p><strong>Correo electrónico:\u00a0</strong><a href="mailto:{email_suporte}">{email_suporte}</a></p><p><strong>WhatsApp:</strong>\u00a0<a href="https://wa.me/{whatsapp_completo}">{
                    cod_pais[pais]} {whatsapp}</a></p><p><strong>Horario de atención: </strong>{antend}</p>"""
                titlezap = f"Contáctanos \n<span>\n¡El mejor soporte es con {empresa}!\n</span>"
                ##############################
                ##############################
                if estilo == "Genérica":
                    df_path = "static/home/csv/ESPANHOL/30GENER_ESPANHOL.csv"

                    pri = "Casa"
                    seg = "Electrónica"
                    ter = "Pets"
                    qua = "Fitness"
                    qui = "Kids"
                    sex = "Más Vendidos"

                    primeiro = "Casa"
                    segundo = "Electrónica"
                    terceiro = "Pets"
                    quarto = "Fitness"
                    quinto = "Kids"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Casa",
                        "Electrónica",
                        "Pets",
                        "Fitness",
                        "Kids",
                        "Más Vendidos",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                    
                elif estilo == "Eletrônicos":
                    df_path = "static/home/csv/ESPANHOL/30ELETRONICOS.csv"
                    pri = "Auriculares"
                    seg = "Drones"
                    ter = "Smartwatch"
                    qua = "Accesorios"
                    qui = "Productos para Jugadores"
                    sex = "Más Vendidos"

                    primeiro = "Auriculares"
                    segundo = "Drones"
                    terceiro = "Smartwatch"
                    quarto = "Accesorios"
                    quinto = "Productos-para-Jugadores"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Auriculares", 
                        "Drones", 
                        "Smartwatch", 
                        "Accesorios", 
                        "Productos para Jugadores", 
                        "Más Vendidos"
                        ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                    
                elif estilo == "Kids":
                    df_path = "static/home/csv/ESPANHOL/30KIDS.csv"
                    pri = "Moda"
                    seg = "Accesorios"
                    ter = "Juguetes"
                    qua = "Cuidado"
                    qui = "Maternidad"
                    sex = "Más Vendidos"

                    primeiro = "Moda"
                    segundo = "Accesorios"
                    terceiro = "Juguetes"
                    quarto = "Cuidado"
                    quinto = "Maternidad"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Moda", 
                        "Accesorios", 
                        "Juguetes", 
                        "Cuidado", 
                        "Maternidad", 
                        "Más Vendidos"]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Casa":
                    df_path = "static/home/csv/ESPANHOL/"
                    
                    pri = "Accesorios"
                    seg = "Iluminación"
                    ter = "Decoración del ambiente"
                    qua = "Organización y almacenamiento"
                    qui = "Aromaterapia y Bienestar"
                    sex = "Productos de Temporada"

                    primeiro = "Accesorios"
                    segundo = "Iluminación"
                    terceiro = "Decoración-del-ambiente"
                    quarto = "Organización-y-almacenamiento"
                    quinto = "Aromaterapia-y-Bienestar"
                    sexto = "Productos-de-Temporada"
                    colecoes = [
                        "Accesorios",
                        "Iluminación",
                        "Decoración del ambiente",
                        "Organización y almacenamiento",
                        "Aromaterapia y Bienestar",
                        "Productos de Temporada",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Pet":
                    df_path = "static/home/csv/ESPANHOL/PETPORTUGAL.csv"

                    pri = "Camas"
                    seg = "Accesorios"
                    ter = "Comederos"
                    qua = "Juguetes"
                    qui = "Ropa"
                    sex = "Más Vendidos"

                    primeiro = "Camas"
                    segundo = "Accesorios"
                    terceiro = "Comederos"
                    quarto = "Juguetes"
                    quinto = "Ropa"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Camas",
                        "Accesorios",
                        "Comederos",
                        "Juguetes",
                        "Ropa",
                        "Más Vendidos",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Fitness":
                    df_path = "static/home/csv/ESPANHOL/FITNESSPORTUGAL.csv"

                    pri = "Moda"
                    seg = "Accesorios"
                    ter = "Zapatos"
                    qua = "Correctores"
                    qui = "Productos Deportivos"
                    sex = "Más Vendidos"

                    primeiro = "Moda"
                    segundo = "Accesorios"
                    terceiro = "Zapatos"
                    quarto = "Correctores"
                    quinto = "Productos-Deportivos"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Moda",
                        "Accesorios",
                        "Zapatos",
                        "Correctores",
                        "Productos Deportivos",
                        "Más Vendidos",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Masculino":
                    df_path = "static/home/csv/PT_BR/30MASCULINO.csv"

                    pri = "Accesorios"
                    seg = "Herramientas"
                    ter = "Pesca"
                    qua = "Ropa"
                    qui = "Automotor"
                    sex = "Más Vendidos"

                    primeiro = "Accesorios"
                    segundo = "Herramientas"
                    terceiro = "Pesca"
                    quarto = "Ropa"
                    quinto = "Automotor"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Accesorios", 
                        "Herramientas", 
                        "Pesca", 
                        "Ropa", 
                        "Automotor", 
                        "Más Vendidos"
                        ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                        
                elif estilo == "Feminino":
                    df_path = "static/home/csv/PT_BR/30FEMININO.csv"

                    pri = "Moda"
                    seg = "Accesorios"
                    ter = "Make"
                    qua = "Cepillos Alisadores"
                    qui = "Salud"
                    sex = "Más Vendidos"

                    primeiro = "Moda"
                    segundo = "Accesorios"
                    terceiro = "Make"
                    quarto = "Cepillos-Alisadores"
                    quinto = "Salud"
                    sexto = "Más-Vendidos"
                    colecoes = [
                        "Moda", 
                        "Accesorios", 
                        "Make", 
                        "Cepillos Alisadores", 
                        "Salud", 
                        "Más Vendidos"
                        ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }
                    
                elif estilo == "Moda":
                    df_path = "static/home/csv/ESPANHOL/"
                    pri = "Accesorios"
                    seg = "Femenino"
                    ter = "Masculino"
                    qua = "Niños"
                    qui = "Joyas"
                    sex = "Promociones"

                    primeiro = "Accesorios"
                    segundo = "Femenino"
                    terceiro = "Masculino"
                    quarto = "Niños"
                    quinto = "Joyas"
                    sexto = "Promociones"
                    colecoes = [
                        "Accesorios",
                        "Femenino",
                        "Masculino",
                        "Niños",
                        "Joyas",
                        "Promociones",
                    ]
                    
                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

                elif estilo == "Saúde e Beleza":
                    df_path = "static/home/csv/ESPANHOL/"
                    
                    
                    pri = "Fitness y Ejercicio"
                    seg = "Belleza"
                    ter = "Más Vendidos"
                    qua = "Relajación"
                    qui = "Cuidado Personal"
                    sex = "Monitoreo de Salud"

                    primeiro = "Fitness-y-Ejercicio"
                    segundo = "Belleza"
                    terceiro = "Más-Vendidos"
                    quarto = "Relajación"
                    quinto = "Cuidado-Personal"
                    sexto = "Monitoreo-de-Salud"
                    colecoes = [
                        "Fitness y Ejercicio",
                        "Belleza",
                        "Más Vendidos",
                        "Relajación",
                        "Cuidado Personal",
                        "Monitoreo de Salud",
                    ]

                    # urls = {
                        # "desktop1": "",
                        # "desktop2": "",
                        # "desktop3": "",
                        # "mobile1": "",
                        # "mobile2": "",
                        # "mobile3": "",
                        # "capa1": "",
                        # "capa2": "",
                        # "capa3": "",
                        # "capa4": "",
                        # "capa5": "",
                        # "capa6": "",
                    # }

            ##########################
            
            ##########################
            ##########################

            shopify.Session.setup(api_key=API_KEY, secret=API_SECRET)
            session = shopify.Session(SHOP_URL, API_VERSION, PRIVATE_APP_PASSWORD)
            shopify.ShopifyResource.activate_session(session)

            df = pd.read_csv(df_path)

            def importar_produtos(df, pais):
                try:
                    erros = []  # Lista para acumular erros de cada produto

                    # Dicionário com as taxas de câmbio para cada país em
                    # relação ao Brasil
                    taxa_cambio = {
                        "Brasil": 1.0,
                        "Argentina": 0.04,
                        "Bolívia": 0.14,
                        "Chile": 0.14,
                        "Colômbia": 0.72,
                        "Equador": 0.27,
                        "Paraguai": 0.14,
                        "Peru": 0.24,
                        "Uruguai": 0.25,
                        "Venezuela": 0.11,
                        "EUA": 5.0,
                        "Portugal": 5.5,
                        "Espanha": 5.5,
                    }

                    # Dicionário de formatação de moeda para cada país
                    formato_moeda = {
                        "Brasil": "{:.2f}",
                        "Argentina": "{:.2f}",
                        "Bolívia": "{:.2f}",
                        "Chile": "{:.2f}",
                        "Colômbia": "{:.2f}",
                        "Equador": "{:.2f}",
                        "Paraguai": "{:.2f}",
                        "Peru": "S/. {:.2f}",
                        "Uruguai": "{:.2f}",
                        "Venezuela": "{:.2f}",
                        "EUA": "{:.2f}",
                        "Portugal": "{:.2f}",
                        "Espanha": "{:.2f}",
                    }

                    for index, row in df.iterrows():
                        try:
                            handle = row["Handle"]
                            title = (
                                row["Title"]
                                if pd.notna(row["Title"]) and row["Title"].strip()
                                else "Título padrão"
                            )
                            image_src = row["Image Src"]

                            # Obter a taxa de câmbio com base no país
                            taxa = taxa_cambio.get(pais, 1.0)

                            # Ajuste de preços dependendo do país
                            if pais in ["EUA", "Portugal", "Espanha"]:
                                variant_price = (
                                    format(row["Variant Price"] / taxa, ".2f")
                                    if pd.notna(row["Variant Price"])
                                    else "0.00"
                                )
                                variant_compare_at_price = (
                                    format(
                                        row["Variant Compare At Price"] / taxa,
                                        ".2f",
                                    )
                                    if pd.notna(row["Variant Compare At Price"])
                                    else "0.00"
                                )
                            else:
                                variant_price = (
                                    format(row["Variant Price"] * taxa, ".2f")
                                    if pd.notna(row["Variant Price"])
                                    else "0.00"
                                )
                                variant_compare_at_price = (
                                    format(
                                        row["Variant Compare At Price"] * taxa,
                                        ".2f",
                                    )
                                    if pd.notna(row["Variant Compare At Price"])
                                    else "0.00"
                                )

                            # Aplicar formatação de moeda
                            variant_price = formato_moeda.get(
                                pais, "${:.2f}"
                            ).format(float(variant_price))
                            variant_compare_at_price = formato_moeda.get(
                                pais, "${:.2f}"
                            ).format(float(variant_compare_at_price))

                            # Criação do produto
                            product = shopify.Product(
                                {
                                    "title": title,
                                    "body_html": (
                                        row["Body (HTML)"]
                                        if pd.notna(row["Body (HTML)"])
                                        else ""
                                    ),
                                    "vendor": (
                                        row["Vendor"]
                                        if pd.notna(row["Vendor"])
                                        else ""
                                    ),
                                    "product_type": (
                                        row["Type"] if pd.notna(row["Type"]) else ""
                                    ),
                                    "tags": (
                                        unidecode.unidecode(row["Tags"])
                                        .lower()
                                        .replace(" ", "-")
                                        if pd.notna(row["Tags"])
                                        else ""
                                    ),
                                    "published": (
                                        row["Published"]
                                        if pd.notna(row["Published"])
                                        else True
                                    ),
                                }
                            )

                            if not product.save():
                                erros.append(
                                    f"Erro ao salvar produto '{title}': {product.errors.full_messages()}")
                                continue

                            # Adicionar variante ao produto
                            variant = shopify.Variant(
                                {
                                    "taxable": "false",
                                    "price": variant_price,
                                    "compare_at_price": variant_compare_at_price,
                                    "inventory_quantity": 1000,
                                    "inventory_management": "shopify",
                                    "inventory_policy": "continue",
                                }
                            )
                            product.variants = [variant]

                            if not product.save():
                                erros.append(
                                    f"Erro ao salvar variante para o produto '{title}': {product.errors.full_messages()}")
                                continue

                            # Adicionar imagem ao produto
                            if pd.notna(image_src):
                                image = shopify.Image({"src": image_src})
                                image.product_id = product.id
                                if not image.save():
                                    erros.append(
                                        f"Erro ao salvar imagem para o produto '{title}': {image.errors.full_messages()}")

                            print(f"Produto '{title}' carregado com sucesso!")

                        except Exception as e:
                            erros.append(
                                f"Erro ao importar produto '{title}': {str(e)}")

                    # Retorno consolidado com a lista de erros ou sucesso
                    if erros:
                        return "Erro: " + " | ".join(erros)
                    else:
                        return (
                            "Sucesso: Todos os produtos foram importados com êxito."
                        )

                except Exception as e:
                    return f"Erro: Exceção geral ao importar produtos - {str(e)}"

            def importar_tema(Tema):
                try:
                    # Configuração do novo tema na Shopify
                    new_theme = shopify.Theme()
                    new_theme.name = "ShopBotfy"
                    new_theme.src = Tema
                    new_theme.role = "main"

                    # Tentativa de salvar o tema e verificação de erros
                    if new_theme.save():
                        return f"Sucesso: Tema '{new_theme.name}' criado com êxito. ID: {new_theme.id}"
                    else:
                        erros = new_theme.errors.full_messages()
                        return f"Erro: Falha ao criar o tema. Detalhes: {' | '.join(erros)}"

                except Exception as e:
                    return f"Erro: Exceção ao importar tema - {str(e)}"

    
            # imagens
            def download_svg(url):
                response = requests.get(url)
                if response.status_code == 200:
                    local_path = url.split("/")[-1]
                    with open(local_path, "wb") as file:
                        file.write(response.content)
                    return local_path
                else:
                    raise Exception(f"Erro ao baixar o SVG: {response.status_code}")

            def edit_svg_colors(svg_path, color_mapping):
                try:
                    tree = ET.parse(svg_path)
                    root = tree.getroot()
                    namespace = {"svg": "http://www.w3.org/2000/svg"}

                    for element in root.findall(".//svg:*", namespaces=namespace):
                        fill = element.attrib.get("fill")
                        if fill and fill in color_mapping:
                            element.set("fill", color_mapping[fill])

                    edited_path = "edited_" + svg_path
                    tree.write(edited_path)
                    return edited_path
                except Exception as e:
                    raise Exception(f"Erro ao editar o SVG: {e}")

            def validate_svg(svg_path):
                try:
                    tree = ET.parse(svg_path)
                    root = tree.getroot()
                    return root.tag == "{http://www.w3.org/2000/svg}svg"
                except Exception:
                    return False

            def convert_svg_to_png(svg_path):
                try:
                    response = cloudinary.uploader.upload(svg_path, resource_type="image", format="png", folder="banners_e_capas")
                    return response["secure_url"], response["public_id"]
                except Exception as e:
                    raise Exception(f"Erro ao converter o SVG para PNG: {e}")

            def delete_image_from_cloudinary(public_id):
                try:
                    cloudinary.uploader.destroy(public_id)
                    print(f"Imagem {public_id} deletada com sucesso no Cloudinary.")
                except Exception as e:
                    print(f"Erro ao deletar imagem no Cloudinary: {e}")

            def upload_png_to_shopify(png_url, file_name):
                mutation = """
                    mutation fileCreate($files: [FileCreateInput!]!) {
                        fileCreate(files: $files) {
                            files {
                                id
                                alt
                            }
                            userErrors {
                                message
                            }
                        }
                    }
                """
                file_data = {
                    "files": [
                        {"alt": file_name, "originalSource": png_url}
                    ]
                }

                response = requests.post(
                    f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json",
                    json={"query": mutation, "variables": file_data},
                    headers={"X-Shopify-Access-Token": PRIVATE_APP_PASSWORD},
                )
                if response.status_code != 200:
                    raise Exception(f"Erro ao enviar para Shopify: {response.text}")

            public_ids_to_delete = cache.get(f"public_ids_{username}", [])

            urls_shopify = {}

            for name, url in urls.items():
                try:
                    svg_path = download_svg(url)
                    edited_svg_path = edit_svg_colors(svg_path, {"#000000": cor1})

                    if validate_svg(edited_svg_path):
                        png_url, public_id = convert_svg_to_png(edited_svg_path)
                        upload_png_to_shopify(png_url, name)
                        urls_shopify[name] = png_url
                        public_ids_to_delete.append(public_id)
                    else:
                        raise Exception(f"SVG inválido: {edited_svg_path}")
                except Exception as e:
                    print(f"Erro no processamento de {name}: {e}")

            def criar_colecoes(colecoes):
                try:
                    erros = []  # Lista para acumular mensagens de erro

                    # Iteração sobre cada coleção a ser criada
                    for nome in colecoes:
                        try:
                            nova_colecao = shopify.CustomCollection()
                            nova_colecao.title = nome

                            # Tentativa de salvar a nova coleção
                            if not nova_colecao.save():
                                erros.append(
                                    f"Erro ao salvar a coleção '{nome}': {nova_colecao.errors.full_messages()}"
                                )
                            else:
                                print(f"Coleção '{nome}' criada com sucesso!")

                        except Exception as e:
                            erros.append(
                                f"Erro ao criar a coleção '{nome}': {str(e)}"
                            )

                    # Retorna uma mensagem final com base no resultado
                    if erros:
                        return "Erro: " + " | ".join(erros)
                    else:
                        return "Sucesso: Todas as coleções foram criadas com êxito."

                except Exception as e:
                    return f"Erro: Exceção geral ao criar coleções - {str(e)}"

            def criar_paginas(paginas):
                try:
                    erros = []  # Lista para acumular erros de cada página

                    # Iteração sobre as páginas a serem criadas
                    for titulo, caminho_arquivo in paginas.items():
                        try:
                            # Leitura do conteúdo com base no tipo de
                            # arquivo
                            if caminho_arquivo.endswith(".docx"):
                                doc = Document(caminho_arquivo)
                                conteudo = "\n".join(
                                    [p.text for p in doc.paragraphs]
                                )
                            elif caminho_arquivo.endswith(".txt"):
                                with open(caminho_arquivo, "r") as f:
                                    conteudo = f.read()
                            else:
                                erros.append(
                                    f"Formato de arquivo não suportado: {caminho_arquivo}"
                                )
                                continue

                            # Criação da nova página na Shopify
                            nova_pagina = shopify.Page()
                            nova_pagina.title = titulo
                            nova_pagina.body_html = conteudo

                            # Salvando a página e verificando sucesso
                            if not nova_pagina.save():
                                erros.append(
                                    f"Erro ao salvar a página '{titulo}': {nova_pagina.errors.full_messages()}"
                                )

                        except Exception as e:
                            erros.append(
                                f"Erro ao criar a página '{titulo}': {str(e)}"
                            )

                    # Verificação final para consolidar a resposta
                    if erros:
                        return "Erro: " + " | ".join(erros)
                    else:
                        return "Sucesso: Todas as páginas foram criadas com êxito."

                except Exception as e:
                    return f"Erro: Exceção geral ao criar páginas - {str(e)}"

            def ler_documento(caminho):
                document = Document(caminho)
                texto = ""
                for paragraph in document.paragraphs:
                    texto += paragraph.text + "\n"
                return texto

            def atualizar_politicas():
                try:
                    # Endpoint para atualizar as políticas
                    endpoint = (
                        f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json"
                    )

                    # Cabeçalhos da requisição com autenticação
                    headers = {
                        "X-Shopify-Access-Token": PRIVATE_APP_PASSWORD,
                        "Content-Type": "application/json",
                    }

                    # Mutation para atualizar a política
                    mutation = """
                            mutation shopPolicyUpdate($shopPolicy: ShopPolicyInput!) {
                                shopPolicyUpdate(shopPolicy: $shopPolicy) {
                                    shopPolicy {
                                        id
                                    }
                                    userErrors {
                                        field
                                        message
                                    }
                                }
                            }
                            """

                    # Lista para coletar erros (se houver) durante o
                    # processo
                    erros = []

                    # Iteração para cada política e tentativa de
                    # atualização
                    for (
                        titulo,
                        caminho_documento,
                    ) in politicas_para_atualizar.items():
                        corpo = ler_documento(caminho_documento)
                        politica_para_atualizar = {
                            "body": corpo,
                            "type": titulo.upper().replace(" ", "_"),
                        }

                        variables = {"shopPolicy": politica_para_atualizar}

                        # Envia a requisição POST com a mutação
                        response = requests.post(
                            endpoint,
                            json={"query": mutation, "variables": variables},
                            headers=headers,
                        )

                        # Verifica a resposta da requisição
                        if response.status_code == 200:
                            data = response.json()
                            if "errors" in data:
                                for error in data["errors"]:
                                    erros.append(
                                        f"Erro ao atualizar a política '{titulo}': {error['message']}"
                                    )
                            elif (
                                data.get("data", {})
                                .get("shopPolicyUpdate", {})
                                .get("userErrors")
                            ):
                                # Verifica os erros de usuário
                                for user_error in data["data"]["shopPolicyUpdate"][
                                    "userErrors"
                                ]:
                                    erros.append(
                                        f"Erro ao atualizar a política '{titulo}': {user_error['message']}"
                                    )
                            else:
                                print(
                                    f"Política '{titulo}' atualizada com sucesso!"
                                )
                        else:
                            erros.append(
                                f"Erro ao enviar a requisição para a política '{titulo}': Status Code - {response.status_code}, Response Text - {response.text}"
                            )

                    # Se houver erros, retorna todos como uma string única
                    if erros:
                        return "Erro: " + " | ".join(erros)
                    else:
                        return "Sucesso: Todas as políticas foram atualizadas com êxito."

                except Exception as e:
                    return f"Erro: Exceção durante a atualização das políticas - {str(e)}"

            def adicionar_produtos_colecoes(colecoes, produtos):
                try:
                    colecoes = shopify.CustomCollection.find()
                    produtos = shopify.Product.find()
                    headers = {
                        "X-Shopify-Access-Token": PRIVATE_APP_PASSWORD,
                        "Content-Type": "application/json",
                    }
                    mutation = """
                                mutation collectionAddProductsV2($id: ID!, $productIds: [ID!]!) {
                                    collectionAddProductsV2(id: $id, productIds: $productIds) {
                                        job {
                                            done
                                            id
                                        }
                                        userErrors {
                                            field
                                            message
                                        }
                                    }
                                }
                                """

                    # Itera sobre cada coleção e produto
                    for colecao in colecoes:
                        collection_id = colecao.id
                        collection_handle = colecao.handle

                        # Filtra os produtos que correspondem à coleção
                        # pelo handle
                        ids_produtos_correspondentes = [
                            produto.id
                            for produto in produtos
                            if collection_handle in produto.tags
                        ]

                        # Adiciona todos os produtos correspondentes à
                        # coleção
                        for product_id in ids_produtos_correspondentes:
                            # Criação das variáveis da mutação
                            variables = {
                                "id": f"gid://shopify/Collection/{collection_id}",
                                "productIds": [
                                    f"gid://shopify/Product/{product_id}"
                                ],
                            }

                            # Criação do corpo da requisição
                            body = {
                                "query": mutation,
                                "variables": variables,
                            }
                            # Realização da requisição à API da Shopify
                            response = requests.post(
                                f"https://{SHOP_URL}/admin/api/{API_VERSION}/graphql.json",headers=headers,data=json.dumps(body),
                            )

                            # Verificação da resposta para cada produto
                            if response.status_code != 200:
                                raise Exception(
                                    f"Erro ao adicionar produto {product_id} à coleção {collection_id}: {response.content}"
                                )

                    return "Sucesso: Todos os produtos foram adicionados às coleções com êxito."

                except Exception as e:
                    return f"Erro: {str(e)}"

            
            def editar_tema(urls_shopify, max_tentativas=5):
                tentativa = 0
                while tentativa < max_tentativas:
                    try:
                        temas = shopify.Theme.find()
                        tema_publicado = [tema for tema in temas if tema.role == "main"][0]
                        id_tema_publicado = tema_publicado.id
                        theme = shopify.Theme.find(id_tema_publicado)
                        settings_asset = shopify.Asset.find(
                            "config/settings_data.json", theme_id=theme.id
                        )

                        print("Arquivo de configurações encontrado.")

                        settings_data = settings_asset.value

                        # Verifica se settings_data é uma string e tenta carregar como JSON
                        if isinstance(settings_data, str):
                            settings_data = json.loads(settings_data)

                        print("Configurações carregadas com sucesso.")


                        # Armazenando os valores em variáveis
                        desktop1 = f"shopify://shop_images/{urls_shopify['desktop1'].split('/')[-1]}"
                        desktop2 = f"shopify://shop_images/{urls_shopify['desktop2'].split('/')[-1]}"
                        desktop3 = f"shopify://shop_images/{urls_shopify['desktop3'].split('/')[-1]}"
                        mobile1 = f"shopify://shop_images/{urls_shopify['mobile1'].split('/')[-1]}"
                        mobile2 = f"shopify://shop_images/{urls_shopify['mobile2'].split('/')[-1]}"
                        mobile3 = f"shopify://shop_images/{urls_shopify['mobile3'].split('/')[-1]}"
                        capa1 = f"shopify://shop_images/{urls_shopify['capa1'].split('/')[-1]}"
                        capa2 = f"shopify://shop_images/{urls_shopify['capa2'].split('/')[-1]}"
                        capa3 = f"shopify://shop_images/{urls_shopify['capa3'].split('/')[-1]}"
                        capa4 = f"shopify://shop_images/{urls_shopify['capa4'].split('/')[-1]}"
                        capa5 = f"shopify://shop_images/{urls_shopify['capa5'].split('/')[-1]}"
                        capa6 = f"shopify://shop_images/{urls_shopify['capa6'].split('/')[-1]}"

                        if tema_base == "Evolution":
                            #######################################################
                            # CORES DO TEMA ###########
                            #######################################################

                            ###### HEADER #########
                            settings_data["current"]["header_accent_color"] = acent
                            settings_data["current"]["header_background"] = header_background
                            settings_data["current"]["header_text_color"] = header_text_color
                            settings_data["current"]["accent_color"] = accent_color
                            settings_data["current"]["link_color"] = link_color
                            settings_data["current"]["header_light_text_color"] = header_light_text_color
                            settings_data["current"]["primary_button_background"] = primary_button_background
                            settings_data["current"]["secondary_button_background"] = secondary_button_background
                            settings_data["current"]["product_on_sale_accent"] = product_on_sale_accent
                            settings_data["current"]["product_cor_do_preco_semdesc"] = product_cor_do_preco_semdesc
                            settings_data["current"]["product_cor_do_preco"] = product_cor_do_preco
                            settings_data["current"]["product_cor_dos_titles"] = product_cor_dos_titles
                            settings_data["current"]["product_in_stock_color"] = product_in_stock_color

                            settings_data["current"]["sections"]["header"]["settings"]["background1"] = degrade1
                            settings_data["current"]["sections"]["header"]["settings"]["background2"] = degrade2
                            settings_data["current"]["sections"]["header"]["settings"]["background3"] = degrade3
                            settings_data["current"]["sections"]["header"]["settings"]["background4"] = degrade4
                            settings_data["current"]["sections"]["header"]["settings"]["backgroundlocal1"] = cor1
                            settings_data["current"]["sections"]["header"]["settings"]["backgroundlocal2"] = cor2
                            settings_data["current"]["sections"]["header"]["settings"]["colortextlocal"] = textcolor
                            settings_data["current"]["sections"]["header"]["settings"]["navigation_phone_number"] = whatsapp
                            settings_data["current"]["sections"]["header"]["settings"]["navigation_email"] = email_suporte

                            ###### FOOTER #########

                            #######################################################
                            ####    DADOS DA LOJA   ###########
                            #######################################################
                            ### CABEÇALHO ######
                            settings_data["current"]["sections"]["header"]["settings"]["navigation_phone_number"] = whatsapp
                            settings_data["current"]["sections"]["header"]["settings"]["navigation_email"] = email_suporte

                            """#### WHATSAPP #######
                            settings_data['current']['addzap'] = True
                            settings_data['current']['numberzap'] = tel_whats
                            settings_data['current']['titlezap'] = titlezap
                            """

                            ##### EDITA FOOTER ######
                            settings_data["current"]["sections"]["footer"]["blocks"]["text_image_t6X8qj"]["settings"]["content"] = footer
                            settings_data["current"]["footer_background_color"] = footer_background
                            settings_data["current"]["footer_heading_text_color"] = footer_text_color
                            settings_data["current"]["footer_body_text_color"] = footer_text_color

                            ##### INSERIR BANNERS ##########
                            settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-1"]["settings"]["image"] = desktop1
                            settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-1"]["settings"]["mobile_image"] = mobile1
                            settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-2"]["settings"]["image"] = desktop2
                            settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-2"]["settings"]["mobile_image"] = mobile2
                            settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-3"]["settings"]["image"] = desktop3
                            settings_data["current"]["sections"]["slideshow"]["blocks"]["slide-3"]["settings"]["mobile_image"] = mobile3

                            #### COLEÇÔES E CAPAS  ###########
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_6EMmhW"]["settings"]["collection"] = primeiro
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_6EMmhW"]["settings"]["image"] = capa1
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_Anzt8E"]["settings"]["collection"] = segundo
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_Anzt8E"]["settings"]["image"] = capa2
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_NqzGkT"]["settings"]["collection"] = terceiro
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_NqzGkT"]["settings"]["image"] = capa3
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_X6edXj"]["settings"]["collection"] = quarto
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_X6edXj"]["settings"]["image"] = capa4
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_pPAM48"]["settings"]["collection"] = quinto
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_pPAM48"]["settings"]["image"] = capa5
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_j4N3Rp"]["settings"]["collection"] = sexto
                            settings_data["current"]["sections"]["collection_list_yr77bz"]["blocks"]["collection_j4N3Rp"]["settings"]["image"] = capa6

                            #### FEATURED COLECC #####################
                            settings_data["current"]["sections"]["featured_collection_EXHCLn"]["settings"]["collection"] = primeiro
                            settings_data["current"]["sections"]["featured_collection_EXHCLn"]["settings"]["title"] = primeiro
                            settings_data["current"]["sections"]["featured_collection_jLrtgd"]["settings"]["collection"] = segundo
                            settings_data["current"]["sections"]["featured_collection_jLrtgd"]["settings"]["title"] = segundo
                            settings_data["current"]["sections"]["featured_collection_YAcpkY"]["settings"]["collection"] = terceiro
                            settings_data["current"]["sections"]["featured_collection_YAcpkY"]["settings"]["title"] = terceiro
                            settings_data["current"]["sections"]["featured_collection_NRCkDA"]["settings"]["collection"] = quarto
                            settings_data["current"]["sections"]["featured_collection_NRCkDA"]["settings"]["title"] = quarto
                            settings_data["current"]["sections"]["featured_collection_YhPALQ"]["settings"]["collection"] = quinto
                            settings_data["current"]["sections"]["featured_collection_YhPALQ"]["settings"]["title"] = quinto
                            settings_data["current"]["sections"]["collection_with_image_6pAgRL"]["settings"]["collection"] = sexto
                            settings_data["current"]["sections"]["collection_with_image_6pAgRL"]["settings"]["title"] = sexto
                            ##### COLEÇÃO image_6pAgRL#########
                            settings_data["current"]["sections"]["collection_with_image_6pAgRL"]["settings"]["background"] = cor1
                            settings_data["current"]["sections"]["collection_with_image_6pAgRL"]["settings"]["text_color"] = textcolor

                            settings_asset.value = json.dumps(settings_data)
                            settings_asset.save()

                            return "Sucesso: Estilização aplicada com êxito."
                        
                        elif tema_base == "Dropmeta":

                            #######################################################
                            ### CORES DO TEMA ###########
                            #######################################################
                            ###### BARRA DE ANUNCIOS #########
                            
                            settings_data["current"]["sections"]["announcement-bar"]["settings"]["background1"] = cor1
                            settings_data["current"]["sections"]["announcement-bar"]["settings"]["background2"] = cor2
                            settings_data["current"]["sections"]["announcement-bar"]["settings"]["text_color"] = textcolor
                            
                            
                            # Header
                            settings_data["current"]["header_accent_color"] = acent
                            settings_data["current"]["accent_color"] = accent_color                       
                            settings_data["current"]["link_color"] = link_color
                            settings_data["current"]["header_background"] = header_background
                            settings_data["current"]["header_text_color"] = header_text_color
                            settings_data["current"]["header_light_text_color"] = header_light_text_color
                            settings_data["current"]["primary_button_background"] = primary_button_background
                            settings_data["current"]["secondary_button_background"] = secondary_button_background
                            settings_data["current"]["product_on_sale_accent"] = product_on_sale_accent
                            settings_data["current"]["product_cor_do_preco_semdesc"] = product_cor_do_preco_semdesc
                            settings_data["current"]["product_cor_do_preco"] = product_cor_do_preco
                            settings_data["current"]["product_cor_dos_titles"] = product_cor_dos_titles
                            settings_data["current"]["product_in_stock_color"] = product_in_stock_color
                            settings_data["current"]["sections"]["header"]["settings"]["background1"] = degrade1
                            settings_data["current"]["sections"]["header"]["settings"]["background2"] = degrade2
                            settings_data["current"]["sections"]["header"]["settings"]["background3"] = degrade3 
                            settings_data["current"]["sections"]["header"]["settings"]["background4"] = degrade4
                            
                                
                            ###### FOOTER #########
                            settings_data["current"]["sections"]["footer"]["settings"]["background1"] = degradefooter1
                            settings_data["current"]["sections"]["footer"]["settings"]["background2"] = degradefooter2
                            settings_data["current"]["sections"]["footer"]["settings"]["background3"] = degradefooter3
                            settings_data["current"]["sections"]["footer"]["settings"]["background4"] = degradefooter4
                            settings_data["current"]["footer_background"] = footer_background
                            settings_data["current"]["footer_text_color"] = footer_text_color

                            ##### COLEÇÃO #########
                            
                            settings_data["current"]["sections"]["1649913264519d35a7"]["settings"]["background1"] = cor1
                            settings_data["current"]["sections"]["1649913264519d35a7"]["settings"]["background2"] = cor2
                            settings_data["current"]["sections"]["1649913264519d35a7"]["settings"]["text_color"] = textcolor
                            
                                
                            #######################################################
                            ####    DADOS DA LOJA   ###########
                            #######################################################
                            ### CABEÇALHO ######
                            
                            settings_data["current"]["sections"]["header"]["settings"]["navigation_phone_number"] = whatsapp
                            settings_data["current"]["sections"]["header"]["settings"]["navigation_email"] = email_suporte

                            #### WHATSAPP #######
                            settings_data["current"]["addzap"] = True
                            settings_data["current"]["numberzap"] = tel_whats
                            settings_data["current"]["titlezap"] = titlezap
                            
                            ##### EDITA FOOTER ######
                            settings_data["current"]["sections"]["footer"]["blocks"]["c9a6c378-573f-4c54-9fc6-5ca6e279f3f2"]["settings"]["content"] = footer

                            ##### INSERIR BANNERS ##########
                            # banner 1
                            settings_data["current"]["sections"]["slideshow"]["blocks"]["4665d2ed-db3b-479e-8984-d272fdfab8d8"]["settings"]["image"] = desktop1
                            settings_data["current"]["sections"]["slideshow"]["blocks"]["4665d2ed-db3b-479e-8984-d272fdfab8d8"]["settings"]["mobile_image"] = mobile1
                            # banner 2
                            settings_data["current"]["sections"]["slideshow"]["blocks"]["bbea030d-de91-4721-9ea1-9660f647bbd7"]["settings"]["image"] = desktop2
                            settings_data["current"]["sections"]["slideshow"]["blocks"]["bbea030d-de91-4721-9ea1-9660f647bbd7"]["settings"]["mobile_image"] = mobile2
                            
                            # banner 3
                            if pais == "Espanha":
                                settings_data["current"]["sections"]["slideshow_7n83kx"]["blocks"]["image_H6WbRH"]["settings"]["image"] = desktop3
                                settings_data["current"]["sections"]["slideshow_7n83kx"]["blocks"]["image_H6WbRH"]["settings"]["mobile_image"] = mobile3
                                
                            elif pais == "EUA":
                                settings_data["current"]["sections"]["slideshow_37CJTG"]["blocks"]["image_CQGpKw"]["settings"]["image"] = desktop3
                                settings_data["current"]["sections"]["slideshow_37CJTG"]["blocks"]["image_CQGpKw"]["settings"]["mobile_image"] = mobile3


                            elif pais == "Portugal":
                                settings_data["current"]["sections"]["slideshow_DViiTn"]["blocks"]["image_WiCjtP"]["settings"]["image"] = desktop3
                                settings_data["current"]["sections"]["slideshow_DViiTn"]["blocks"]["image_WiCjtP"]["settings"]["mobile_image"] = mobile3


                            elif pais == "Brasil":
                                settings_data["current"]["sections"]["slideshow_NcxY8c"]["blocks"]["image_8RNyne"]["settings"]["image"] = desktop3
                                settings_data["current"]["sections"]["slideshow_NcxY8c"]["blocks"]["image_8RNyne"]["settings"]["mobile_image"] = mobile3

                            else:
                                settings_data["current"]["sections"]["slideshow_LR4JhV"]["blocks"]["image_7Y9Pdh"]["settings"]["image"] = desktop3
                                settings_data["current"]["sections"]["slideshow_LR4JhV"]["blocks"]["image_7Y9Pdh"]["settings"]["mobile_image"] = mobile3
                            
                            #### COLEÇÔES E CAPAS  ###########
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-1"]["settings"]["collection"] = primeiro
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-1"]["settings"]["image"] = capa1
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-2"]["settings"]["collection"] = segundo
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-2"]["settings"]["image"] = capa2
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-3"]["settings"]["collection"] = terceiro
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-3"]["settings"]["image"] = capa3
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-5"]["settings"]["collection"] = quarto
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["1649489657b78ca727-5"]["settings"]["image"] = capa4
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["d5a5159b-779d-45d2-871b-4a5e622bb3b2"]["settings"]["collection"] = quinto
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["d5a5159b-779d-45d2-871b-4a5e622bb3b2"]["settings"]["image"] = capa5
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["0c56a93b-d66f-4823-ae91-bd01bb634a00"]["settings"]["collection"] = sexto
                            settings_data["current"]["sections"]["1649489657b78ca727"]["blocks"]["0c56a93b-d66f-4823-ae91-bd01bb634a00"]["settings"]["image"] = capa6
                            
                            #### FEATURED COLECC #####################
                            settings_data["current"]["sections"]["16153539675ef431a7"]["settings"]["collection"] = sexto
                            settings_data["current"]["sections"]["16153539675ef431a7"]["settings"]["title"] = sex
                            settings_data["current"]["sections"]["16153541158fc73818"]["settings"]["collection"] = primeiro
                            settings_data["current"]["sections"]["16153541158fc73818"]["settings"]["title"] = pri
                            settings_data["current"]["sections"]["1649913264519d35a7"]["settings"]["collection"] = segundo      
                            settings_data["current"]["sections"]["1649913264519d35a7"]["settings"]["title"] = seg
                            settings_data["current"]["sections"]["7196b9c9-7ec2-46f0-87be-f84dcc1d14b5"]["settings"]["collection"] = terceiro
                            settings_data["current"]["sections"]["7196b9c9-7ec2-46f0-87be-f84dcc1d14b5"]["settings"]["title"] = ter
                            settings_data["current"]["sections"]["9dd41389-f9fe-453e-b65e-1cc7383b8c1f"]["settings"]["collection"] = quarto
                            settings_data["current"]["sections"]["9dd41389-f9fe-453e-b65e-1cc7383b8c1f"]["settings"]["title"] = qua
                            settings_data["current"]["sections"]["329bc2be-34eb-46d8-8601-39d8eda721d9"]["settings"]["collection"] = quinto
                            settings_data["current"]["sections"]["329bc2be-34eb-46d8-8601-39d8eda721d9"]["settings"]["title"] = qui
                            
                            
                            
                            settings_asset.value = json.dumps(settings_data)
                            settings_asset.save()

                            return "Sucesso: Estilização aplicada com êxito."

                    except Exception as e:
                        tentativa += 1
                        print(f"Tentativa {tentativa} falhou. Erro: {str(e)}")
                        time.sleep(2)  # Aguarda 2 segundos antes de tentar novamente
                
                return "Erro: Número máximo de tentativas atingido."
        

            # Chamar funções
            if not formulario_user.produtos:
                resposta = importar_produtos(df, pais)
                formulario_user.produtos = True
                formulario_user.resposta_produtos = resposta
                formulario_user.save()
            if not formulario_user.tema:
                resposta = importar_tema(Tema)
                formulario_user.tema = True
                formulario_user.resposta_tema = resposta
                formulario_user.save()
            
            # if not formulario_user.banners:
            #     resposta = subir_banners(urls)
            #     formulario_user.banners = True
            #     formulario_user.resposta_banners = resposta
            #     formulario_user.save()
                
            if not formulario_user.colecoes:
                resposta = criar_colecoes(colecoes)
                formulario_user.colecoes = True
                formulario_user.resposta_colecoes = resposta
                formulario_user.save()
            if not formulario_user.paginas:
                resposta = criar_paginas(paginas)
                formulario_user.paginas = True
                formulario_user.resposta_paginas = resposta
                formulario_user.save()
            if not formulario_user.politicas:
                resposta = atualizar_politicas()
                formulario_user.politicas = True
                formulario_user.resposta_politicas = resposta
                formulario_user.save()
            if not formulario_user.prod_col:
                resposta = adicionar_produtos_colecoes(colecoes, produtos)
                formulario_user.prod_col = True
                formulario_user.resposta_prod_col = resposta
                formulario_user.save()
            if not formulario_user.estilizacao:
                resposta = editar_tema(urls_shopify)
                formulario_user.estilizacao = True
                formulario_user.lojaproduzida = True
                formulario_user.resposta_estilizacao = (
                    resposta  # Campo novo no modelo para armazenar a resposta
                )
                formulario_user.save()

            

            # Deletar imagens no Cloudinary após edição
            for public_id in public_ids_to_delete:
                delete_image_from_cloudinary(public_id)

            # Limpar a cache após deletar
            cache.delete(f"public_ids_{username}")
            shopify.ShopifyResource.clear_session()
            
            
            # Atualiza o status para 'completed'
            formulario_user.status = 'finalizado'  # Ao final de tudo
            formulario_user.save()
            print(f"loja de {empresa} concluida com sucesso!")
            
        else:
            print("Usuário não encontrado")
            return 'user_not_found'

    except Exception as e:
        # Tenta novamente se ocorrer erro
        formulario_user.status = 'erro'
        formulario_user.save()
        
        try:
            print(f"Erro ao processar {username}, tentando novamente...: {str(e)}")
            self.retry(exc=e)  # Tenta novamente
        except MaxRetriesExceededError:
            print(f"Falha ao processar {username} após várias tentativas.")

        return 'error'


@shared_task
def processar_fila_de_espera():
    logger.info("Iniciando a tarefa 'processar_fila_de_espera'")
    filas = FilaDeEspera.objects.filter(processado=False)
    logger.info(f"Encontradas {filas.count()} filas de espera não processadas.")
    
    for fila in filas:
        logger.info(f"Processando fila: {fila.id}")
        agencia = fila.agencia
        formulario_disponivel = AAFormuUserer.objects.filter(
            agencia=agencia, status="reservado"
        ).first()

        if formulario_disponivel:
            logger.info(f"Formulario encontrado para a agencia: {agencia}")
            dados = fila.dados_formulario
            formulario_disponivel.link_store = dados.get("link_store", "")
            formulario_disponivel.url_loja = dados.get("url_loja", "")
            formulario_disponivel.pais = dados.get("pais", "")
            formulario_disponivel.nome = dados.get("nome", "")
            formulario_disponivel.telefone = dados.get("telefone", "")
            formulario_disponivel.email_suporte = dados.get("email_suporte", "")
            formulario_disponivel.empresa = dados.get("empresa", "")
            formulario_disponivel.email = dados.get("email", "")
            formulario_disponivel.senha = dados.get("senha", "")
            formulario_disponivel.business_hours = dados.get("business_hours", "")
            formulario_disponivel.cor = dados.get("cor", "")
            formulario_disponivel.nicho = dados.get("nicho", "")
            formulario_disponivel.status = "respondido"
            formulario_disponivel.save()

            # Marca a fila como processada
            fila.processado = True
            fila.save()
            logger.info(f"Fila {fila.id} processada com sucesso")
        else:
            logger.warning(f"Não foi encontrado formulário reservado para a agencia {agencia}")
